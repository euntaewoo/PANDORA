import docx
from html2image import Html2Image
import html
import os
import sys
from PIL import Image
import re

def apply_semantic_line_breaks(text):
    """
    공식 규격서 4-2. 복합 기술 항목의 의미 단위 줄바꿈 규칙 (Semantic Line-Break Rule)
    긴 1열 항목명을 의미 단위로 분할하여 <br> 적용
    """
    if not text:
        return text
        
    replacements = {
        "사용기한 또는 개봉 후 사용기간": "사용기한 또는<br>개봉 후 사용기간",
        "사용기한 또는\n개봉 후 사용기간": "사용기한 또는<br>개봉 후 사용기간",
        "화장품제조업자 / 책임판매업자": "화장품제조업자 /<br>책임판매업자",
        "화장품제조업자/책임판매업자": "화장품제조업자 /<br>책임판매업자",
        "화장품제조업자 및 화장품책임판매업자": "화장품제조업자 및<br>화장품책임판매업자",
        "기능성 화장품 심사 필 유무": "기능성 화장품<br>심사 필 유무",
        "기능성화장품 심사 필 유무": "기능성 화장품<br>심사 필 유무",
        "사용할 때의 주의사항": "사용할 때의<br>주의사항",
        "소비자 상담 전화번호": "소비자 상담<br>전화번호",
        "소비자상담 관련 전화번호": "소비자상담 관련<br>전화번호"
    }
    
    # 텍스트 내에서 치환
    for key, val in replacements.items():
        if key in text:
            return val
            
    # 그 외 너무 긴 단어는 띄어쓰기 기준으로 중간에 <br> 삽입 시도
    if len(text) >= 11 and "<br>" not in text:
        words = text.split()
        if len(words) >= 2:
            mid = len(words) // 2
            return " ".join(words[:mid]) + "<br>" + " ".join(words[mid:])
            
    return text

def render_docx_to_png(doc_path, output_dir, png_filename="010_PDP_상품상세정보_테이블.png"):
    html_path = os.path.join(output_dir, "temp_formatted.html")
    full_png_path = os.path.join(output_dir, png_filename)
    
    doc = docx.Document(doc_path)
    
    html_content = '''<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="utf-8">
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
body {
    font-family: 'Pretendard', sans-serif;
    width: 860px;
    margin: 0;
    padding: 40px;
    box-sizing: border-box;
    background-color: #ffffff;
}
.spec-title {
    font-size: 64px;
    font-weight: 700;
    color: #222;
    text-align: center;
    margin-bottom: 40px;
    letter-spacing: -1px;
}
.spec-table {
    width: 100%;
    border-collapse: collapse;
    border-top: 3px solid #222;
    border-bottom: 3px solid #222;
}
.spec-table th, .spec-table td {
    padding: 24px 20px;
    border-bottom: 1px solid #e0e0e0;
    font-size: 32px;
    line-height: 1.5;
}
.spec-table th {
    width: 275px; /* 공식 가이드 기반 최적화 너비 */
    background-color: #f8f9fa;
    color: #333;
    font-weight: 700;
    text-align: left;
    vertical-align: middle;
}
.spec-table td {
    color: #555;
    vertical-align: middle;
    word-break: keep-all;
}
.spec-table tr:last-child th, .spec-table tr:last-child td {
    border-bottom: none;
}
/* 주의사항 번호 매기기 등은 깔끔하게 처리 */
.formatted-list {
    margin: 0;
    padding-left: 0;
    list-style-type: none;
}
.formatted-list li {
    margin-bottom: 8px;
}
</style>
</head>
<body id="capture-area">
<div class="spec-title">상품 상세 정보</div>
'''

    if doc.tables:
        table = doc.tables[0]
        html_content += '<table class="spec-table">\n'
        start_row = 1 if table.rows[0].cells[0].text.strip() == "항목" else 0
        for row in table.rows[start_row:]:
            cells = row.cells
            if len(cells) >= 2:
                # 항목(라벨) 처리
                raw_key = cells[0].text.strip()
                escaped_key = html.escape(raw_key)
                key_with_br = apply_semantic_line_breaks(escaped_key)
                
                # 값(내용) 처리 - 줄바꿈을 <br>로 변환, 리스트 형식 개선
                val = html.escape(cells[1].text.strip())
                # 1) 2) 등의 번호를 <br>과 함께 깔끔하게 마진 적용
                val = re.sub(r'([1-9]\))', r'<br>\1', val)
                val = re.sub(r'(㈎|㈏|㈐)', r'<br>&nbsp;&nbsp;\1', val)
                val = val.replace('\\n', '<br>').replace('<br><br>', '<br>').strip('<br>')
                
                html_content += f"<tr>\n<th>{key_with_br}</th>\n<td>{val}</td>\n</tr>\n"
        html_content += '</table>\n'

    html_content += '''
</body>
</html>
'''

    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)

    try:
        hti = Html2Image(size=(860, 4000))
        hti.output_path = output_dir
        hti.screenshot(html_file=html_path, save_as=png_filename)

        img = Image.open(full_png_path)
        bbox = img.convert("RGB").getbbox()
        if bbox:
            img_cropped = img.crop((0, 0, 860, bbox[3] + 40))
            img_cropped.save(full_png_path)
        print(f"Success! Saved to {full_png_path}")
    finally:
        if os.path.exists(html_path):
            os.remove(html_path)

if __name__ == "__main__":
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
        output_dir = os.path.dirname(doc_path)
        render_docx_to_png(doc_path, output_dir)
    else:
        doc_path = r"D:\Users\euntaewoo\Desktop\이미지번역워크스페이스\로지컬리스킨_변역대상\05_Multi Corrective Eye cream\05. 멀티코렉티브 아이크림 -Multi Corrective 일렉트 -Multi Corrective Eye cream\010_PDP_상품상세정보.docx"
        # path fixing due to typo in fallback
        doc_path = r"D:\Users\euntaewoo\Desktop\이미지번역워크스페이스\로지컬리스킨_변역대상\05_Multi Corrective Eye cream\05. 멀티코렉티브 아이크림 -Multi Corrective Eye cream\010_PDP_상품상세정보.docx"
        output_dir = r"D:\Users\euntaewoo\Desktop\이미지번역워크스페이스\로지컬리스킨_변역대상\05_Multi Corrective Eye cream\05. 멀티코렉티브 아이크림 -Multi Corrective Eye cream"
        render_docx_to_png(doc_path, output_dir)
