import os
def get_recursive_files(base_dir):
    try:
        res = []
        for root, _, files in os.walk(base_dir):
            for f in files:
                res.append(os.path.relpath(os.path.join(root, f), base_dir))
        return res
    except Exception:
        return []

r"""
===================================================================================
🌐 multilingual_text_in_image_translatio_agy_sdk.py
-----------------------------------------------------------------------------------
• Purpose: Antigravity SDK & Vertex AI Async Transcreation Engine (Full Specification Integrated)
• Location: C:\Users\euntaewoo\Desktop\multilingual_text_in_image_translatio_agy_sdk\multilingual_text_in_image_translatio_agy_sdk.py
• Features:
    1. 단일 공통 인풋 폴더(01_번역대상_원본) 기준 구동 (서브폴더/낱개 파일 무한 Depth 재귀 탐색 지원)
    2. Google Cloud Vertex AI 비동기(client.aio.models.generate_content) 전용 엔드포인트 100% 완전 통합
    3. 실행 시 도착 언어(EN, JP, CN, TW, ALL) 대화형 및 CLI 인자(--lang, --product_name) 지원
    4. 도착어별 규정/법률(영어 초월번역 MoCRA, 일본 약기법 56종, 중국 신광고법/NMPA, 대만 TFDA) 자동 적용
    5. 표(고시표/성분표) 자동 감지 시 HTML 표준 헤드리스 Chromium 렌더러로 고선명 860px 분기 처리
    6. 상품 혼재 방지: [최초 번역대상 상품명]_[번역국가언어] 전용 서브폴더 자동 생성 및 저장
    7. 초월번역 품질 자동 평가(Transcreation QA Evaluator) 및 원클릭 HTML 시각 뷰어 리포트 자동 발행
• Models:
    - Pass 1: gemini-3.1-pro-preview (추론, 번역, 법률 필터링, 대용량 토큰 8192)
    - Pass 2: gemini-3.1-flash-image (시각적 신경망 인페인팅 렌더링, 1:1 해상도 복원)
• API Standard: Google Cloud Vertex AI (location="global", Async Mode)
===================================================================================
"""

import io
import json
import os
import re
import sys
import time
import argparse
import asyncio
import subprocess
from typing import Any, Dict, List, Optional, Tuple

from google import genai
from google.genai import types
from PIL import Image

sys.stdout.reconfigure(encoding="utf-8")

# =================================================================================
# 1. 시스템 기본 경로 및 인증 초기화 (서브폴더 구조 지원)
# =================================================================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if os.path.basename(SCRIPT_DIR) == "multilingual_text_in_image_translatio_agy_sdk_core":
    PROJECT_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, ".."))
else:
    PROJECT_ROOT = SCRIPT_DIR

DEFAULT_INPUT_DIR = os.path.join(PROJECT_ROOT, "01_번역대상_원본")
DEFAULT_OUTPUT_BASE = os.path.join(PROJECT_ROOT, "02_번역결과_최종")
DEFAULT_REMEDIATION_BASE = os.path.join(PROJECT_ROOT, "04_번역교정")

MODEL_PRO = "gemini-3.1-pro-preview"
MODEL_FLASH_IMAGE = "gemini-3.1-flash-image"

# =================================================================================
# 0. 전역 글로벌 컴플라이언스(법무) & 럭셔리 초월번역 시스템 인스트럭션
# =================================================================================
GLOBAL_COMPLIANCE_SYSTEM_INSTRUCTION = """[SYSTEM INSTRUCTION: Global Cross-Border E-Commerce Compliance & Prestige Beauty Transcreation Expert]
당신은 미국 FDA(MoCRA), 일본 후생노동성(약기법), 중국 NMPA/신광고법, 대만 TFDA 규정을 완벽히 준수하는 15년 차 글로벌 뷰티 법무 감사관이자, 세포라(Sephora)·백화점 럭셔리 브랜드의 수석 카피라이터입니다.

[엄격 실행 4대 대원칙 (STRICT EXECUTION RULES)]
1. [컴플라이언스 & 의약품 오인 원천 차단 (First Principles Heuristic)]:
   - 인체 구조, 생리적 기능, 세포(Cell/Cellular) 단위의 생화학적 변화나 치료·재생을 암시하는 클레임(예: cellular vitality, cellular resilience, cell metabolism, collagen synthesis)을 100% 원천 차단하십시오.
   - 사전에 등록되지 않은 신규 성분/어휘라도 세포/생리기능 직접 관여 뉘앙스가 있다면 무조건 '피부 표면의 미용적 외관 개선(-looking, appearance of, natural moisture barrier)'으로 안전하게 우회하십시오.
2. [타겟 권역별 문화적 어댑테이션 및 콩글리시 배제]:
   - 영미권(EN): 콩글리시('Complex skin issues' -> 'Multiple skin concerns', 'Troubled skin' -> 'Blemish-prone skin') 전면 배제 및 직관적 뷰티 표준어 적용. 노화 서술 시 'combats the signs of premature aging'으로 징후(signs) 한정.
   - 일본(JP): 후생노동성 56종 허용 효능(Positive List) 엄격 준수 ('肌を整える', 'うるおいを与える' 등).
   - 중화권(CN/TW): 신광고법 8대 절대화 금지어('最', '第一', '顶级' 등) 배제 및 NMPA/TFDA 화장품 규정 준수.
3. [디자인 & 레이아웃 최적화]:
   - 이미지 레이아웃에 텍스트가 위화감 없이 안착하도록 글자 수 길이를 최적화하고, 백화점 럭셔리 브랜드 수준의 세련된 어휘로 초월번역을 수행하십시오.
4. [고시정보 및 법정 필수 조항]:
   - 고객센터 전화번호(+82 국제번호 통일), 기능성화장품 심사 상태, 주의사항 3대 조항을 정확히 표준화하십시오.
"""


# 언어별 설정 매핑
LANG_CONFIGS = {
    "EN": {
        "name": "영어 (English - Amazon/Shopee US)",
        "folder_name": "영어",
        "code": "EN",
        "tag": "_EN_Translated.png"
    },
    "JP": {
        "name": "일본어 (日本語 - Qoo10 Japan / 56종 약기법)",
        "folder_name": "일본어",
        "code": "JP",
        "tag": "_JP_Translated.png"
    },
    "CN": {
        "name": "중국어 간체 (简体中文 - 중국 본토 신광고법)",
        "folder_name": "중국어_간체",
        "code": "CN",
        "tag": "_CN_Simp_Translated.png"
    },
    "TW": {
        "name": "중국어 번체 (繁體中文 - 대만/홍콩 TFDA)",
        "folder_name": "중국어_번체",
        "code": "TW",
        "tag": "_TW_Trad_Translated.png"
    },
    "KR": {
        "name": "한국어 (한국 - 네이버/쿠팡 최적화)",
        "folder_name": "한국어",
        "code": "KR",
        "tag": "_KR_Translated.png"
    }
}


def load_credentials() -> genai.Client:
    """
    ⛔ [HARD STOP — global_rules.md §5 강제]
    에이전트가 작성하는 모든 스크립트에서 genai.Client(vertexai=True, ...) 직접 작성 절대 금지.
    반드시 이 load_credentials() 함수를 import하여 사용할 것.
    이 함수가 location="global" 및 인증 경로를 공식 가이드 규격으로 자동 보장합니다.
    Vertex AI 서비스 계정 키 및 API 키를 탐색하여 genai.Client를 초기화합니다.
    """
    env_paths = [
        os.path.join(PROJECT_ROOT, ".env"),
        os.path.join(PROJECT_ROOT, "영어", ".env"),
        os.path.join(PROJECT_ROOT, "일본어", ".env"),
        os.path.join(SCRIPT_DIR, ".env"),
    ]
    api_key = os.environ.get("GEMINI_API_KEY")
    gcp_json_key = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")

    for p in env_paths:
        if os.path.exists(p):
            with open(p, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line.startswith("#"):
                        continue
                    if line.startswith("GEMINI_API_KEY=") and not api_key:
                        api_key = line.split("=", 1)[1].strip()
                    elif line.startswith("GOOGLE_APPLICATION_CREDENTIALS=") and not gcp_json_key:
                        gcp_json_key = line.split("=", 1)[1].strip().strip('"').strip("'")

    key_candidates = [
        gcp_json_key,
        os.path.join(PROJECT_ROOT, "00_공통자료", "APIs_KEY", "인증키_및_계정", "김차장_vertex api_key", "vertex_ai_auth_key.json"),
        os.path.join(PROJECT_ROOT, "00_공통자료", "인증키_및_계정", "김차장_vertex api_key", "vertex_ai_auth_key.json"),
        os.path.join(PROJECT_ROOT, "일본어", "vertex_service_account.json"),
    ]

    for kpath in key_candidates:
        if kpath and os.path.exists(kpath) and kpath.endswith(".json"):
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = kpath
            with open(kpath, "r", encoding="utf-8") as f:
                key_data = json.load(f)
                project_id = key_data.get("project_id")
            client = genai.Client(vertexai=True, project=project_id, location="global")
            print(f"[AUTH SUCCESS] Vertex AI Client 연결 완료 (Project: {project_id}, Location: global)", flush=True)
            return client

    if api_key:
        if api_key.startswith("AQ."):
            print("[AUTH] Agent Platform API 키 감지 -> Vertex AI 모드로 전환", flush=True)
            return genai.Client(vertexai=True, api_key=api_key)
        print("[AUTH] Gemini API 키 모드로 연결", flush=True)
        return genai.Client(api_key=api_key)

    print("[ERROR] GEMINI_API_KEY 또는 GOOGLE_APPLICATION_CREDENTIALS가 설정되지 않았습니다.", flush=True)
    sys.exit(1)


# =================================================================================
# 2. 일본 약기법 56종 목록 로드
# =================================================================================

# =================================================================================
# 2-0. [DYNAMIC-COMPLIANCE-LEXICON-LOADER] 국가별 표준 렉시콘 JSON 동적 로더
# =================================================================================
def load_dynamic_compliance_lexicon(lang_code: str) -> Tuple[Dict[str, str], str]:
    """
    00_공통자료/compliance_lexicons 디렉터리에서 해당 국가 표준 렉시콘 JSON을 실시간 로드하여
    (1) Python 정규식 강제 치환 딕셔너리, (2) Pass 1 프롬프트 주입문으로 변환합니다.
    """
    lex_dir_candidates = [
        os.path.join(PROJECT_ROOT, "00_공통자료", "compliance_lexicons"),
        os.path.join(SCRIPT_DIR, "..", "00_공통자료", "compliance_lexicons"),
        os.path.join(SCRIPT_DIR, "compliance_lexicons")
    ]
    lex_dir = None
    for cand in lex_dir_candidates:
        if os.path.exists(cand):
            lex_dir = cand
            break
            
    mapping = {
        "EN": "en_fda_mocra_lexicon.json",
        "JP": "jp_pmda_pharm_lexicon.json",
        "CN": "cn_nmpa_adlaw_lexicon.json",
        "TW": "tw_tfda_lexicon.json",
    }
    fname = mapping.get(lang_code.upper(), "en_fda_mocra_lexicon.json")
    fpath = os.path.join(lex_dir, fname) if lex_dir else None
    
    replacements = {}
    prompt_lines = []
    
    if fpath and os.path.exists(fpath):
        try:
            with open(fpath, "r", encoding="utf-8") as f:
                lex_data = json.load(f)
                cats = lex_data.get("categories", {})
                for c_key, c_val in cats.items():
                    desc = c_val.get("description", "")
                    banned_list = c_val.get("banned_terms", [])
                    if banned_list:
                        prompt_lines.append(f"### [COMPLIANCE LEXICON: {desc}]")
                        for item in banned_list:
                            b = item.get("banned", "").strip()
                            p = item.get("preferred", "").strip()
                            r = item.get("reason", "").strip()
                            if b and p:
                                replacements[b] = p
                                prompt_lines.append(f"- BANNED: `{b}` -> MUST USE: `{p}` ({r})")
        except Exception as e:
            print(f"  ⚠️ [WARN] 렉시콘 로드 실패 ({fpath}): {e}")
            
    prompt_str = "\n".join(prompt_lines)
    return replacements, prompt_str

def load_jp_efficacy_list() -> str:
    efficacy_paths = [
        os.path.join(PROJECT_ROOT, "일본어", "cosmetics_efficacy_56.json"),
        os.path.join(PROJECT_ROOT, "cosmetics_efficacy_56.json")
    ]
    for ep in efficacy_paths:
        if os.path.exists(ep):
            try:
                with open(ep, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    return "\n".join([f"{it['id']}. {it['claim_jp']} ({it['claim_ko']})" for it in data])
            except Exception:
                pass
    return "1. 肌を整える\n2. 肌荒れを防ぐ\n3. 皮膚にうるおいを与える 등 56종"


# =================================================================================
# 2-1. [QA-FEEDBACK-INJECTION-LOCK] 03_번역품질평가 진단 결과 자동 로더 & 결정론적 보정 게이트
# =================================================================================
def load_qa_feedback_and_transcreation_rules(source_folder: str, product_name: str = "", lang_code: str = "EN") -> Dict[str, Any]:
    """
    03_번역품질평가 진단 결과(Transcreation_QA_Report.json 및 correction_feedbacks / transcreation_comparisons)를
    자동 탐색·로드하여 Pass 1 프롬프트 주입문 및 결정론적(Deterministic) 단어/구문 치환 딕셔너리로 변환합니다.
    """
    qa_data = {}
    found_path = None

    candidates = [
        os.path.join(source_folder, "transcreation_guide.json"),
        os.path.join(source_folder, "Transcreation_QA_Report.json"),
        os.path.join(source_folder, "qa_report.json"),
    ]
    if product_name:
        candidates.extend([
            os.path.join(PROJECT_ROOT, "03_번역품질평가", "02_진단결과", product_name, "Transcreation_QA_Report.json"),
            os.path.join(PROJECT_ROOT, "03_번역품질평가", "02_진단결과", f"{product_name}_{lang_code}", "Transcreation_QA_Report.json"),
            os.path.join(PROJECT_ROOT, "03_번역품질평가", "02_진단결과", f"{product_name}_EN", "Transcreation_QA_Report.json"),
        ])

    qa_results_dir = os.path.join(PROJECT_ROOT, "03_번역품질평가", "02_진단결과")
    if os.path.exists(qa_results_dir):
        for sub in os.listdir(qa_results_dir):
            sub_p = os.path.join(qa_results_dir, sub, "Transcreation_QA_Report.json")
            if os.path.exists(sub_p) and sub_p not in candidates:
                candidates.append(sub_p)

    for c_path in candidates:
        if os.path.exists(c_path):
            try:
                with open(c_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, dict) and (data.get("correction_feedbacks") or data.get("transcreation_comparisons")):
                        qa_data = data
                        found_path = c_path
                        break
            except Exception:
                pass

    prompt_lines = []
    spelling_dict = {}
    phrase_dict = {}

    known_typos = {
        r"\benurgy\b": "energy",
        r"\bdeley\b": "delay",
        r"\bocne\b": "acne",
        r"\bmetabailism\b": "metabolism",
        r"\bLIGHTWEGHT\b": "LIGHTWEIGHT",
        r"\bCosmetis\b": "Cosmetics",
        r"\bPynidoxine\b": "Pyridoxine",
    }
    for k, v in known_typos.items():
        spelling_dict[k] = v

    if qa_data:
        print(f"  🎯 [QA 피드백 연동 성공] {os.path.relpath(found_path, PROJECT_ROOT)} 에서 교정 규칙 100% 로드 완료", flush=True)

        feedbacks = qa_data.get("correction_feedbacks", [])
        if feedbacks:
            prompt_lines.append("### A. [MANDATORY SPELLING & COMPLIANCE CORRECTIONS (STRICT OVERRIDE)]")
            for fb in feedbacks:
                orig = fb.get("original", "").strip()
                cur = fb.get("current_translation", "").strip()
                rec = fb.get("recommended_correction", "").strip()
                reason = fb.get("reason", "").strip()
                if rec:
                    prompt_lines.append(f"- **Issue**: `{orig or cur}` ➔ **MUST BE CORRECTED TO**: `{rec}` (Reason: {reason})")
                    if cur and rec:
                        phrase_dict[cur] = rec
                    match = re.search(r'([a-zA-Z]+)\s*(?:->|➔|to)\s*([a-zA-Z]+)', reason)
                    if match:
                        typo_w, correct_w = match.group(1), match.group(2)
                        spelling_dict[rf"\b{typo_w}\b"] = correct_w

        comparisons = qa_data.get("transcreation_comparisons", [])
        if comparisons:
            prompt_lines.append("\n### B. [GOLD-STANDARD TRANSCREATION PAIRS (Pre-Approved Sephora-Grade Copy)]")
            for comp in comparisons:
                orig_ko = comp.get("original", "").strip()
                lit = comp.get("literal_translation", "").strip()
                trans = comp.get("transcreation", "").strip()
                val = comp.get("value_analysis", "").strip()
                if orig_ko and trans:
                    prompt_lines.append(f"- **Source (KO)**: `{orig_ko}`")
                    prompt_lines.append(f"  ➔ **Approved Transcreation**: `{trans}`")
                    if val:
                        prompt_lines.append(f"  ➔ **Key Reason**: {val}")
                    phrase_dict[orig_ko] = trans
                    if lit:
                        phrase_dict[lit] = trans

    override_text = "\n".join(prompt_lines)
    return {
        "found_path": found_path,
        "raw_data": qa_data,
        "prompt_override_text": override_text,
        "spelling_dict": spelling_dict,
        "phrase_dict": phrase_dict
    }

def apply_deterministic_qa_overrides(t_map: List[Dict[str, Any]], qa_rules: Dict[str, Any], lang_code: str) -> List[Dict[str, Any]]:
    """
    [DETERMINISTIC-QA-OVERRIDE-GATE]
    Pass 1 LLM 응답 후, QA 진단 결과에서 도출된 오타/MoCRA 금지어/초월번역 쌍을
    Python 정규식 및 치환 규칙으로 100% 강제 검증 및 보정합니다.
    """
    if not isinstance(t_map, list):
        return t_map

    spelling_dict = qa_rules.get("spelling_dict", {})
    phrase_dict = qa_rules.get("phrase_dict", {})

    # 동적 렉시콘 JSON에서 실시간 치환 규칙 병합
    dynamic_replacements, _ = load_dynamic_compliance_lexicon(lang_code)
    
    mocra_banned_replacements = {
        r"\bPrescribe\b": "Targeted Solution for",
        r"\bprescribe\b": "targeted solution for",
        r"\bBio-Immunity\b": "Skin Defense",
        r"\bbio-immunity\b": "skin defense",
        r"\bfed directly\b": "infused daily",
        r"\bKyel-Tan-Tone\b": "Texture, Elasticity & Luminosity",
        r"\bkyel-tan-tone\b": "texture, elasticity & luminosity",
        # 5대 법적 리스크 및 콩글리시 완벽 강제 치환
        r"\bComplex skin issues\b": "Multiple skin concerns",
        r"\bcomplex skin issues\b": "multiple skin concerns",
        r"\bTroubled skin\b": "Blemish-prone skin",
        r"\btroubled skin\b": "blemish-prone skin",
        r"\bnutrients for cellular vitality\b": "hydration for a resilient-looking complexion",
        r"\bcellular vitality\b": "resilient-looking complexion",
        r"\breinforces cellular resilience\b": "reinforces the skin's natural moisture barrier",
        r"\bcellular resilience\b": "skin's natural moisture barrier",
        r"\bcombats premature aging\b": "combats the signs of premature aging",
        r"\bcombats aging\b": "combats the signs of aging",
        r"\bcellular metabolism\b": "natural skin vitality",
    }
    for dyn_b, dyn_p in dynamic_replacements.items():
        dyn_pat = rf"\b{re.escape(dyn_b)}\b"
        if dyn_pat not in mocra_banned_replacements:
            mocra_banned_replacements[dyn_pat] = dyn_p

    updated_count = 0
    for item in t_map:
        if not isinstance(item, dict):
            continue
        kor = item.get("kor", "")
        target = item.get("target_text", "")
        if not target:
            continue

        orig_target = target

        # 1. 구문 단위 정확 매칭 (Approved Transcreation Override)
        for p_k, p_v in phrase_dict.items():
            if p_k and (p_k in kor or p_k in target or p_k.lower() in target.lower()):
                if len(p_k) > 8 and target != p_v:
                    target = p_v
                    break

        # 2. MoCRA 금지어 치환
        if lang_code == "EN":
            for pat, repl in mocra_banned_replacements.items():
                target = re.sub(pat, repl, target)

        # 3. 오타 정규식 치환 (대소문자 무관 단어 단위)
        for typo_pat, correct_word in spelling_dict.items():
            target = re.sub(typo_pat, correct_word, target, flags=re.IGNORECASE)

        # 4. 고객센터 전화번호 국제 규격 통일 (+82)
        if any(k in kor.lower() or k in target.lower() for k in ["customer", "contact", "phone", "상담", "문의"]):
            if "02-" in target or "02)" in target or "6743-3206" in target:
                target = re.sub(r"(?:02|\+82-2)[-.)\s]*6743[-.]?3206", "+82-2-6743-3206", target)

        if target != orig_target:
            item["target_text"] = target
            item["qa_overridden"] = True
            updated_count += 1
            print(f"     ⚡ [QA 규격 자동 보정] `{orig_target[:35]}` ➔ `{target[:35]}`", flush=True)

    if updated_count > 0:
        print(f"  🛡️ [DETERMINISTIC GATE] 총 {updated_count}개 텍스트 블록에 대해 1단계 QA 진단 교정안 100% 강제 반영 완료")

    return t_map


# =================================================================================
# 3. 언어별 프롬프트 생성기 (Pass 1 & Pass 2 - Luxury Beauty Transcreation Architecture)
# =================================================================================
def build_prompts(lang_code: str, qa_override_prompt: str = "") -> Tuple[str, str]:
    qa_block = ""
    if qa_override_prompt and qa_override_prompt.strip():
        qa_block = f"""
## 0. MANDATORY QA DIAGNOSIS OVERRIDES & PRE-APPROVED TRANSCREATION (HIGHEST PRIORITY)
The following issues were identified during Step 1 Transcreation QA Diagnosis and MUST BE 100% ENFORCED.
Any violation or misspelling from the list below will cause total rejection:
{qa_override_prompt}
"""

    if lang_code == "EN":
        pass1 = f"""
[SYSTEM PROMPT] Global Luxury Beauty Transcreation & Compliance Automator (English Mode)
{qa_block}
## 1. System Role & Persona
You are a Senior Creative Director and Elite Copywriter with 10+ years of experience specializing in localizing global high-end cosmetic brands (e.g., Estée Lauder, Lancôme, Sisley, SK-II) for US and global luxury beauty markets.
Your mission is NOT literal translation. You must perform 'Transcreation'—rewriting the source text into a sophisticated, natural, and persuasive marketing copy that aligns perfectly with luxury Sephora and prestige department store consumer psychology and advertising regulations.

## 2. Core Transcreation Principles & Mandatory Glossary (Strict Compliance)
### A. 고유 명사 및 브랜드명 필수 표기 지침 (Brand & Ingredient Glossary)
- **Brand Name**: Maintain proprietary brand name **`Logicall Skin`** in original English. Never transliterate.
- **Key Ingredients**: Maintain **`Aquatide`**, **`Aquatide 5000`**, **`LiftDerm`**, **`Lifting Logic for eye`** in original English.
- **Product Title**: Standardized luxury Sephora US product title format (e.g. `Logicall Skin Aquatide Resurface Serum`).
### B. Eliminate Translationese & 1:1 Matching
- Never translate source adverbs literally (e.g., Do not translate '확실히', '진짜', '정말' into stiff equivalents like 'Definitely', 'Truly', 'Really', 'Certainly').
- Capture the underlying scientific efficacy or emotional benefit, and recreate it using active, premium verbs native to the luxury beauty market.
### C. Natural Sentence Flow & Syntactic Restructuring
- Ensure seamless syntactic connectivity. If a sentence mentions active ingredients or percentages (e.g., "10% LiftDerm" or "Aquatide"), restructure the sentence gracefully so that it flows naturally into the product name or efficacy claim without sounding fragmented or awkward.
### D. Use Premium Beauty & Biotech Terminology
- Skin deeper layers: Deep within the skin layers / Dermal matrix
- Multi-corrective / Repair: Multi-Corrective Repair / Advanced Total Revitalizing Care / Resurface
- Firming / Elasticity: Rebuilding skin elasticity / Restoring visible firmness & bounce
- Fine lines / Wrinkles: Fine lines and wrinkles / Micro-creases
- Active ingredients: Maintain proprietary names like 'Aquatide', 'LiftDerm', 'Lifting Logic for eye' in original English.

## 3. Global Cosmetic Regulatory Screening (Mandatory Guardrails)
### A. Ban on Absolute & Unverifiable Claims
- Do not use absolute expressions such as "World's First", "No.1", "Best", or "The Ultimate".
- Rephrase them into compliant luxury terms of innovation and advanced care (e.g., `Innovative formula engineered for delicate eye areas`, `Advanced Multi-Corrective Solution`, `Targeted Precision Care`).
### B. Ban on Medical/Clinical Misinterpretation & 4 Compliance Safe Verbs
- Do not use phrases implying permanent wrinkle deletion or medical procedures (e.g., Botox-like, Filler-like effects).
- Strictly enforce the 4 compliance-safe verbs: **`Smooth`**, **`Diminish`**, **`Alleviate`**, **`Care / Repair`**.

## 4. 이미지 텍스트 전수 추출 & 출력 포맷
이미지 내의 모든 한국어 텍스트는 단 하나도 빠짐없이 100% 추출하십시오. (고시표 테이블인 경우 is_table: true 설정)
출력은 반드시 순수 JSON이어야 합니다:
{{
  "is_table": false,
  "translation_map": [
    {{
      "kor": "한국어 원문",
      "target_text": "세포라/백화점 톤앤매너 최고급 영문 카피",
      "reasoning": "절대표현 순화 및 럭셔리 초월번역 근거"
    }}
  ]
}}
"""
        pass2_tmpl = """
당신은 정밀한 시각적 로컬라이제이션을 수행하는 이미지 인페인팅 AI입니다.
첨부된 원본 이미지의 배경, 텍스처, 제품 누끼, 색상 톤을 1픽셀의 왜곡 없이 보존하세요.
아래 [번역 매핑 데이터 JSON]을 바탕으로 단일 이미지를 생성하세요.

[엄격 렌더링 규칙 - 프리미엄 영문 럭셔리 뷰티]
1. (KOR ERASING) 원본 한국어 텍스트는 배경색/텍스처로 완벽히 덮어써서 100% 제거할 것.
2. (JSON APPLY) 지워진 그 자리에 오직 [번역 매핑 데이터 JSON]의 'target_text'만 렌더링할 것.
3. (FONT DIRECTIVE) 세련된 프리미엄 산세리프(Montserrat 100% 단일 서체)로 정갈하고 모던하게 렌더링할 것.
4. (FULL REGENERATION) 패칭(덧칠)하지 말고 전체 캔버스를 완벽하게 새로 렌더링할 것.
5. (PACKAGE PRESERVATION) 제품 본품(용기 표면)의 영문 및 로고는 100% 완벽 보존할 것.

[번역 매핑 데이터 JSON]
{json_data}
"""
    elif lang_code == "JP":
        efficacy_str = load_jp_efficacy_list()
        pass1 = f"""
[SYSTEM PROMPT] Global Luxury Beauty Transcreation & Compliance Automator (Japanese Mode)

## 1. 시스템 역할 및 콘셉트 (Role & Context)
당신은 시슬리, SK-II, 데코르테 등 일본 하이엔드 프레스티지 뷰티 시장을 총괄하는 10년 차 수석 크리에이티브 디렉터이자 @cosme 전문 엘리트 카피라이터입니다.
일본 소비자의 감성을 깊게 자극하는 정중하고 품격 있는 뷰티 카피(美肌, ハリ, 潤い)로 초월번역(Transcreation)하세요.

## 2. 초월번역 핵심 원칙 및 필수 용어집 (Core Principles & Glossary)
### A. 고유 명사 및 브랜드명 필수 표기 지침 (Brand & Glossary)
- **브랜드명**: 고유 영문 명칭인 **`Logicall Skin`**을 그대로 유지하십시오.
- **핵심 성분명**: **`Aquatide`**, **`Aquatide 5000`**, **`LiftDerm`** 등 독자 성분명은 영문 고유 표기를 유지하여 임상적 신뢰도를 극대화하십시오.
- **제품명**: `Logicall Skin アクアタイド リサーフェス セラム` 또는 영문 병기 표기.
### B. 번역투 및 직역 부사 전면 금지 (Eliminate Translationese)
- '確実に', '本当に', '絶対に' 등 딱딱한 부사 직역을 전면 금지하고, 피부 감촉과 효능을 섬세하게 묘사하는 프리미엄 어휘로 재창조하십시오.
### C. 자연스러운 구문 결속 및 제형 감성 묘사 (Natural Sentence Flow)
- "10% LiftDerm"이나 "Aquatide" 등 성분 비율이 문장 중간에 어색하게 끊기지 않고 매끄러운 뷰티 서사로 이어지도록 구조를 재조정하십시오.
### D. 4대 기능성 뷰티 전문 어휘 사전 채택
- 피부 속/기저층: 肌の奥・角質層のすみずみまで
- 토탈 케어/멀티 코렉티브: 高機能トータルリペア / 多機能エイジングケア
- 탄력 복원/강화: ハリ・弾力を呼び覚ます / 弾むようなハリ感
- 눈가 잔주름/건조주름: 目元の小ジワ・乾燥ジワ
- 독자 성분명 영문 보존: 'LiftDerm', 'Aquatide', 'Lifting Logic for eye' 등은 억지로 가타카나로 뭉개지 않고 영문 고유 표기를 유지하여 임상적 신뢰도를 극대화하십시오.

## 3. 후생노동성 약기법(약사법) 규제 준수 (Regulatory Compliance)
### A. 절대적/과대 표현 전면 금지 (Ban on Absolute Claims)
- '世界初', 'No.1', '最高', '究極' 등 검증 불가능한 절대 표현 사용을 금지하고, '目元のために開発された先進テクノロジー', '高機能トータルケア' 등 프리미엄 케어 표현으로 순화하십시오.
### B. 의료 시술 오인 금지 및 약기법 안전 표현 (Compliance-Safe Verbs)
[일본 후생노동성 공인 56종 허용 효능 목록 엄격 준수]
{efficacy_str}
- '치료/재생/소염/보톡스/필러 효과' 등 의료 행위 및 성형 시술 연상 표현 절대 금지 -> **`肌を整える`**, **`乾燥による小ジワを目立たなくする`**, **`ハリを与える`**, **`うるおいを与える`** 등 포지티브 리스트 표현으로 순화.
- '무자극' -> '低刺激処方', '미백' -> 'うるおいを与え、透明感のある肌へ'.

## 4. 이미지 텍스트 전수 추출 & 출력 포맷
이미지 내의 모든 한국어 텍스트는 단 하나도 빠짐없이 100% 추출하십시오. (고시표인 경우 is_table: true 설정)
출력은 반드시 순수 JSON이어야 합니다:
{{
  "is_table": false,
  "translation_map": [
    {{
      "kor": "한국어 원문",
      "target_text": "일본 프레스티지 뷰티 톤앤매너 번역문",
      "reasoning": "약기법 검열 및 감성 초월번역 사유"
    }}
  ]
}}
"""
        pass2_tmpl = """
당신은 정밀한 시각적 로컬라이제이션을 수행하는 이미지 인페인팅 AI입니다.
첨부된 원본 이미지의 디자인 레이아웃과 제품을 그대로 유지하세요.
아래 [번역 매핑 데이터 JSON]을 바탕으로 단일 이미지를 생성하세요.

[엄격 렌더링 규칙 - 일본 럭셔리 뷰티]
1. (KOR ERASING) 원본 한국어 텍스트는 배경으로 덮어써서 100% 지울 것.
2. (JSON APPLY) 지워진 자리에 오직 [번역 매핑 데이터 JSON]의 'target_text'만 렌더링할 것.
3. (FONT DIRECTIVE) 일본 최고급 표준 서체(Noto Sans JP / Gothic) 스타일로 정갈하게 렌더링할 것.
4. (FULL REGENERATION) 전체 캔버스를 결점 없이 완벽히 새로 렌더링할 것.
5. (PACKAGE PRESERVATION) 제품 본품(용기)의 영문 및 로고는 100% 보존할 것.

[번역 매핑 데이터 JSON]
{json_data}
"""
    elif lang_code == "CN":
        pass1 = f"""
[SYSTEM PROMPT] Global Luxury Beauty Transcreation & Compliance Automator (China Simplified Mode)
{qa_block}
## 1. System Role & Persona
당신은 에스티로더, 랑콤, 헬레나 루빈스타인 등 중국 본토 하이엔드 럭셔리 뷰티 시장을 총괄하는 10년 차 수석 크리에이티브 디렉터이자 샤오홍슈/티몰 럭셔리 전문 엘리트 카피라이터입니다.
단순 직역을 배제하고, 지적이고 고급스러운 하이테크 바이오 뷰티 서사로 초월번역(Transcreation)하세요.

## 2. Core Transcreation Principles & Mandatory Glossary (Strict Compliance)
### A. 고유 명사 및 브랜드명 필수 표기 지침 (용어집/Glossary)
- **브랜드명**: 고유 영문 명칭인 **`Logicall Skin`**을 그대로 유지하십시오.
- **Aquatide & Serum 표기**: **Aquatide 와 Serum은 중국어 고유명칭과 함께 반드시 영문을 병기**하십시오:
  * Aquatide -> **`阿夸肽 (Aquatide)`**
  * Serum -> **`精华液 (Serum)`**
  * 전체 제품명 -> **`阿夸肽修护精华液 (Aquatide Resurface Serum)`** 또는 문맥에 맞춘 영문 병기
### B. Eliminate Translationese & 1:1 Matching
- '确实', '真正', '非常', '绝对' 등 딱딱한 부사 직역을 전면 금지하고, 프리미엄 뷰티 전문 어휘로 세련되게 재창조하십시오.
### C. Natural Sentence Flow & Syntactic Restructuring
- "10% LiftDerm"이나 "Aquatide" 등 활성 성분 비율/성분명이 문맥과 끊기지 않고 제품 효능 및 서사로 매끄럽게 연결되도록 문장 구조를 재조정하십시오.
### D. Use Premium Beauty & Biotech Terminology
- 피부 속/기저층: 肌底深处 / 充盈肌底
- 토탈 케어/멀티 코렉티브/리서페이스: 多效修护 / 焕活肌底 / 抚平粗糙
- 탄력 복원/강화: 赋活肌底弹力 / 提升紧实度
- 눈가 및 피부 잔주름/건조주름: 细纹・干纹 / 抚平细纹
- 독자 성분명 영문 보존: 'LiftDerm', 'Aquatide' 등 글로벌 독자 성분명은 영문 또는 상기 영문 병기 규격을 엄수하십시오.

## 3. Global Cosmetic Regulatory Screening & NMPA Compliance (Mandatory Guardrails)
### A. Ban on Absolute & Unverifiable Claims (8대 절대화 금지어 전면 차단)
- 'World's First', 'No.1', 'Best', 'The Ultimate' 등 검증 불가능한 절대 표현(`全球首创`, `第一`, `最`, `顶级`, `极品`, `终极对策`) 전면 금지.
- 반드시 `卓越`, `优异`, `专为修护研发的创新科技`, `高端多效`, `精准修护` 등으로 순화하십시오.
### B. Ban on Medical/Clinical Misinterpretation & 4 Compliance Safe Verbs
- '주름 박멸', '영구 삭제', '보톡스/필러 효과', '치료(治疗)', '소염(消炎)', '재생(修复疤痕)' 등 의료 시술 오인 단어 전면 배제.
- 반드시 화장품 규정 내 안전 동사인 **`抚平` (Smooth), `淡化` (Diminish), `舒缓` (Alleviate), `修护` (Care/Repair)**만을 사용하여 표현하십시오.
- 순수 간체자(Simplified Chinese, zh-CN)만 사용할 것.

## 4. 이미지 텍스트 전수 추출 & 출력 포맷
이미지 내의 모든 한국어 텍스트는 단 하나도 빠짐없이 100% 추출하십시오. (고시표인 경우 is_table: true 설정)
출력은 반드시 순수 JSON이어야 합니다:
{{
  "is_table": false,
  "translation_map": [
    {{
      "kor": "한국어 원문",
      "target_text": "중국 신광고법 및 용어집 준수 럭셔리 간체 카피",
      "reasoning": "광고법 순화, 용어집 반영 및 럭셔리 초월번역 사유"
    }}
  ]
}}
"""
        pass2_tmpl = """
당신은 정밀한 시각적 로컬라이제이션을 수행하는 이미지 인페인팅 AI입니다.
첨부된 원본 이미지 레이아웃을 보존하며 아래 [번역 매핑 데이터 JSON]으로 단일 이미지를 생성하세요.

[엄격 렌더링 규칙 - 중국 본토 럭셔리 뷰티 간체]
1. (KOR ERASING) 원본 한국어 텍스트를 100% 지울 것.
2. (JSON APPLY) 오직 [번역 매핑 데이터 JSON]의 'target_text'만 렌더링할 것.
3. (FONT & LAYOUT) Noto Sans SC (思源黑体) 산세리프 스타일로 렌더링하되, 한자 특성을 고려하여 한국어 대비 폰트 크기 10% 축소, 행간 15% 확장 적용.
4. (PACKAGE PRESERVATION) 본품 표면 영문/로고 100% 보존.

[번역 매핑 데이터 JSON]
{json_data}
"""
    else:  # TW
        pass1 = f"""
[SYSTEM PROMPT] Global Luxury Beauty Transcreation & Compliance Automator (Taiwan Traditional Mode)
{qa_block}
## 1. System Role & Persona
당신은 시슬리, SK-II, 랑콤 등 대만/홍콩 프레스티지 더마 뷰티 시장을 총괄하는 10년 차 수석 크리에이티브 디렉터이자 Shopee TW / momo 전문 엘리트 카피라이터입니다.
단순 직역을 배제하고, 대만 현지 소비자가 열광하는 우아하고 지적인 메디컬 코스메틱(더마) 스타일의 프리미엄 카피로 초월번역(Transcreation)하세요.

## 2. Core Transcreation Principles & Mandatory Glossary (Strict Compliance)
### A. 고유 명사 및 브랜드명 필수 표기 지침 (Brand & Ingredient Glossary)
- **品牌名稱 (Brand Name)**: 保留高固有英文名稱 **`Logicall Skin`**，嚴禁音譯。
- **Aquatide & Serum 標記 (必須中英並記)**:
  * Aquatide ➔ **`阿夸肽 (Aquatide)`**
  * Serum ➔ **`精華液 (Serum)`**
  * 產品官方品名 ➔ **`阿夸肽修護精華液 (Aquatide Resurface Serum)`**
- **獨家專利成分原形保留**: 'LiftDerm', 'Aquatide', 'Lifting Logic for eye' 等國際專利成分名維持英文原形。
### B. Eliminate Translationese & 1:1 Matching
- '確實', '真正', '非常', '絕對' 등 딱딱한 직역 부사를 전면 금지하고, 대만 럭셔리 뷰티 전문 어휘로 매끄럽게 재창조하십시오.
### C. Natural Sentence Flow & Syntactic Restructuring
- "10% LiftDerm"이나 "Aquatide" 등 활성 성분 비율이나 특정 수치가 문장 중간에 끊기지 않고 자연스러운 제품명 및 효능 서사로 이어지도록 문장 구조를 우아하게 재조정하십시오.
### D. Use Premium Beauty & Biotech Terminology
- 피부 속/기저층: 肌底 / 肌底深層
- 토탈 케어/멀티 코렉티브/리서페이스: 多效修護 / 全方位全效修護 / 煥活肌底
- 탄력 복원/강화: 賦活肌底彈力 / 喚醒肌膚澎潤彈性
- 눈가 잔주름/건조주름: 細紋・乾紋 / 撫平細紋
- 보습/장벽/에센스: 保濕 / 鎖水保護膜 / 保濕屏障 / 精華液 / 菸鹼醯胺
- 독자 성분명 영문 보존: 'LiftDerm', 'Aquatide', 'Lifting Logic for eye' 등 글로벌 성분명은 영문 원형 유지.

## 3. Global Cosmetic Regulatory Screening & TFDA Compliance (Mandatory Guardrails)
### A. Ban on Absolute & Unverifiable Claims (절대적/과대 표현 전면 금지)
- 'World's First', 'No.1', 'Best', 'The Ultimate' 등 검증 불가능한 절대 표현(`全球首創`, `第一`, `最佳`, `終極對策`, `極致` 등) 사용을 엄격히 금지합니다.
- 반드시 혁신 기술 및 프리미엄 케어 용어로 의무 순화하십시오 (예: `專為眼周修護研發的創新科技`, `頂級多效`, `精準修護對策`).
### B. Ban on Medical/Clinical Misinterpretation (의료 시술 오인 금지 및 4대 안전 동사)
- '주름 박멸', '영구 삭제' 등 의학적 시술(보톡스/필러)을 연상시키는 표현 전면 금지.
- 반드시 화장품 규정 내 안전 동사인 **`撫平` (Smooth), `淡化` (Diminish), `舒緩` (Alleviate), `修護` (Care/Repair)**만을 사용하여 표현하십시오.
- 대만 정체자(繁體中文, zh-TW) 100% 필수.

## 4. 이미지 텍스트 전수 추출 & 출력 포맷
이미지 내의 모든 한국어 텍스트는 단 하나도 빠짐없이 100% 추출하십시오. (고시표인 경우 is_table: true 설정)
출력은 반드시 순수 JSON이어야 합니다:
{{
  "is_table": false,
  "translation_map": [
    {{
      "kor": "한국어 원문",
      "target_text": "대만 TFDA 및 럭셔리 초월번역 준수 정체자 카피",
      "reasoning": "TFDA 절대표현 순화 및 럭셔리 초월번역 사유"
    }}
  ]
}}
"""
        pass2_tmpl = """
당신은 정밀한 시각적 로컬라이제이션을 수행하는 이미지 인페인팅 AI입니다.
아래 [번역 매핑 데이터 JSON]을 바탕으로 단일 이미지를 생성하세요.

[엄격 렌더링 규칙 - 대만 정체자(繁體中文) 전용]
1. (KOR ERASING) 원본 한국어 텍스트 100% 지울 것.
2. (JSON APPLY) 오직 'target_text'의 대만 정체자만 렌더링할 것.
3. (FONT & LAYOUT) Noto Sans TC / 思源黑體 스타일로 렌더링하되, 한자 특성을 고려하여 한국어 대비 폰트 크기 10% 축소, 행간 15% 확장 적용.
4. (PACKAGE PRESERVATION) 본품 표면 영문/로고 100% 보존.
5. (🚨 CRITICAL: ABSOLUTE TRADITIONAL GLYPH ENFORCEMENT)
   - 절대 중국 본토 간체자(Simplified Chinese) 획수를 그리지 마십시오! 
   - 확산 모델의 간체자 쏠림(Drift)을 엄격히 차단하고, 반드시 획수가 많은 순수 번체자(繁體字/正體字) 글리프를 정확히 그리십시오:
     * 養 (O) vs 养 (X - 절대금지) -> 保養, 營養, 調理
     * 對 (O) vs 对 (X - 절대금지) -> 對策, 針對
     * 護 (O) vs 护 (X - 절대금지) -> 修護, 護理
     * 創 (O) vs 创 (X - 절대금지) -> 首創, 創造
     * 變 (O) vs 变 (X - 절대금지) -> 變得, 變化
     * 顯 (O) vs 显 (X - 절대금지) -> 明顯
     * 實 (O) vs 实 (X - 절대금지) -> 確實, 實驗
     * 體 (O) vs 体 (X - 절대금지) -> 體驗, 身體
     * 驗 (O) vs 验 (X - 절대금지) -> 體驗, 實驗
     * 緊 (O) vs 紧 (X - 절대금지) -> 緊緻
     * 緻 (O) vs 致 (X - 절대금지) -> 細緻, 緊緻
     * 膚 (O) vs 肤 (X - 절대금지) -> 肌膚, 膚質
     * 雙 (O) vs 双 (X - 절대금지) -> 雙重
     * 氣 (O) vs 气 (X - 절대금지) -> 空氣
     * 隊 (O) vs 队 (X - 절대금지) -> 團隊
     * 劃 (O) vs 划 (X - 절대금지) -> 企劃
     * 歲 (O) vs 岁 (X - 절대금지) -> 30歲
     * 乾 (O) vs 干 (X - 절대금지) -> 乾燥, 乾性
     * 華 (O) vs 华 (X - 절대금지) -> 精華液
     * 濕 (O) vs 湿 (X - 절대금지) -> 保濕
     * 鎖 (O) vs 锁 (X - 절대금지) -> 鎖水
     * 膠 (O) vs 胶 (X - 절대금지) -> 膠原蛋白
     * 纖 (O) vs 纤 (X - 절대금지) -> 纖維

[번역 매핑 데이터 JSON]
{json_data}
"""
    return pass1, pass2_tmpl


# =================================================================================
# 4. 고시정보 표 자동 HTML 렌더러 연동 및 DOCX 파싱 지원
# =================================================================================
def parse_docx_content(docx_path: str) -> List[Tuple[str, str]]:
    """DOCX 파일에서 항목-내용 쌍을 추출합니다."""
    import zipfile
    import xml.etree.ElementTree as ET

    items = []
    try:
        with zipfile.ZipFile(docx_path) as z:
            xml_content = z.read("word/document.xml")
            tree = ET.fromstring(xml_content)
            
            # 테이블 탐색
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            tables = tree.findall(".//w:tbl", ns)
            if tables:
                for tbl in tables:
                    rows = tbl.findall(".//w:tr", ns)
                    for tr in rows:
                        cells = tr.findall(".//w:tc", ns)
                        if len(cells) >= 2:
                            c0_texts = [t.text for t in cells[0].findall(".//w:t", ns) if t.text]
                            c1_texts = [t.text for t in cells[1].findall(".//w:t", ns) if t.text]
                            lbl = "".join(c0_texts).strip()
                            val = "\n".join(["".join([t.text for t in p.findall(".//w:t", ns) if t.text]) for p in cells[1].findall(".//w:p", ns)]).strip()
                            if lbl and lbl != "항목":
                                items.append((lbl, val))
            
            # 테이블이 없거나 비어있는 경우 단락 파싱
            if not items:
                paragraphs = []
                for p in tree.findall(".//w:p", ns):
                    txt = "".join([t.text for t in p.findall(".//w:t", ns) if t.text]).strip()
                    if txt:
                        paragraphs.append(txt)
                
                # 라인 단위로 항목-내용 파싱
                i = 0
                while i < len(paragraphs):
                    p = paragraphs[i]
                    if p in ["상품상세정보", "항목", "내용"]:
                        i += 1
                        continue
                    if i + 1 < len(paragraphs):
                        items.append((p, paragraphs[i+1]))
                        i += 2
                    else:
                        items.append((p, ""))
                        i += 1
    except Exception as e:
        print(f"  ❌ [DOCX ERROR] 파싱 실패: {e}")
    return items


def standardize_notice_table_items(items: List[Dict[str, str]], lang_code: str) -> List[Dict[str, str]]:
    """고시정보표 4대 필수 법률 조항(고객상담번호 +82, 기능성 심사필, 주의사항 3대 조항, 공정위 품질보증기준)을 전역 표준화합니다."""
    upper_lang = str(lang_code).upper()

    for item in items:
        lbl = item.get("label", "").lower()
        val = item.get("value", "").strip()

        # 1) 고객상담 전화번호 국제번호(+82) 표준화
        if any(k in lbl for k in ["customer", "contact", "phone", "电话", "電話", "상담", "문의"]):
            if "02-" in val or "02)" in val or "02." in val or "6743-3206" in val:
                val_cleaned = re.sub(r"^02[-.)\s]*", "+82-2-", val)
                if "+82" not in val_cleaned and "6743-3206" in val_cleaned:
                    val_cleaned = "+82-2-6743-3206"
                item["value"] = val_cleaned

        # 2) 기능성화장품 심사필 표준화
        if any(k in lbl for k in ["functional cosmetics", "special use", "특수용도", "기능성", "심사", "review"]):
            if not val or val.lower() in ["not specified", "해당없음", "none", "null", ""]:
                val = "화장품법에 따른 기능성 화장품 심사(또는 보고)를 필함"

            if "TW" in upper_lang or "HK" in upper_lang:
                prefix = "已依韓國化粧品法完成韓國食品醫藥品安全處(MFDS)特定用途(功能性)化粧品審查(或報告)"
            elif "CN" in upper_lang or "ZH" in upper_lang:
                prefix = "已依据韩国化妆品法完成韩国食品医药品安全处(MFDS)特殊用途(功能性)化妆品审查(或报告)"
            elif "JP" in upper_lang or "JA" in upper_lang:
                prefix = "韓国化粧品法に基づき韓国食品医薬品安全処(MFDS)の機能性化粧品審査(または報告)済"
            elif "KO" in upper_lang:
                prefix = "화장품법에 따른 식품의약품안전처(MFDS) 기능성화장품 심사(또는 보고)를 필함"
            else:
                prefix = "Completed Functional Cosmetics Review (or Report) with the Ministry of Food and Drug Safety (MFDS, Republic of Korea) in accordance with the Cosmetics Act"

            match = re.match(r'^([Yy](?:es)?|[Oo]|심사필|해당(?:있음)?|是|已完成(?:審查|审查)?)\b', val, re.IGNORECASE)
            if match:
                remainder = val[match.end(1):].strip()
                if not remainder:
                    item["value"] = prefix
                elif remainder.startswith('(') or remainder.startswith('（'):
                    item["value"] = f"{prefix} {remainder}"
                else:
                    clean_rem = remainder.lstrip("- ")
                    item["value"] = f"{prefix} - {clean_rem}"
            elif val.lower() in ["not specified", "해당없음", "none", "null", "화장품법에 따른 기능성 화장품 심사(또는 보고)를 필함"]:
                item["value"] = prefix

        # 3) 사용할 때의 주의사항 3대 조항 표준화 (누락 또는 축약 시 법정 필수 조항으로 자동 보강)
        if any(k in lbl for k in ["precaution", "caution", "주의사항", "注意事项", "注意事項"]):
            if not val or len(val) < 25 or "not specified" in val.lower() or "상세" in val or "별도" in val:
                if "TW" in upper_lang or "HK" in upper_lang:
                    item["value"] = "1) 使用時或使用後因陽光直射導致使用部位出現紅斑、腫脹或搔癢等異常症狀或副作用時，請諮詢專業醫師。 2) 請勿使用於傷口等異常部位。 3) 保存及處理注意事項：A) 請置於嬰幼兒無法取得之處。B) 請避免陽光直射保存。"
                elif "CN" in upper_lang or "ZH" in upper_lang:
                    item["value"] = "1) 使用时或使用后因直射光线导致使用部位出现红斑、肿胀或瘙痒等异常症状或副作用时，请咨询专业医生。 2) 伤口等异常部位请勿使用。 3) 保管及处理注意事项：A) 请置于儿童接触不到的地方保管。B) 请避开直射光线保管。"
                elif "JP" in upper_lang or "JA" in upper_lang:
                    item["value"] = "1) お肌に異常が生じていないかよく注意して使用してください。化粧品がお肌に合わないとき即ち次のような場合には、使用を中止し、皮膚科専門医等にご相談されることをおすすめします。 2) 傷やはれもの、しっしん等、異常のある部位にはお使いにならないでください。 3) 保管及び取扱い上の注意：A) 乳幼児の手の届かないところに保管してください。B) 直射日光のあたる場所には保管しないでください。"
                elif "KO" in upper_lang:
                    item["value"] = "1) 화장품 사용 시 또는 사용 후 직사광선에 의하여 사용부위가 붉은 반점, 부어오름 또는 가려움증 등의 이상 증상이나 부작용이 있는 경우 전문의 등과 상담할 것 2) 상처가 있는 부위 등에는 사용을 자제할 것 3) 보관 및 취급 시의 주의사항: 가) 어린이의 손이 닿지 않는 곳에 보관할 것 나) 직사광선을 피해서 보관할 것"
                else:
                    item["value"] = "1) Consult a specialist if there are abnormal symptoms or side effects such as red spots, swelling, or itching caused by direct sunlight during or after use. 2) Refrain from using on wounded areas. 3) Precautions for storage and handling: A) Keep out of reach of children. B) Store away from direct sunlight."

        # 4) 품질보증기준 표준화 (공정거래위원회 소비자분쟁해결기준)
        if any(k in lbl for k in ["quality assurance", "warranty", "품질보증", "质量保证", "品質保證", "品質保証"]):
            if not val or len(val) < 15 or "not specified" in val.lower() or "공정거래위원회" in val or "fair trade" in val.lower():
                if "TW" in upper_lang or "HK" in upper_lang:
                    item["value"] = "本產品如有任何異常，將依公平交易委員會公告之「消費者爭議解決基準」予以賠償。"
                elif "CN" in upper_lang or "ZH" in upper_lang:
                    item["value"] = "若本产品出现质量问题，依据公平交易委员会告示“消费者纷争解决标准”提供补偿。"
                elif "JP" in upper_lang or "JA" in upper_lang:
                    item["value"] = "本商品に異常がある場合、公正取引委員会告示（消費者紛争解決基準）に基づき補償いたします。"
                elif "KO" in upper_lang:
                    item["value"] = "본 제품에 이상이 있을 경우 공정거래위원회 고시 '소비자분쟁해결기준'에 의거하여 보상해 드립니다."
                else:
                    item["value"] = "Compensation will be provided in accordance with the Fair Trade Commission's Consumer Dispute Settlement Standards."

    return items


async def render_notice_table(client: genai.Client, mapping_data: Dict[str, Any], lang_code: str, out_path: str, orig_width: int = 860) -> bool:
    """render_notice_table_standard.py 모듈을 활용하여 1열(번역 라벨) 및 2열(번역 값) 고선명 표 이미지를 렌더링합니다."""
    std_script_dir = os.path.join(PROJECT_ROOT, "00_공통자료")
    if std_script_dir not in sys.path:
        sys.path.insert(0, std_script_dir)

    try:
        import render_notice_table_standard as rnts
        
        prompt_en = """
You are a senior regulatory affairs and e-commerce localization expert.
Format the following cosmetic product details into structured table items for Amazon / Sephora US.
Map standard labels: Volume -> Size/Net Wt., Skin Type, Shelf Life, Directions, Manufacturer/Distributed by, Country of Origin, Ingredients, Functional Cosmetics Review Status, Precautions for Use, Quality Assurance Standard, Customer Service (+82-2-6743-3206).

[CRITICAL TRANSCREATION RULES FOR VALUES (2nd Column)]:
- Functional Cosmetics Review Status: NEVER output a raw letter 'Y'. If Korean says 'Y (미백, 주름개선 등)' or '해당', ALWAYS transcreate into 'MFDS-Certified Functional Cosmetic (Brightening, Wrinkle Improvement, UV Protection)'.
- Country of Origin: Always output 'Republic of Korea'.
- Skin Type: Transcreate '모든피부용' to 'All Skin Types'.
- Customer Service: Format with '+82-2-6743-3206'.

Output MUST be a JSON object: {"title": "PRODUCT DETAILS", "items": [{"label": "Size / Net Wt.", "value": "..."}, ...]}
"""
        prompt_cn = """
你是资深化妆品法规与电商本地化专家。请将以下韩国化妆品产品详细信息（中文告示表）整理并翻译为规范的简体中文告示表格（严禁将韩文原标签作为第1列，第1列必须是规范的中文项目名，第2列是对应的中文内容值）。
严格对照标准项目名：
- 净含量 / 容量
- 适用肤质
- 使用期限 / 保质期
- 使用方法
- 化妆品生产企业 / 责任销售商
- 原产国
- 全成分 (INCI/NMPA 标准全成分名称)
- 特殊用途化妆品审查状态 (如：已完成审查 (美白、改善皱纹双重功效) / 不适用)
- 使用注意事项
- 质量保证标准
- 消费者咨询电话 (格式为：+82-2-6743-3206)

输出必须为纯 JSON 格式：
{
  "title": "商品基本信息",
  "items": [
    {"label": "净含量", "value": "25g x 3片"},
    {"label": "适用肤质", "value": "适合所有肤质"},
    ...
  ]
}
"""
        prompt_jp = """
あなたは日本の化粧品薬機法およびQoo10 Japanの専門家です。以下の商品情報を整理し、規範的な日本語の告示表項目にマッピングしてください（第1列は日本語の項目名、第2列は対応する値）。
標準項目名：内容量、お肌のタイプ、使用期限、ご使用方法、製造販売元、原産国、全成分、医薬部外品承認/機能性化粧品審査、ご使用上の注意、品質保証基準、お客様相談窓口 (+82-2-6743-3206)。
出力形式：{"title": "商品基本情報", "items": [{"label": "内容量", "value": "..."}, ...]}
"""
        prompt_tw = """
你是資深化妝品法規與電商本地化專家。請將以下產品詳細資訊整理為規範的繁體中文告示表格（第1列為規範的繁體中文項目名，第2列為對應的值）。
標準項目名：淨含量 / 容量、適用膚質、保存期限、使用方法、製造商 / 責任銷售商、產地、全成分、特殊用途化妝品審查、注意事項、質量保證、客服專線 (+82-2-6743-3206)。
輸出格式：{"title": "商品基本資訊", "items": [{"label": "淨含量", "value": "..."}, ...]}
"""
        p_map = {"EN": prompt_en, "CN": prompt_cn, "JP": prompt_jp, "TW": prompt_tw}
        selected_prompt = p_map.get(lang_code, prompt_cn)

        t_map = mapping_data.get("translation_map", [])
        input_lines = []
        for row in t_map:
            k = row.get("kor", "").strip()
            t = row.get("target_text", "").strip()
            if k and t:
                input_lines.append(f"[{k}] {t}")
            elif t:
                input_lines.append(t)
            elif k:
                input_lines.append(k)

        full_prompt = f"{selected_prompt}\n\n[입력된 고시정보 표 원문 및 번역 텍스트 목록]\n" + "\n".join(input_lines)
        
        print(f"  🔍 [표 구조 정제] Gemini 3.1 Pro로 항목 라벨과 값을 2열 구조화 중 ({MODEL_PRO})...", flush=True)
        # [수정 1] 표 렌더러 - 대용량 데이터 안전망 (8192 토큰) 및 황금 비율 하이퍼파라미터 적용
        resp = await client.aio.models.generate_content(
            model=MODEL_PRO,
            contents=[full_prompt],
            config=types.GenerateContentConfig(
                system_instruction=GLOBAL_COMPLIANCE_SYSTEM_INSTRUCTION,
                response_mime_type="application/json",
                temperature=0.6,
                top_p=0.9,
                max_output_tokens=8192
            )
        )
        res_json = json.loads(resp.text.strip())
        items = res_json.get("items", [])
        title = res_json.get("title", "商品基本信息")

        # [4대 필수 법률 조항 전역 표준화 게이트] (고객상담번호 +82, 기능성 심사필, 주의사항 3대 조항, 공정위 품질보증기준)
        items = standardize_notice_table_items(items, lang_code)

        return rnts.render_notice_table_to_png(title, items, out_path, lang=lang_code)
    except Exception as e:
        print(f"  -> [TABLE RENDER FAIL] 표 렌더러 예외: {e}")
        return False


async def process_docx_notice_table(client: genai.Client, docx_path: str, out_path: str, lang_code: str) -> bool:
    """DOCX 고시정보표를 번역하여 표준 고선명 HTML 테이블 PNG로 렌더링합니다."""
    print(f"\n================================================================================")
    print(f"📄 [DOCX 고시정보표 번역] {os.path.basename(docx_path)} -> [{LANG_CONFIGS[lang_code]['name']}]")
    print(f"================================================================================")
    
    raw_items = parse_docx_content(docx_path)
    if not raw_items:
        print(f"  ❌ [ERROR] DOCX 파일에서 고시정보 항목을 추출하지 못했습니다.")
        return False

    prompt_en = """
You are a senior regulatory affairs and e-commerce localization expert.
Translate the following Korean cosmetic product details (specifications table) into professional English for Amazon / Sephora US.

[CRITICAL INSTRUCTION: STANDARD FIELD NAME MAPPING]
You MUST strictly map the following Korean labels to these standardized global beauty e-commerce terms. Do not use direct translations if they deviate from this list:
- 용량 또는 중량 (내용물의 용량): Size / Net Wt.
- 제품 주요 사양: Skin Type
- 사용기한 또는 개봉 후 사용기간: Shelf Life / PAO
- 사용방법: Directions
- 화장품제조업자 및 책임판매업자: Manufacturer / Distributed by
- 제조국: Country of Origin
- 전성분: Ingredients (국제화장품원료집(INCI) 및 한국화장품성분사전(KCID) 표준 기반 미국 FDA/PCPC 표기법 및 띄어쓰기 규격 강제 적용)
- 기능성 화장품 심사 필 유무: Functional Cosmetics Review Status
- 사용할 때의 주의사항: Precautions for Use / Cautions
- 품질보증기준: Quality Assurance Standard
- 소비자 상담 전화번호: Customer Service / Contact (MUST format Korean phone number with international country code, e.g. +82-2-6743-3206)

Output MUST be a JSON object:
{
  "title": "PRODUCT DETAILS",
  "items": [
    {"label": "Volume / Net Weight", "value": "25ml"},
    ...
  ]
}
"""

    prompt_cn = """
你是资深化妆品法规与电商本地化专家。请将以下韩国化妆品产品详细信息（中文告示表）翻译为规范的简体中文，严格遵守中国国家药监局(NMPA)及新广告法规范。

[CRITICAL INSTRUCTION: STANDARD FIELD NAME MAPPING]
You MUST strictly map the following Korean labels to these standardized Chinese e-commerce/legal terms. Do not use direct translations if they deviate from this list:
- 용량 또는 중량 (내용물의 용량): 净含量 / 容量
- 제품 주요 사양: 适用肤质 / 产品规格
- 사용기한 또는 개봉 후 사용기간: 使用期限 / 保质期
- 사용방법: 使用方法
- 화장품제조업자 및 책임판매업자: 化妆品生产企业 / 责任销售商
- 제조국: 原产国 / 产地
- 전성분: 全成分 (INCI 및 KCID 기반 中国国家药品监督管理局(NMPA) 标准名称 및 띄어쓰기 규격 강제 적용)
- 기능성 화장품 심사 필 유무: 特殊用途化妆品审查状态 (如：已完成审查 (美白、改善皱纹双重功效))
- 사용할 때의 주의사항: 使用注意事项
- 품질보증기준: 质量保证标准
- 소비자 상담 전화번호: 消费者咨询电话 (电话号码必须带有韩国国际区号，格式为：+82-2-6743-3206)

输出必须为纯 JSON 格式：
{
  "title": "商品基本信息",
  "items": [
    {"label": "净含量", "value": "25ml"},
    ...
  ]
}
"""

    prompt_jp = """
あなたは日本の化粧品薬機法およびQoo10 Japanの専門家です。以下の韓国化粧品の商品基本情報表を自然で正確な日本語に翻訳してください。

[CRITICAL INSTRUCTION: STANDARD FIELD NAME MAPPING]
You MUST strictly map the following Korean labels to these standardized Japanese e-commerce/legal terms. Do not use direct translations if they deviate from this list:
- 용량 또는 중량 (내용물의 용량): 内容量
- 제품 주요 사양: お肌のタイプ / 対象肌
- 사용기한 또는 개봉 후 사용기간: 使用期限
- 사용방법: ご使用方法
- 화장품제조업자 및 책임판매업자: 製造販売元
- 제조국: 原産国
- 전성분: 全成分 (INCI 및 KCID 기반 일본 화장품공업연합회(JCIA) 표준 명칭 및 띄어쓰기 규격 강제 적용)
- 기능성 화장품 심사 필 유무: 医薬部外品承認 / 機能性化粧品審査
- 사용할 때의 주의사항: ご使用上の注意
- 품질보증기준: 品質保証基準
- 소비자 상담 전화번호: お客様相談窓口 (お客様相談電話番号は必ず韓国国際国番号付き +82-2-6743-3206 で表記してください)

出力は純粋なJSONオブジェクトである必要があります：
{
  "title": "商品基本情報",
  "items": [
    {"label": "内容量", "value": "25ml"},
    ...
  ]
}
"""
    prompt_tw = """
你是資深化妝品法規與電商本地化專家。請將以下產品詳細資訊翻譯為規範的繁體中文（台灣/香港TFDA標準）。

[CRITICAL INSTRUCTION: STANDARD FIELD NAME MAPPING]
You MUST strictly map the following Korean labels to these standardized Taiwanese e-commerce/legal terms. Do not use direct translations if they deviate from this list:
- 용량 또는 중량 (내용물의 용량): 淨含量 / 容量
- 제품 주요 사양: 適用膚質
- 사용기한 또는 개봉 후 사용기간: 保存期限
- 사용방법: 使用方法
- 화장품제조업자 및 책임판매업자: 製造商 / 責任銷售商
- 제조국: 產地
- 전성분: 全成分 (INCI 및 KCID 기반 대만 위생복리부(TFDA) 표준 명칭 및 띄어쓰기 규격 강제 적용)
- 기능성 화장품 심사 필 유무: 含藥化妝品許可 / 特殊用途化妝品審查
- 사용할 때의 주의사항: 注意事項
- 품질보증기준: 售後服務 / 質量保證
- 소비자 상담 전화번호: 客服專線 (客服諮詢電話必須帶有韓國國際區號，例如：+82-2-6743-3206)

輸出必須為純 JSON 格式：
{
  "title": "商品基本資訊",
  "items": [
    {"label": "容量 / 淨含量", "value": "25ml"},
    ...
  ]
}
"""
    if lang_code == "KR":
        print(f"  🎨 [한국어 원본 고시표 렌더러 가동] Gemini 3.1 Pro 지능형 정제 및 860px 렌더링...", flush=True)
        try:
            import render_notice_table_korean as rntk
        except ImportError:
            sys.path.insert(0, SCRIPT_DIR)
            import render_notice_table_korean as rntk
        items_kr = [{"label": lbl, "value": val} for lbl, val in raw_items if lbl != "항목"]
        rntk.render_korean_notice_table("상품 상세 정보", items_kr, out_path, max_height=2580, use_gemini=True)
        base_name, ext = os.path.splitext(out_path)
        part1_path = f"{base_name}_Part1{ext}"
        part2_path = f"{base_name}_Part2{ext}"
        if os.path.exists(out_path) or (os.path.exists(part1_path) and os.path.exists(part2_path)):
            print(f"  🎉 [SUCCESS] 한국어 원본 고시정보 표 PNG 렌더링 완료: {os.path.basename(out_path)}")
            return True
        else:
            print(f"  ❌ [ERROR] 한국어 원본 고시정보 표 PNG 렌더링 실패")
            return False

    p_map = {"EN": prompt_en, "CN": prompt_cn, "JP": prompt_jp, "TW": prompt_tw}
    selected_prompt = p_map.get(lang_code, prompt_en)

    input_text = "\n".join([f"[{lbl}]\n{val}" for lbl, val in raw_items])
    full_prompt = f"{selected_prompt}\n\n[입력 고시정보 표 데이터]\n{input_text}"

    print(f"  🔍 [PASS 1] 고시정보 표 텍스트 다국어 번역 및 표준화 중 ({MODEL_PRO})...", flush=True)
    try:
        # [수정 2] DOCX 고시표 번역 - 대용량 데이터 안전망 (8192 토큰) 및 황금 비율 하이퍼파라미터 적용
        resp = await client.aio.models.generate_content(
            model=MODEL_PRO,
            contents=[full_prompt],
            config=types.GenerateContentConfig(
                system_instruction=GLOBAL_COMPLIANCE_SYSTEM_INSTRUCTION,
                response_mime_type="application/json",
                temperature=0.6,
                top_p=0.9,
                max_output_tokens=8192
            )
        )
        res_json = json.loads(resp.text.strip())
        items = res_json.get("items", [])
        title = res_json.get("title", LANG_CONFIGS[lang_code]["name"])
        
        # [사용자 규칙 강제]: 고객상담 전화번호는 반드시 +82 국제전화 국가번호로 표기
        # [4대 필수 법률 조항 전역 표준화 게이트] (고객상담번호 +82, 기능성 심사필, 주의사항 3대 조항, 공정위 품질보증기준)
        items = standardize_notice_table_items(items, lang_code)

        print(f"  ✅ [PASS 1 완료] 총 {len(items)}개 고시 항목 번역 완료 (국제 전화번호 +82 규격 동기화)")
    except Exception as e:
        print(f"  ❌ [ERROR] DOCX 번역 실패: {e}")
        return False

    # 렌더링 호출
    print(f"  🎨 [PASS 2] 고선명 HTML 표준 헤드리스 렌더러 가동...", flush=True)
    std_script_dir = os.path.join(PROJECT_ROOT, "00_공통자료")
    if std_script_dir not in sys.path:
        sys.path.insert(0, std_script_dir)
    import render_notice_table_standard as rnts
    success = rnts.render_notice_table_to_png(title, items, out_path, lang=lang_code)
    
    base_name, ext = os.path.splitext(out_path)
    part1_path = f"{base_name}_Part1{ext}"
    part2_path = f"{base_name}_Part2{ext}"

    if success and (os.path.exists(out_path) or (os.path.exists(part1_path) and os.path.exists(part2_path))):
        print(f"  🎉 [SUCCESS] 고시정보 표 PNG 렌더링 완료: {os.path.basename(out_path)}")
        return True
    else:
        print(f"  ❌ [ERROR] 고시정보 표 PNG 렌더링 실패")
        return False


# =================================================================================
# 5. 핵심 번역 파이프라인
# =================================================================================
async def process_single_image(client: genai.Client, in_path: str, out_path: str, lang_code: str, product_name: str = "") -> bool:
    print(f"\n================================================================================")
    print(f"🖼️ [번역 시작] {os.path.basename(in_path)} -> [{LANG_CONFIGS[lang_code]['name']}]")
    print(f"================================================================================")

    try:
        original_image = Image.open(in_path)
        if original_image.mode != "RGB":
            original_image = original_image.convert("RGB")
        original_image.load()
        orig_w, orig_h = original_image.size
    except Exception as e:
        print(f"  ❌ [ERROR] 이미지 파일 로드 실패: {e}")
        return False

    source_dir = os.path.dirname(in_path)
    if not product_name:
        product_name = os.path.basename(source_dir)

    # 1단계 QA 진단 결과 탐색 및 프롬프트 주입문 생성
    qa_rules = load_qa_feedback_and_transcreation_rules(source_dir, product_name, lang_code)
    pass1_prompt, pass2_tmpl = build_prompts(lang_code, qa_rules.get("prompt_override_text", ""))

    # PASS 1: 텍스트 추출 및 번역 매핑 (재시도 로직 포함)
    p1_json = None
    for attempt in range(1, 4):
        print(f"  🔍 [PASS 1] 텍스트 스캔 및 다국어 매핑 추출 중 ({MODEL_PRO}, 시도 {attempt}/3)...", flush=True)
        try:
            # [수정 3] 이미지 텍스트 전수 스캔 - 대용량 데이터 안전망 (8192 토큰) 및 황금 비율 하이퍼파라미터 적용
            response_p1 = await client.aio.models.generate_content(
                model=MODEL_PRO,
                contents=[original_image, pass1_prompt],
                config=types.GenerateContentConfig(
                    system_instruction=GLOBAL_COMPLIANCE_SYSTEM_INSTRUCTION,
                    response_mime_type="application/json",
                    temperature=0.6,
                    top_p=0.9,
                    max_output_tokens=8192
                )
            )
            p1_text = response_p1.text.strip()
            p1_json = json.loads(p1_text)
            break
        except Exception as e:
            print(f"  ⚠️ [WARN] PASS 1 시도 {attempt} 실패: {e}")
            if attempt < 3:
                wait_t = 15 * attempt
                print(f"  ⏳ {wait_t}초 대기 후 재시도합니다...", flush=True)
                time.sleep(wait_t)
            else:
                print(f"  ❌ [ERROR] PASS 1 최종 실패")
                return False

    if isinstance(p1_json, list):
        t_map = p1_json
    else:
        t_map = p1_json.get("translation_map", []) if isinstance(p1_json, dict) else []

    # 1단계 QA 진단 결과 결정론적 보정 게이트 (오타 7종 및 MoCRA 100% 반영)
    t_map = apply_deterministic_qa_overrides(t_map, qa_rules, lang_code)
    if isinstance(p1_json, dict):
        p1_json["translation_map"] = t_map
    elif isinstance(p1_json, list):
        p1_json = t_map
    print(f"  ✅ [PASS 1 완료] 총 {len(t_map)}개 텍스트 블록 추출 완료")
    for i, item in enumerate(t_map[:3]):
        print(f"     ({i+1}) 원문: {item.get('kor', '')[:30]} -> 번역: {item.get('target_text', '')[:30]}")
    if len(t_map) > 3:
        print(f"     ... 외 {len(t_map)-3}개 항목")

    # 고시정보표 감지 시 분기 처리
    if isinstance(p1_json, list):
        is_table = False
    else:
        is_table = p1_json.get("is_table", False)
    fname_lower = os.path.basename(in_path).lower()
    if is_table or "notice" in fname_lower or "상세정보" in fname_lower or "spec" in fname_lower or "details" in fname_lower:
        print(f"  📊 [TABLE DETECTED] 고시정보표 레이아웃 감지 -> HTML 표준 렌더러 분기 가동")
        rendered = await render_notice_table(client, p1_json, lang_code, out_path, orig_width=orig_w)
        if rendered:
            print(f"  🎉 [SUCCESS] 고선명 표 이미지 생성 완료: {os.path.basename(out_path)}")
            return True
        else:
            print(f"  ⚠️ [INFO] HTML 표 렌더러 실패 -> 일반 인페인팅 모드로 폴백 진행")

    # PASS 2: 신경망 인페인팅 렌더링 (지수 백오프 재시도 포함)
    clean_json_str = json.dumps(p1_json, ensure_ascii=False, indent=2)
    pass2_prompt = pass2_tmpl.format(json_data=clean_json_str)

    for attempt in range(1, 4):
        print(f"  🎨 [PASS 2] 시각적 신경망 인페인팅 렌더링 가동 ({MODEL_FLASH_IMAGE}, 시도 {attempt}/3)...", flush=True)
        try:
            # [수정 4] 이미지 인페인팅 렌더링 - 일관된 하이퍼파라미터 적용
            response_p2 = await client.aio.models.generate_content(
                model=MODEL_FLASH_IMAGE,
                contents=[original_image, pass2_prompt],
                config=types.GenerateContentConfig(
                    response_modalities=["IMAGE"],
                    temperature=0.6,
                    top_p=0.9
                )
            )

            rendered_img_bytes = None
            for part in response_p2.candidates[0].content.parts:
                if part.inline_data:
                    rendered_img_bytes = part.inline_data.data
                    break

            if not rendered_img_bytes:
                raise ValueError("PASS 2 이미지 데이터 반환 없음")

            img = Image.open(io.BytesIO(rendered_img_bytes))
            if img.size != (orig_w, orig_h):
                img = img.resize((orig_w, orig_h), Image.Resampling.LANCZOS)

            os.makedirs(os.path.dirname(out_path), exist_ok=True)
            img.save(out_path, format="PNG", optimize=True)
            print(f"  🎉 [SUCCESS] 렌더링 완료 및 저장: {os.path.basename(out_path)} ({orig_w}x{orig_h})")
            return True
        except Exception as e:
            print(f"  ⚠️ [WARN] PASS 2 인페인팅 렌더링 실패: {e}")
            if attempt < 3:
                wait_time = 15 * attempt
                print(f"  ⏳ {wait_time}초 대기 후 PASS 2 재시도합니다... ({attempt}/3)")
                time.sleep(wait_time)
            else:
                print("  ❌ [ERROR] PASS 2 인페인팅 최종 실패.")
                return False
    return False


async def run_translation_batch_for_folder(client: genai.Client, current_source_dir: str, target_lang: str, product_name: str, custom_target_dir: Optional[str] = None):
    """지정된 단일 리프 폴더(current_source_dir)에 대해 이미지 및 DOCX 번역 배치를 실행합니다."""
    config = LANG_CONFIGS[target_lang]

    # 0. 1단계 QA 진단 결과 파일 자동 연동 (transcreation_guide.json)
    guide_dst = os.path.join(current_source_dir, "transcreation_guide.json")
    is_remediation_mode = False
    if os.path.exists(guide_dst):
        is_remediation_mode = True
    else:
        qa_src_candidates = [
            os.path.join(PROJECT_ROOT, "03_번역품질평가", "02_진단결과", product_name, "Transcreation_QA_Report.json"),
            os.path.join(PROJECT_ROOT, "03_번역품질평가", "02_진단결과", f"{product_name}_{target_lang}", "Transcreation_QA_Report.json"),
            os.path.join(PROJECT_ROOT, "03_번역품질평가", "02_진단결과", f"{product_name}_EN", "Transcreation_QA_Report.json"),
        ]
        qa_res_dir = os.path.join(PROJECT_ROOT, "03_번역품질평가", "02_진단결과")
        if os.path.exists(qa_res_dir):
            for s in os.listdir(qa_res_dir):
                sp = os.path.join(qa_res_dir, s, "Transcreation_QA_Report.json")
                if os.path.exists(sp) and sp not in qa_src_candidates:
                    qa_src_candidates.append(sp)

        for qsc in qa_src_candidates:
            if os.path.exists(qsc):
                try:
                    import shutil
                    shutil.copy2(qsc, guide_dst)
                    print(f"🔗 [QA 자동 연동] 1단계 QA 진단 리포트를 소스 폴더 가이드로 동봉 완료: {os.path.basename(guide_dst)}")
                    is_remediation_mode = True
                    break
                except Exception:
                    pass

    if custom_target_dir:
        target_dir = custom_target_dir
    else:
        # [4대 마스터 폴더 체계] 교정 모드일 경우 04_번역교정, 신규 런칭일 경우 02_번역결과_최종에 자동 분기
        base_out = DEFAULT_REMEDIATION_BASE if is_remediation_mode else DEFAULT_OUTPUT_BASE
        target_dir = os.path.join(base_out, f"{product_name}_{config['folder_name']}")
    os.makedirs(target_dir, exist_ok=True)

    print(f"================================================================================")
    print(f"📂 [작업 대상 폴더] {current_source_dir}")
    print(f"📦 [상품 식별명] {product_name}")
    print(f"🌐 [도착 언어] {config['name']}")
    print(f"📁 [저장 위치] {target_dir} ({'🎯 04_번역교정' if is_remediation_mode else '💎 02_신규런칭'})")
    print(f"================================================================================\n")

    # 1. 이미지 파일 처리 (.gif 확장자 포함)
    image_files = sorted([f for f in get_recursive_files(current_source_dir) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.webp', '.gif')) and not f.startswith('~')])
    docx_files = sorted([f for f in get_recursive_files(current_source_dir) if f.lower().endswith('.docx') and not f.startswith('~')])

    total_tasks = len(image_files) + len(docx_files)
    if total_tasks == 0:
        print(f"⚠️ [WARNING] '{current_source_dir}' 폴더에 처리할 이미지나 DOCX 파일이 없습니다.")
        return

    current_idx = 0
    success_count = 0

    for img_name in image_files:
        current_idx += 1
        in_path = os.path.join(current_source_dir, img_name)
        stem = os.path.splitext(img_name)[0]
        out_name = f"{stem}{config['tag']}"
        out_path = os.path.join(target_dir, out_name)

        if os.path.exists(out_path):
            print(f"  ⚠️ [SKIP] 이미 결과물이 존재함: {out_name}")
            success_count += 1
            continue

        print(f"\n[{current_idx}/{total_tasks}] 이미지 작업 시작: {img_name}")
        success = await process_single_image(client, in_path, out_path, target_lang, product_name=product_name)
        if success:
            success_count += 1

        if current_idx < total_tasks:
            print("⏳ API 쿼터 안전 대기 (12초)...", flush=True)
            time.sleep(12)

    # 2. DOCX 고시정보표 처리
    for docx_name in docx_files:
        current_idx += 1
        in_path = os.path.join(current_source_dir, docx_name)
        stem = os.path.splitext(docx_name)[0]
        out_name = f"{stem}{config['tag']}"
        out_path = os.path.join(target_dir, out_name)

        if os.path.exists(out_path):
            print(f"  ⚠️ [SKIP] 이미 결과물이 존재함: {out_name}")
            success_count += 1
            continue

        print(f"\n[{current_idx}/{total_tasks}] DOCX 작업 시작...")
        success = await process_docx_notice_table(client, in_path, out_path, target_lang)
        if success:
            success_count += 1

        if current_idx < total_tasks:
            print("⏳ API 쿼터 안전 대기 (12초)...", flush=True)
            time.sleep(12)

    # 3. SEO / GEO / AEO 메타데이터 TXT 자동 생성
    await generate_seo_geo_aeo_txt(client, current_source_dir, target_dir, target_lang, product_name)

    print(f"\n🏁 [{config['name']}] 번역 및 SEO/GEO/AEO 생성 완료: 총 {total_tasks}개 중 {success_count}개 성공!")
    print(f"📂 저장 경로: {target_dir}\n")


async def generate_seo_geo_aeo_txt(client: genai.Client, current_source_dir: str, target_dir: str, target_lang: str, product_name: str):
    """4-Core 마이크로-써머리 SEO/GEO/AEO (TXT, HTML 뷰어, DOCX, MD) 4종 파일을 자동 생성합니다. (URL/고시표 듀얼 인제스천 및 4개국 법무 렉시콘 100% 결합)"""
    print(f"\n🌐 [SEO/GEO/AEO 4-Core] 정밀 팩트 인제스천 및 4종 포맷(DOCX/HTML/TXT/MD) 생성 중 ({target_lang})...", flush=True)

    # 1. 듀얼 인제스천: url.txt 실시간 웹 스크래핑
    url_fact_context = ""
    url_file_candidates = [
        os.path.join(current_source_dir, "url.txt"),
        os.path.join(current_source_dir, "product_url.txt"),
        os.path.join(current_source_dir, "URL.txt")
    ]
    for u_path in url_file_candidates:
        if os.path.exists(u_path):
            try:
                with open(u_path, "r", encoding="utf-8") as uf:
                    raw_url = uf.read().strip()
                if raw_url.startswith("http"):
                    print(f"  🔗 [URL INGESTION 감지] {raw_url}", flush=True)
                    import urllib.request
                    req = urllib.request.Request(raw_url, headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"})
                    with urllib.request.urlopen(req, timeout=5) as resp:
                        html_bytes = resp.read()
                        raw_html = html_bytes.decode("utf-8", errors="ignore")
                        clean_text = re.sub(r'<script.*?</script>', '', raw_html, flags=re.DOTALL)
                        clean_text = re.sub(r'<style.*?</style>', '', clean_text, flags=re.DOTALL)
                        clean_text = re.sub(r'<[^>]+>', ' ', clean_text)
                        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                        url_fact_context = f"\n[Live Web Product Page Facts (Ground Truth)]\nURL: {raw_url}\nScraped Text Extract: {clean_text[:1200]}\n"
                        print(f"  ✅ [URL 팩트 스크래핑 완료] {len(clean_text)} 자 추출", flush=True)
                        break
            except Exception as ue:
                print(f"  ⚠️ [URL 스크래핑 실패 / 폴백 진행]: {ue}", flush=True)

    # 2. 이미지 팩트 앵커링 & 초월번역 컨텍스트 로드
    guide_path = os.path.join(current_source_dir, "transcreation_guide.json")
    qa_path = os.path.join(target_dir, "Transcreation_QA_Report.json")
    guide_context = ""
    if os.path.exists(guide_path):
        try:
            with open(guide_path, "r", encoding="utf-8") as f:
                gdata = json.load(f)
                guide_context = json.dumps(gdata.get("transcreation_comparisons", [])[:6], ensure_ascii=False)
        except Exception:
            pass
    elif os.path.exists(qa_path):
        try:
            with open(qa_path, "r", encoding="utf-8") as f:
                qdata = json.load(f)
                guide_context = json.dumps(qdata.get("transcreation_comparisons", [])[:6], ensure_ascii=False)
        except Exception:
            pass

    # 3. 4개국 법무 렉시콘 로드 (COMPLIANCE-FIRST)
    lexicon_rules_text = ""
    lexicon_map = {
        "EN": "en_fda_mocra_lexicon.json",
        "JP": "jp_pmda_pharm_lexicon.json",
        "CN": "cn_nmpa_adlaw_lexicon.json",
        "TW": "tw_tfda_lexicon.json"
    }
    lex_file = lexicon_map.get(target_lang)
    if lex_file:
        lex_path = os.path.join(PROJECT_ROOT, "00_공통자료", "compliance_lexicons", lex_file)
        if os.path.exists(lex_path):
            try:
                with open(lex_path, "r", encoding="utf-8") as lf:
                    lex_data = json.load(lf)
                    lexicon_rules_text = f"\n[MANDATORY COMPLIANCE LEXICON ({lex_data.get('jurisdiction', '')})]\n"
                    cats = lex_data.get("categories", {})
                    for c_name, c_val in cats.items():
                        banned = c_val.get("banned_terms", [])
                        for b in banned[:5]:
                            lexicon_rules_text += f"- Banned: '{b.get('banned')}' -> Must use: '{b.get('preferred')}' ({b.get('reason')})\n"
            except Exception:
                pass

    lang_names = {
        "EN": "English for Amazon / Sephora US",
        "JP": "Japanese for Qoo10 Japan / Cosme",
        "CN": "Simplified Chinese for Tmall / Xiaohongshu",
        "TW": "Traditional Chinese for Shopee Taiwan / Momo",
        "KR": "Korean for Naver Smartstore / Coupang"
    }
    target_lang_desc = lang_names.get(target_lang, "English")

    prompt = f"""[SYSTEM PROMPT] Global E-Commerce SEO/GEO/AEO 4-Core Master Copy Generator
Product Name: {product_name}
Target Market & Language: {target_lang_desc}

{url_fact_context}
[Verified Transcreation Context & Ingredients Data]:
{guide_context}

{lexicon_rules_text}

[CRITICAL INSTRUCTION - ZERO META COMMENTARY]
Never output developer metadata, markdown headers '##', explanation notes, character counters, or words like 'GEO', 'AEO', 'RAG'.
Output purely customer-facing content structured in 4 distinct sections.

Strict 4-Core Structure:
1. Official Product Title
(Under 100 characters. Noun Phrase: Brand 'Logicall Skin' + Product Title + Key Efficacy + Volume)

2. Core Value & Active Ingredient Summary
(5 concise bullet points containing quantitative metrics e.g. ppm, %, non-irritation score 0.00):
- Brand: Logicall Skin
- Core Actives & Concentration: (e.g. Multi-Vitamin 10% / 100,000ppm, Aquatide 3%)
- Key Benefits: (Efficacy claims strictly compliant with target country cosmetics law)
- Texture & Absorption: (Hydra-watery, non-greasy, fast-absorbing)
- Skin Compatibility: (Dermatologist tested, 0.00 irritation index)

3. Product Specifications & Comparison Table
(Output a clean HTML <table> comparing Logicall Skin vs Generic Market Standard):
<table>
  <tr><th>Dimensions</th><th>Logicall Skin</th><th>Standard Market Benchmark</th></tr>
  <tr><td>Active Concentration</td><td>High-Potency Multi-Vitamin 100,000ppm (10%)</td><td>Diluted extract 1,000~5,000ppm</td></tr>
  <tr><td>Patented Science</td><td>Aquatide 5000 (3%)</td><td>Generic purified water base</td></tr>
  <tr><td>Formula Stability</td><td>High-stability oxidation-free formula</td><td>Prone to discoloration / oxidation</td></tr>
  <tr><td>Irritation Index</td><td>0.00 Low-Irritation Certified</td><td>May cause stinging or redness</td></tr>
</table>

4. Product Usage Guide & Frequently Asked Questions (FAQ)
(5 high-conversion B2C customer FAQs):
Q1: When should I apply this serum?
A: ...
Q2: Is it suitable for sensitive skin?
A: ...
Q3: How does the active formula benefit the skin?
A: ...
Q4: What is the main efficacy of the Multi-Vitamin complex?
A: ...
Q5: Can I layer this with other skincare products?
A: ...

Generate the complete 4-Core content in {target_lang_desc} now.
"""
    try:
        resp = await client.aio.models.generate_content(
            model=MODEL_PRO,
            contents=[prompt],
            config=types.GenerateContentConfig(
                temperature=0.6,
                top_p=0.9,
                max_output_tokens=8192
            )
        )
        content_text = resp.text.strip()
    except Exception as e:
        print(f"  ⚠️ [WARN] SEO 텍스트 생성 중 오류 발생 -> 기본 템플릿 대체: {e}")
        content_text = f"""1. Logicall Skin {product_name} Multi Vitamin Daily Care Serum 50ml

2. Core Value & Active Ingredient Summary
- Brand: Logicall Skin
- Core Actives & Concentration: High-Potency Multi-Vitamin 10% (100,000ppm) & Aquatide 5000 (3%)
- Key Benefits: Visibly brightens, refines texture, and reinforces natural moisture barrier
- Texture & Absorption: Refreshingly lightweight hydra-watery formula with instant absorption
- Skin Compatibility: Dermatologist-tested, 0.00 skin irritation index suitable for daily use

3. Product Specifications & Comparison Table
<table>
  <tr><th>Specification</th><th>Logicall Skin Multi-Vitamin Serum</th><th>Standard Vitamin Serum</th></tr>
  <tr><td>Active Concentration</td><td>Multi-Vitamin Complex 100,000ppm (10.0%)</td><td>1,000 ~ 5,000ppm</td></tr>
  <tr><td>Patented Technology</td><td>Aquatide 5000 30,000ppm (3.0%)</td><td>Purified water base</td></tr>
  <tr><td>Formula Stability</td><td>High stability against air and light oxidation</td><td>Vulnerable to discoloration</td></tr>
  <tr><td>Skin Irritation Score</td><td>0.00 (Certified Low-Irritation)</td><td>May cause stinging sensation</td></tr>
</table>

4. Product Usage Guide & Frequently Asked Questions (FAQ)
Q1: When should I apply this serum?
A: Apply 3-4 drops evenly morning and evening after cleansing and toner.
Q2: Is it suitable for sensitive skin?
A: Yes, it is dermatologist-tested with a 0.00 skin irritation index.
Q3: How does Aquatide benefit the skin?
A: It reinforces the natural moisture barrier and revitalizes skin appearance.
Q4: What is the main efficacy of the Multi-Vitamin complex?
A: It provides deep hydration for a resilient-looking complexion and combats the signs of premature aging.
Q5: Can I layer this with other skincare products?
A: Yes, its fast-absorbing texture layers smoothly under creams and sunscreens.
"""

    # 4. 결정론적 법무 후처리 게이트 통과 (apply_deterministic_qa_overrides)
    if target_lang == "EN":
        for b_pat, p_val in [
            (r"\bnutrients for cellular vitality\b", "hydration for a resilient-looking complexion"),
            (r"\bcellular vitality\b", "resilient-looking complexion"),
            (r"\breinforces cellular resilience\b", "reinforces the skin's natural moisture barrier"),
            (r"\bcellular resilience\b", "skin's natural moisture barrier"),
            (r"\bcellular metabolism\b", "natural skin vitality"),
            (r"\bcombats premature aging\b", "combats the signs of premature aging"),
            (r"\bComplex skin issues\b", "Multiple skin concerns"),
            (r"\bTroubled skin\b", "Blemish-prone skin"),
            (r"\bcellular autophagy\b", "targeted skin nourishment")
        ]:
            content_text = re.sub(b_pat, p_val, content_text, flags=re.IGNORECASE)

    # 5. 4종 멀티 포맷 일괄 익스포트 (TXT, HTML, DOCX, MD)
    txt_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO.txt"
    html_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO_VIEWER.html"
    docx_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO.docx"
    md_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO.md"

    txt_path = os.path.join(target_dir, txt_filename)
    html_path = os.path.join(target_dir, html_filename)
    docx_path = os.path.join(target_dir, docx_filename)
    md_path = os.path.join(target_dir, md_filename)

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(content_text)
    print(f"  📄 [TXT 저장 완료]: {txt_path}")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(f"# {product_name} ({target_lang}) SEO/GEO/AEO PDP 원고\n\n" + content_text)
    print(f"  📑 [MD 저장 완료]: {md_path}")

    _generate_web_copier_html_file(f"{product_name} ({target_lang})", content_text, html_path, target_lang=target_lang)
    print(f"  🌐 [HTML 뷰어 저장 완료]: {html_path}")

    _generate_docx_file(f"{product_name} ({target_lang})", content_text, docx_path, target_lang=target_lang)
    print(f"  📄 [DOCX 서식 문서 저장 완료]: {docx_path}")


def _generate_docx_file(title: str, text_content: str, out_docx_path: str, target_lang: str = "EN"):
    """MS Word 서식(.docx)으로 4-Core 상세페이지 완성 원고를 렌더링합니다."""
    try:
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        doc = docx.Document()
        for s in doc.sections:
            s.top_margin = Inches(0.8)
            s.bottom_margin = Inches(0.8)
            s.left_margin = Inches(0.8)
            s.right_margin = Inches(0.8)

        # Title Header
        p_title = doc.add_paragraph()
        r_t = p_title.add_run(f"🛒 {title} - E-Commerce PDP Master Copy")
        r_t.font.name = "맑은 고딕"
        r_t.font.size = Pt(16)
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(16, 44, 87)

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(14)
        r_sub = p_sub.add_run("이 문서는 각국 전자상거래 플랫폼(스마트스토어, 쿠팡, 아마존, 큐텐 등) 상세페이지 에디터에 그대로 복사하여 등록할 수 있는 100% 고객 노출용 원고입니다.")
        r_sub.font.name = "맑은 고딕"
        r_sub.font.size = Pt(9.5)
        r_sub.font.color.rgb = RGBColor(100, 116, 139)

        # Parse sections
        lines = [l.strip() for l in text_content.splitlines() if l.strip()]
        cur_sec = 0
        sec_buffers = {1: [], 2: [], 3: [], 4: []}

        for l in lines:
            if l.startswith("1."):
                cur_sec = 1
                sec_buffers[1].append(l)
            elif l.startswith("2."):
                cur_sec = 2
                sec_buffers[2].append(l)
            elif l.startswith("3."):
                cur_sec = 3
                sec_buffers[3].append(l)
            elif l.startswith("4."):
                cur_sec = 4
                sec_buffers[4].append(l)
            else:
                if cur_sec in sec_buffers:
                    sec_buffers[cur_sec].append(l)

        # Section 1: Title
        h1 = doc.add_heading(level=1)
        r_h1 = h1.add_run("1. 공식 상품명 (Official Product Title)")
        r_h1.font.name = "맑은 고딕"
        r_h1.font.size = Pt(12.5)
        r_h1.font.bold = True
        r_h1.font.color.rgb = RGBColor(37, 99, 235)

        t_val = " ".join([l for l in sec_buffers[1] if not l.startswith("1.")]).strip()
        if not t_val and sec_buffers[1]:
            t_val = sec_buffers[1][-1].replace("1.", "").strip()
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_after = Pt(10)
        r_tv = p_t.add_run(t_val)
        r_tv.font.name = "맑은 고딕"
        r_tv.font.size = Pt(10.5)
        r_tv.font.bold = True

        # Section 2: Summary
        h2 = doc.add_heading(level=1)
        r_h2 = h2.add_run("2. 핵심 가치 및 제품 안내 (Core Value & Summary)")
        r_h2.font.name = "맑은 고딕"
        r_h2.font.size = Pt(12.5)
        r_h2.font.bold = True
        r_h2.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[2]:
            if l.startswith("2."):
                continue
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(l.lstrip("- •*"))
            r.font.name = "맑은 고딕"
            r.font.size = Pt(9.5)

        # Section 3: Comparison Table
        h3 = doc.add_heading(level=1)
        r_h3 = h3.add_run("3. 제품 상세 스펙 비교 (Product Specifications & Comparison Table)")
        r_h3.font.name = "맑은 고딕"
        r_h3.font.size = Pt(12.5)
        r_h3.font.bold = True
        r_h3.font.color.rgb = RGBColor(37, 99, 235)

        raw_s3_text = "\n".join(sec_buffers[3])
        if "<table>" in raw_s3_text:
            rows_data = []
            tr_matches = re.findall(r'<tr>(.*?)</tr>', raw_s3_text, flags=re.DOTALL)
            for tr in tr_matches:
                cols = re.findall(r'<t[hd]>(.*?)</t[hd]>', tr, flags=re.DOTALL)
                if cols:
                    rows_data.append([re.sub(r'<[^>]+>', '', c).strip() for c in cols])
            
            if rows_data:
                col_cnt = max(len(r) for r in rows_data)
                tbl = doc.add_table(rows=len(rows_data), cols=col_cnt)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r_idx, r_cols in enumerate(rows_data):
                    for c_idx, c_val in enumerate(r_cols):
                        cell = tbl.rows[r_idx].cells[c_idx]
                        p = cell.paragraphs[0]
                        r = p.add_run(c_val)
                        r.font.name = "맑은 고딕"
                        r.font.size = Pt(9)
                        if r_idx == 0:
                            r.font.bold = True
                            tcPr = cell._element.get_or_add_tcPr()
                            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
                            tcPr.append(shd)
        else:
            for l in sec_buffers[3]:
                if not l.startswith("3."):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    r = p.add_run(l)
                    r.font.name = "맑은 고딕"
                    r.font.size = Pt(9.5)

        # Section 4: FAQ
        h4 = doc.add_heading(level=1)
        r_h4 = h4.add_run("4. 자주 묻는 질문 (Frequently Asked Questions)")
        r_h4.font.name = "맑은 고딕"
        r_h4.font.size = Pt(12.5)
        r_h4.font.bold = True
        r_h4.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[4]:
            if l.startswith("4."):
                continue
            if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = RGBColor(30, 58, 138)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(51, 65, 85)

        doc.save(out_docx_path)
    except Exception as de:
        print(f"  ⚠️ [DOCX 생성 폴백]: {de}", flush=True)


def _generate_web_copier_html_file(title: str, text_content: str, out_html_path: str, target_lang: str = "EN"):
    """지마켓/스마트스토어/쿠팡/해외몰 등 전 이커머스 플랫폼에 맞춘 4-Core '에디터에 붙여넣을 텍스트 복사' 및 'HTML로 텍스트 복사' 뷰어를 생성합니다."""
    
    section_titles = {
        "EN": {
            "sec2": "2. Core Value & Active Ingredient Summary",
            "sec3": "3. Product Specifications & Comparison Table",
            "sec4": "4. Product Usage Guide & Frequently Asked Questions (FAQ)"
        },
        "JP": {
            "sec2": "2. コアバリュー＆成分サイエンス要約",
            "sec3": "3. 製品仕様・他社比較テーブル",
            "sec4": "4. 使用ガイド＆よくある質問 (FAQ)"
        },
        "CN": {
            "sec2": "2. 核心价值与成分科技摘要",
            "sec3": "3. 产品规格与竞品对比表",
            "sec4": "4. 商品使用指南与常见问题解答 (FAQ)"
        },
        "TW": {
            "sec2": "2. 核心價值與成分科技摘要",
            "sec3": "3. 產品規格與競品對比表",
            "sec4": "4. 使用指南與常見問題解答 (FAQ)"
        },
        "KR": {
            "sec2": "2. 제품 핵심 안내 및 성분 요약",
            "sec3": "3. 제품 상세 비교 스펙 테이블",
            "sec4": "4. 자주 묻는 질문 (FAQ)"
        }
    }
    
    detected_lang = target_lang.upper() if target_lang else "EN"
    titles_map = section_titles.get(detected_lang, section_titles["EN"])
    sec2_heading = titles_map["sec2"]
    sec3_heading = titles_map["sec3"]
    sec4_heading = titles_map.get("sec4", "4. FAQ")

    lines = [l.strip() for l in text_content.splitlines() if l.strip()]
    
    s1_lines = []
    s2_lines = []
    s3_lines = []
    s4_lines = []
    cur_sec = 0
    for l in lines:
        if l.startswith("1."):
            cur_sec = 1
            s1_lines.append(l)
        elif l.startswith("2."):
            cur_sec = 2
            s2_lines.append(l)
        elif l.startswith("3."):
            cur_sec = 3
            s3_lines.append(l)
        elif l.startswith("4."):
            cur_sec = 4
            s4_lines.append(l)
        else:
            if cur_sec == 1:
                s1_lines.append(l)
            elif cur_sec == 2:
                s2_lines.append(l)
            elif cur_sec == 3:
                s3_lines.append(l)
            elif cur_sec == 4:
                s4_lines.append(l)

    # 1. Section 1 (Title)
    s1_clean = " ".join([l for l in s1_lines if not l.startswith("1.")]).strip()
    if not s1_clean and s1_lines:
        s1_clean = s1_lines[-1].replace("1.", "").strip()

    # 2. Section 2 (Summary)
    s2_clean_items = [l for l in s2_lines if not l.startswith("2.")]
    s2_text_formatted = f"{sec2_heading}\n\n" + "\n\n".join(s2_clean_items)
    
    s2_html_items = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s2_html_items.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>🔬 {sec2_heading}</h3>")
    for l in s2_clean_items:
        if ":" in l or "：" in l:
            k, v = re.split(r'[:：]', l, 1)
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'><strong style='color:#1e3a8a;'>{k.strip()}:</strong> {v.strip()}</p>")
        else:
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'>{l.strip()}</p>")
    s2_html_items.append("</div>")
    s2_html = "\n".join(s2_html_items)

    # 3. Section 3 (Table)
    s3_clean_items = [l for l in s3_lines if not l.startswith("3.")]
    s3_raw_block = "\n".join(s3_clean_items)
    s3_text_formatted = f"{sec3_heading}\n\n" + s3_raw_block
    
    if "<table>" in s3_raw_block:
        styled_table = s3_raw_block.replace("<table>", "<table style='width:100%; border-collapse:collapse; margin:10px 0; font-size:13.5px;'>")
        styled_table = styled_table.replace("<th>", "<th style='background:#f1f5f9; padding:10px 12px; border:1px solid #cbd5e1; color:#1e3a8a; text-align:left; font-weight:bold;'>")
        styled_table = styled_table.replace("<td>", "<td style='padding:9px 12px; border:1px solid #cbd5e1; color:#334155;'>")
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n{styled_table}\n</div>"
    else:
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n<pre style='white-space:pre-wrap; font-family:inherit; font-size:13.5px;'>{s3_raw_block}</pre>\n</div>"

    # 4. Section 4 (FAQ)
    s4_clean_items = [l for l in s4_lines if not l.startswith("4.")]
    s4_text_blocks = []
    s4_html_blocks = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s4_html_blocks.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 12px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>💬 {sec4_heading}</h3>")
    
    cur_q = ""
    cur_answers = []
    for l in s4_clean_items:
        if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
            if cur_q:
                q_text = cur_q
                a_text = "\n".join(cur_answers)
                s4_text_blocks.append(f"{q_text}\n{a_text}")
                s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
                for ans in cur_answers:
                    s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
            cur_q = l.strip()
            cur_answers = []
        else:
            sentences = [s.strip() for s in re.split(r'(?<=[。！？\.\?!])\s*', l) if s.strip()]
            for s in sentences:
                cur_answers.append(s)
                
    if cur_q:
        q_text = cur_q
        a_text = "\n".join(cur_answers)
        s4_text_blocks.append(f"{q_text}\n{a_text}")
        s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
        for ans in cur_answers:
            s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
    s4_html_blocks.append("</div>")

    s4_text_formatted = f"{sec4_heading}\n\n" + "\n\n".join(s4_text_blocks)
    s4_html = "\n".join(s4_html_blocks)

    full_html_code = f"""<!-- 다국어 E-Commerce 상세페이지 4-Core 마이크로-써머리 & 비교표 & FAQ -->
<div style="font-family:'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; max-width:860px; margin:0 auto; padding:10px 0; color:#1e293b;">
{s2_html}
{s3_html}
{s4_html}
</div>"""

    html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; background: #f1f5f9; color: #0f172a; padding: 25px; margin: 0; line-height: 1.6; }}
  .container {{ max-width: 960px; margin: 0 auto; background: #ffffff; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.06); padding: 35px; border: 1px solid #cbd5e1; }}
  h1 {{ font-size: 22px; color: #0f172a; border-bottom: 2px solid #2563eb; padding-bottom: 12px; margin-top: 0; display: flex; align-items: center; gap: 8px; }}
  .guide-box {{ background: #eff6ff; border: 1px solid #bfdbfe; border-left: 5px solid #2563eb; padding: 18px 20px; border-radius: 8px; margin-bottom: 25px; font-size: 14px; color: #1e40af; line-height: 1.8; }}
  .guide-box strong {{ color: #1e3a8a; }}
  .card {{ background: #f8fafc; border: 1px solid #cbd5e1; border-radius: 10px; padding: 22px; margin-bottom: 25px; }}
  .card-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 14px; border-bottom: 1px solid #e2e8f0; padding-bottom: 10px; }}
  .card-title {{ font-size: 16px; font-weight: bold; color: #1e3a8a; }}
  .btn-group {{ display: flex; gap: 8px; }}
  .copy-btn {{ background: #2563eb; color: #ffffff; border: none; padding: 9px 16px; border-radius: 6px; cursor: pointer; font-size: 13px; font-weight: 700; transition: all 0.2s; box-shadow: 0 2px 4px rgba(37,99,235,0.2); display: inline-flex; align-items: center; gap: 6px; }}
  .copy-btn:hover {{ background: #1d4ed8; transform: translateY(-1px); }}
  .copy-btn.html-mode-btn {{ background: #059669; box-shadow: 0 2px 4px rgba(5,150,105,0.25); }}
  .copy-btn.html-mode-btn:hover {{ background: #047857; }}
  .full-btn {{ background: #7c3aed; padding: 12px 24px; font-size: 15px; width: 100%; justify-content: center; margin-bottom: 20px; box-shadow: 0 3px 6px rgba(124,58,237,0.25); }}
  .full-btn:hover {{ background: #6d28d9; }}
  .text-area {{ width: 100%; box-sizing: border-box; background: #ffffff; border: 1px solid #cbd5e1; border-radius: 6px; padding: 14px 16px; font-size: 13.5px; line-height: 1.75; color: #1e293b; font-family: 'Malgun Gothic', 'Segoe UI', monospace; resize: vertical; outline: none; }}
  .text-area:focus {{ border-color: #2563eb; box-shadow: 0 0 0 3px rgba(37,99,235,0.15); }}
  .toast {{ position: fixed; bottom: 30px; right: 30px; background: #0f172a; color: #ffffff; padding: 14px 28px; border-radius: 8px; font-size: 14px; font-weight: 600; display: none; z-index: 1000; box-shadow: 0 6px 20px rgba(0,0,0,0.25); }}
</style>
</head>
<body>

<div class="container">
  <h1>🌐 {title}</h1>
  
  <div class="guide-box">
    📢 <strong>쇼핑몰 등록 방식별 2대 원클릭 복사 기능 안내:</strong><br>
    • <strong>1. [📋 에디터에 붙여넣을 텍스트 복사] (파란색 버튼)</strong>: 지마켓/스마트스토어/쿠팡/해외몰 <strong>'에디터 작성'</strong> 화면에 붙여넣을 때 사용합니다.<br>
    • <strong>2. [🌐 HTML로 텍스트 복사] (초록색 버튼)</strong>: <strong>'HTML 작성'</strong> 탭이나 HTML 직접 입력 모드에 붙여넣을 때 사용합니다.
  </div>

  <button class="copy-btn full-btn html-mode-btn" onclick="copyFromTextarea('full-html-ta', '🎉 전체 HTML 소스코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🚀 [HTML로 전체 일괄 복사] 4-Core 요약 + 비교표 + FAQ 전체 소스코드 복사</button>
  <textarea id="full-html-ta" style="display:none;">{full_html_code}</textarea>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📌 1. 공식 상품명 (Title)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec1-ta', '✅ 상품명이 복사되었습니다!')">📋 상품명 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec1-ta" rows="2" readonly>{s1_clean}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">🔬 2. 핵심 가치 및 5줄 마이크로 요약</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec2-ta', '✅ 5줄 요약 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec2-html-ta', '🌐 5줄 요약 HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec2-ta" rows="8" readonly>{s2_text_formatted}</textarea>
    <textarea id="sec2-html-ta" style="display:none;">{s2_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📊 3. 제품 상세 스펙 비교표 (HTML Table)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec3-ta', '✅ 비교표 텍스트가 복사되었습니다!')">📋 비교표 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec3-html-ta', '🌐 비교표 HTML 코드가 복사되었습니다!')">🌐 HTML로 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec3-ta" rows="10" readonly>{s3_text_formatted}</textarea>
    <textarea id="sec3-html-ta" style="display:none;">{s3_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">💬 4. 5대 핵심 FAQ & 상세 가이드</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec4-ta', '✅ FAQ 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec4-html-ta', '🌐 FAQ HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec4-ta" rows="18" readonly>{s4_text_formatted}</textarea>
    <textarea id="sec4-html-ta" style="display:none;">{s4_html}</textarea>
  </div>

</div>

<div class="toast" id="toast">✅ 클립보드에 복사되었습니다! 쇼핑몰 에디터에 붙여넣기(Ctrl+V) 하세요.</div>

<script>
function showToast(msg) {{
  const t = document.getElementById('toast');
  t.innerText = msg;
  t.style.display = 'block';
  setTimeout(() => {{ t.style.display = 'none'; }}, 3000);
}}

function copyFromTextarea(id, customMsg) {{
  const ta = document.getElementById(id);
  const text = ta.value;
  
  const temp = document.createElement('textarea');
  temp.value = text;
  temp.style.position = 'fixed';
  temp.style.left = '-9999px';
  document.body.appendChild(temp);
  temp.select();
  temp.setSelectionRange(0, 99999);
  
  try {{
    document.execCommand('copy');
    showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
  }} catch (err) {{
    if (navigator.clipboard) {{
      navigator.clipboard.writeText(text).then(() => {{
        showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
      }});
    }}
  }} finally {{
    document.body.removeChild(temp);
  }}
}}
</script>

</body>
</html>
"""
    with open(out_html_path, "w", encoding="utf-8") as f:
        f.write(html_content)


def extract_product_name_from_files(files: List[str], fallback_name: str) -> str:
    """파일명 목록이나 폴더명에서 상품 식별명을 깔끔하게 정제합니다."""
    clean_name = fallback_name
    clean_name = re.sub(r'^\(한국어\)[\s_-]*', '', clean_name)
    clean_name = re.sub(r'^썸네일[\s_-]*', '', clean_name)
    clean_name = clean_name.strip()
    return clean_name if clean_name else fallback_name


def find_target_leaf_folders(base_dir: str) -> List[Tuple[str, str]]:
    """이미지나 docx가 존재하는 실제 리프 폴더들을 탐색하여 (폴더경로, 상품명) 목록을 반환합니다."""
    leaf_dirs = []
    for root, dirs, files in os.walk(base_dir):
        valid_files = [f for f in files if f.lower().endswith(('.png', '.jpg', '.jpeg', '.webp', '.gif', '.docx')) and not f.startswith('~')]
        if valid_files:
            rel = os.path.relpath(root, base_dir)
            if rel == ".":
                pname = extract_product_name_from_files(valid_files, os.path.basename(base_dir))
            else:
                top_part = rel.split(os.sep)[0]
                pname = extract_product_name_from_files(valid_files, top_part)
            leaf_dirs.append((root, pname))
    return leaf_dirs


async def run_translation_batch(client: genai.Client, source_dir: str, target_lang: str, custom_target_dir: Optional[str] = None, custom_product_name: Optional[str] = None):
    """source_dir 내에 서브폴더나 중첩 폴더를 탐색하여 모든 대상 리프 폴더를 번역합니다."""
    leaf_folders = find_target_leaf_folders(source_dir)
    if not leaf_folders:
        print(f"⚠️ [WARNING] '{source_dir}' 폴더에 처리할 이미지나 DOCX 파일이 없습니다.")
        return

    for sdir, pname in leaf_folders:
        final_pname = custom_product_name if custom_product_name else pname
        await run_translation_batch_for_folder(client, sdir, target_lang, final_pname, custom_target_dir=custom_target_dir)


# =================================================================================
# 6. 메인 실행 및 대화형 사용자 질의응답
# =================================================================================
async def main_async():
    parser = argparse.ArgumentParser(description="multilingual_text_in_image_translatio_agy_sdk_core")
    parser.add_argument("--source", default=DEFAULT_INPUT_DIR, help="원본 이미지 디렉터리")
    parser.add_argument("--lang", choices=["EN", "JP", "CN", "TW", "ALL"], default=None, help="도착 언어 코드")
    parser.add_argument("--out_dir", default=None, help="결과물 저장 디렉터리 직접 지정 (선택)")
    parser.add_argument("--product_name", default=None, help="상품 식별명 직접 지정 (선택)")
    args = parser.parse_args()

    source_dir = os.path.abspath(args.source)
    os.makedirs(source_dir, exist_ok=True)

    chosen_lang = args.lang
    if not chosen_lang:
        image_count = 0
        for root, _, files in os.walk(source_dir):
            image_count += len([f for f in files if f.lower().endswith(('.png', '.jpg', '.jpeg', '.webp', '.gif'))])

        print("\n" + "=" * 76)
        print(" 🌐 multilingual_text_in_image_translatio_agy_sdk_core")
        print("=" * 76)
        print(f" 📂 [공통 인풋 폴더] : {source_dir}")
        print(f" 🖼️ [감지된 원본 이미지] : 총 {image_count}개")
        print("-" * 76)
        print(" 번역할 도착 언어를 선택하세요:")
        print("   [1] 🇺🇸 영어 (EN - Amazon / Shopee US 초월번역 + Montserrat)")
        print("   [2] 🇯🇵 일본어 (JP - Qoo10 Japan / 후생노동성 56종 약기법 준수)")
        print("   [3] 🇨🇳 중국어 간체 (CN - 중국 본토 신광고법 및 NMPA 규정 준수)")
        print("   [4] 🇹🇼 중국어 번체 (TW - 대만/홍콩 TFDA 규정 준수)")
        print("   [5] 🌐 전체 언어 일괄 번역 (EN -> JP -> CN -> TW 순차 실행)")
        print("   [Q] 종료 (Quit)")
        print("=" * 76)

        while True:
            choice = input(" 👉 번호를 입력하세요 (1/2/3/4/5/Q): ").strip().upper()
            if choice == "1":
                chosen_lang = "EN"
                break
            elif choice == "2":
                chosen_lang = "JP"
                break
            elif choice == "3":
                chosen_lang = "CN"
                break
            elif choice == "4":
                chosen_lang = "TW"
                break
            elif choice == "5":
                chosen_lang = "ALL"
                break
            elif choice in ["Q", "QUIT", "EXIT"]:
                print("번역 작업을 취소하고 종료합니다.")
                sys.exit(0)
            else:
                print(" ⚠️ 잘못된 입력입니다. 1, 2, 3, 4, 5 또는 Q를 입력하세요.")

    client = load_credentials()

    if chosen_lang == "ALL":
        for lang in ["EN", "JP", "CN", "TW"]:
            await run_translation_batch(client, source_dir, lang, custom_target_dir=args.out_dir, custom_product_name=args.product_name)
    else:
        await run_translation_batch(client, source_dir, chosen_lang, custom_target_dir=args.out_dir, custom_product_name=args.product_name)


def main():
    asyncio.run(main_async())

if __name__ == '__main__':
    main()


def main():
    asyncio.run(main_async())






