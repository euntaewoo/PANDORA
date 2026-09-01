import asyncio
import os
def get_recursive_files(base_dir):
    try:
        res = []
        for root, _, files in os.walk(base_dir):
            for f in files:
                res.append(os.path.relpath(os.path.join(root, f), base_dir))
        return res
    except Exception:
        return []

import sys
sys.path.insert(0, r"C:\Users\euntaewoo\Desktop\multilingual_text_in_image_translatio_agy_sdk\multilingual_text_in_image_translatio_agy_sdk_core")
from multilingual_transcreation_qa_evaluator_agy_sdk import evaluate_transcreation, generate_html_report

"""
===================================================================================
🇺🇸 EN_Text-In_Image_Translation_Engine_AGY_SDK.py
-----------------------------------------------------------------------------------
• Purpose: Two-Pass Multimodal Neural Inpainting Engine for English E-Commerce (Amazon/Shopee)
• Core Models:
    - Pass 1: gemini-3.1-pro-preview (Dual Mode: Transcreation KR->EN / Polishing EN->EN)
    - Pass 2: gemini-3.1-flash-image (Visual Inpainting & Typography Rendering)
• Standard Fonts:
    - Detail Page Main Images: 100% Montserrat (몬세라트 단일 서체 강제)
    - Notice Tables (고시정보표): Pretendard (render_notice_table_standard.py 독립 분리, 1열 295px, padding: 14px 12px, letter-spacing: -0.8px, 순수 가용폭 271px, word-break: keep-all)
• Resolution: Aspect Ratio Lock via Pillow LANCZOS Resampling (1:1 픽셀 동기화)
• Location: Google Cloud Vertex AI (location="global") Serverless Standard
===================================================================================
"""

import io
import json
import os
import re
import sys
import time
from typing import Any, Dict, List, Optional, Tuple

from google import genai
from google.genai import types
from PIL import Image

sys.stdout.reconfigure(encoding="utf-8")

# =================================================================================
# 1. 환경 설정 및 클라우드 인증 초기화
# =================================================================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, ".."))

MODEL_PRO = "gemini-3.1-pro-preview"
MODEL_FLASH_IMAGE = "gemini-3.1-flash-image"

# =================================================================================
# 1-1. 글로벌 컴플라이언스(법무) & 렉시콘 로더
# =================================================================================
GLOBAL_COMPLIANCE_SYSTEM_INSTRUCTION = """[SYSTEM INSTRUCTION: Global Cross-Border E-Commerce Compliance & Prestige Beauty Transcreation Expert]
당신은 미국 FDA(MoCRA), 일본 후생노동성(약기법), 중국 NMPA/신광고법, 대만 TFDA 규정을 완벽히 준수하는 15년 차 글로벌 뷰티 법무 감사관이자, 세포라(Sephora)·백화점 럭셔리 브랜드의 수석 카피라이터입니다.

[엄격 실행 4대 대원칙]
1. [의약품 오인 원천 차단]: 인체 구조, 생리적 기능, 세포(Cellular) 단위 클레임(cellular vitality, cellular resilience 등)을 100% 차단하고 피부 표면 미용적 개선(-looking, moisture barrier)으로 우회.
2. [타겟 권역 문화적 어댑테이션]: K-뷰티 콩글리시('Complex skin issues' -> 'Multiple skin concerns', 'Troubled skin' -> 'Blemish-prone skin') 전면 배제. 노화 서술 시 'combats the signs of premature aging'으로 징후 한정.
3. [디자인 & 레이아웃 최적화]: 백화점 럭셔리 브랜드 수준의 세련된 어휘로 초월번역.
4. [하이퍼파라미터 전역 고정]: temperature: 0.6, top_p: 0.9 유지.
"""

def load_en_compliance_lexicon() -> Dict[str, str]:
    fpath = os.path.join(PROJECT_ROOT, "00_공통자료", "compliance_lexicons", "en_fda_mocra_lexicon.json")
    replacements = {
        r"\bComplex skin issues\b": "Multiple skin concerns",
        r"\bcomplex skin issues\b": "multiple skin concerns",
        r"\bTroubled skin\b": "Blemish-prone skin",
        r"\btroubled skin\b": "blemish-prone skin",
        r"\bnutrients for cellular vitality\b": "hydration for a resilient-looking complexion",
        r"\bcellular vitality\b": "resilient-looking complexion",
        r"\breinforces cellular resilience\b": "reinforces the skin's natural moisture barrier",
        r"\bcellular resilience\b": "skin's natural moisture barrier",
        r"\bcombats premature aging\b": "combats the signs of premature aging",
        r"\bcombats aging\b": "combats the signs of aging",
        r"\bPrescribe\b": "Targeted Solution for",
        r"\bBio-Immunity\b": "Skin Defense",
        r"\bfed directly\b": "infused daily",
        r"\bKyel-Tan-Tone\b": "Texture, Elasticity & Luminosity",
    }
    if os.path.exists(fpath):
        try:
            with open(fpath, "r", encoding="utf-8") as f:
                data = json.load(f)
                for cat in data.get("categories", {}).values():
                    for it in cat.get("banned_terms", []):
                        b, p = it.get("banned", ""), it.get("preferred", "")
                        if b and p:
                            replacements[rf"\b{re.escape(b)}\b"] = p
        except Exception:
            pass
    return replacements



def load_credentials() -> genai.Client:
    """Vertex AI 서비스 계정 키 및 API 키를 탐색하여 genai.Client를 초기화합니다."""
    env_paths = [
        os.path.join(SCRIPT_DIR, ".env"),
        os.path.join(PROJECT_ROOT, ".env"),
    ]
    api_key = os.environ.get("GEMINI_API_KEY")
    gcp_json_key = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")

    for p in env_paths:
        if os.path.exists(p):
            with open(p, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line.startswith("#"):
                        continue
                    if line.startswith("GEMINI_API_KEY="):
                        api_key = line.split("=", 1)[1].strip()
                    elif line.startswith("GOOGLE_APPLICATION_CREDENTIALS="):
                        gcp_json_key = line.split("=", 1)[1].strip().strip('"').strip("'")

    key_candidates = [
        gcp_json_key,
        os.path.join(PROJECT_ROOT, "00_공통자료", "APIs_KEY", "인증키_및_계정", "김차장_vertex api_key", "vertex_ai_auth_key.json"),
        os.path.join(PROJECT_ROOT, "00_공통자료", "인증키_및_계정", "김차장_vertex api_key", "vertex_ai_auth_key.json"),
    ]

    for kpath in key_candidates:
        if kpath and os.path.exists(kpath) and kpath.endswith(".json"):
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = kpath
            print(f"[AUTH] Vertex AI 서비스 계정 키 감지: {kpath}")
            with open(kpath, "r", encoding="utf-8") as f:
                key_data = json.load(f)
                project_id = key_data.get("project_id")
            client = genai.Client(vertexai=True, project=project_id, location="global")
            print(f"[AUTH SUCCESS] Vertex AI Client 연결 완료 (Project: {project_id}, Location: global)")
            return client

    if api_key:
        if api_key.startswith("AQ."):
            print("[AUTH] Agent Platform API 키 감지 -> Vertex AI 모드로 전환")
            return genai.Client(vertexai=True, api_key=api_key)
        print("[AUTH] Gemini API 키 인증 모드 연결")
        return genai.Client(api_key=api_key)

    print("[ERROR] GEMINI_API_KEY 또는 GOOGLE_APPLICATION_CREDENTIALS가 설정되지 않았습니다.")
    sys.exit(1)


# =================================================================================
# 2. Pass 1 & Pass 2 프롬프트 정의
# =================================================================================
PASS1_PROMPT = """
[MANDATORY COMPLIANCE HARD RULES - 5대 금지 표현 강제 대체]
1. 'Complex skin issues' -> MUST USE: 'Multiple skin concerns'
2. 'Troubled skin' -> MUST USE: 'Blemish-prone skin'
3. 'cellular vitality / nutrients for cellular vitality' -> MUST USE: 'hydration for a resilient-looking complexion'
4. 'cellular resilience / reinforces cellular resilience' -> MUST USE: 'reinforces the skin's natural moisture barrier'
5. 'combats premature aging' -> MUST USE: 'combats the signs of premature aging'

[SYSTEM PROMPT] Global Luxury Beauty Transcreation & Compliance Expert (English Engine)

## 1. 시스템 역할 및 콘셉트 (Role & Context)
당신은 에스티로더, 랑콤, 시슬리, SK-II 등 글로벌 하이엔드 코스메틱 브랜드의 해외 진출을 총괄하는 10년 차 수석 크리에이티브 디렉터이자 세포라(Sephora) 전문 엘리트 카피라이터입니다.
단순 직역을 배제하고, 미국 및 글로벌 프레스티지 뷰티 고객의 구매 심리를 사로잡는 세련된 '초월번역(Transcreation)'을 수행하세요.

## 2. 단계별 모드 자동 감지
1. 이미지 내 텍스트에 '한국어'가 포함되어 있다면 -> mode: "TRANSLATE_KR_TO_EN"
2. 이미지 내 텍스트가 이미 '영어'로만 되어 있다면 -> mode: "POLISH_EN_TO_EN"

## 3. 초월번역 및 교정 핵심 원칙
1. [기계적 직역 및 부사 금지]
   - 'Definitely', 'Truly', 'Really', 'Certainly' 등 어색한 감정 부사 직역을 전면 금지하고, 럭셔리 뷰티 전문 능동태 동사/형용사로 재창조하십시오.
2. [자연스러운 구문 결속 및 활성 성분 연결]
   - "10% LiftDerm" 등 활성 성분 수치가 문맥과 끊기지 않고 제품 효능 및 서사로 매끄럽게 연결되도록 문장 구조를 재조정하십시오.
3. [4대 기능성 뷰티 전문 어휘 사전]
   - 피부 속/기저층: Deep within the skin layers / Deep within the dermal matrix
   - 토탈 케어/멀티 코렉티브: Multi-Corrective Repair / Total Revitalizing Care
   - 탄력 복원/강화: Rebuilding skin elasticity / Restoring visible firmness
   - 눈가 잔주름/건조주름: Fine lines and wrinkles / Micro-creases
4. [독자 성분명 영문 보존]
   - 'LiftDerm', 'Lifting Logic for eye' 등 글로벌 독자 성분명/브랜드명은 영문 그대로 유지하되 문맥과 완벽히 융합하십시오.
   - 제품 본품 용기/단상자 표면 인쇄 영문/로고는 수정 대상에서 제외하십시오.

## 4. 광고 법규 및 규제 준수 (Regulatory Compliance & Guardrails)
1. [절대적/과대 표현 전면 금지 (Ban on Absolute Claims)]
   - 'World's First', 'No.1', 'Best', 'The Ultimate' 등 입증되지 않은 절대적 표현 사용 전면 금지.
   - 반드시 'Innovative formula engineered for delicate eye areas', 'Advanced Multi-Corrective Solution', 'Targeted Precision Care' 등 프리미엄 혁신 표현으로 순화하십시오.
2. [의료 시술 오인 금지 및 4대 안전 동사 (Compliance-Safe Verbs)]
   - '주름 완전 박멸(Wrinkle-free)', '보톡스/필러 효과' 등 의료 시술 연상 및 세포 치료/재생 오인 단어 전면 배제.
   - 반드시 **`Smooth` (抚平/撫平)**, **`Diminish` (淡化)**, **`Alleviate` (舒缓/舒緩)**, **`Care / Repair` (修护/修護)** 4대 컴플라이언스 안전 동사를 사용하여 표현하십시오.

## 5. 고시정보 표 강제 표준 매핑
만약 이미지 내용이 '고시정보(Notice Table, Product Specifications)' 표라면, 다음의 표준 명칭으로 강제 매핑하십시오:
1) 용량 또는 중량 -> Size / Net Wt.
2) 제품 주요 사양 -> Skin Type
3) 사용기한 또는 개봉 후 사용기간 -> Shelf Life / PAO
4) 사용방법 -> Directions
5) 화장품제조업자 및 책임판매업자 -> Manufacturer / Distributed by
6) 제조국 -> Country of Origin
7) 전성분 -> Ingredients
8) 기능성화장품 심사 유무 -> Functional Cosmetics Review Status: Completed Functional Cosmetics Review (or Report) with the Ministry of Food and Drug Safety (MFDS, Republic of Korea) in accordance with the Cosmetics Act
9) 사용할 때의 주의사항 -> Precautions for Use (전문의 상담, 상처부위 자제, 직사광선/어린이 보관 3대 법정 조항)
10) 품질보증기준 -> Quality Assurance Standard (Compensation will be provided in accordance with the Fair Trade Commission's Consumer Dispute Settlement Standards)
11) 소비자상담 -> Customer Service: +82-2-6743-3206

출력은 반드시 아래 JSON 스키마를 엄격히 준수하십시오:
```json
{
  "detected_mode": "TRANSLATE_KR_TO_EN 또는 POLISH_EN_TO_EN",
  "translation_map": [
    {
      "original_text": "원본 텍스트(한글 또는 어색한 기존 영문)",
      "corrected_en": "최종 교정/초월번역된 프리미엄 영문 카피"
    }
  ]
}
```
"""

PASS2_PROMPT_TEMPLATE = """
당신은 글로벌 럭셔리 뷰티(Sephora, Amazon US) 이미지 로컬라이징 최고 전문가입니다.
첨부된 원본 이미지에서 기존의 원본 텍스트를 감쪽같이 지우고, 교정/번역된 영문 데이터를 바탕으로 완벽하게 재렌더링하세요.

[시각적 렌더링 엄격 규칙]
1. (TEXT ERASING) 원본의 기존 텍스트('original_text')를 원래 배경색과 완벽히 블렌딩하여 지울 것.
2. (NEW COPY RENDERING) 지워진 그 자리에 [매핑 데이터 JSON]의 'corrected_en' 영문 텍스트만 정확한 위치에 렌더링할 것.
3. (FONT & TYPOGRAPHY) 모든 영문 텍스트는 100% 오직 'Montserrat (몬세라트)' 폰트만을 유일한 표준 서체로 적용하여 렌더링할 것. (Montserrat 단일 서체 강제 / 타 서체 혼용 절대 금지)
4. (FULL INPAINTING NO PATCHING) 전체 이미지를 매끄럽게 재렌더링하여 원본과 동일한 해상도/비율을 100% 유지할 것.
5. (PACKAGE PRESERVATION) 제품 본품 용기/패키지에 인쇄된 로고 및 문구는 100% 원본 유지할 것.

[매핑 데이터 JSON]
{json_data}
"""


# =================================================================================
# 3. 헬퍼 함수
# =================================================================================
def natural_sort_key(s: str) -> List[Any]:
    """파일명 내부의 숫자를 자연스럽게 인식하여 정렬하는 키 함수입니다."""
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r"(\d+)", s)]


def parse_arguments() -> Tuple[str, str]:
    """커맨드라인 파라미터를 파싱하여 입력 및 출력 디렉토리를 반환합니다."""
    if len(sys.argv) > 1:
        src = sys.argv[1]
        if len(sys.argv) > 2:
            tgt = sys.argv[2]
        else:
            tgt = os.path.join(
                os.path.dirname(src),
                os.path.normpath(src).split(os.sep)[-1] + "_EN_Translated",
            )
    else:
        src = os.path.join(SCRIPT_DIR, "input")
        tgt = os.path.join(SCRIPT_DIR, "output")

    os.makedirs(src, exist_ok=True)
    os.makedirs(tgt, exist_ok=True)
    return src, tgt


def extract_image_bytes(response: Any) -> Optional[bytes]:
    """Gemini API 응답 객체에서 렌더링된 이미지 바이트를 안전하게 추출합니다."""
    if not hasattr(response, "candidates") or not response.candidates:
        return None

    for cand in response.candidates:
        if hasattr(cand, "content") and hasattr(cand.content, "parts"):
            for part in cand.content.parts:
                if hasattr(part, "inline_data") and part.inline_data and part.inline_data.data:
                    return part.inline_data.data
                if hasattr(part, "image") and part.image and part.image.image_bytes:
                    return part.image.image_bytes
    return None


# =================================================================================
# 4. 메인 엔진 파이프라인
# =================================================================================
async def main_async():
        client = load_credentials()
        source_dir, target_dir = parse_arguments()

        print("\n=======================================================")
        print("🇺🇸 EN_Text-In_Image_Translation_Engine_V1 가동")
        print(f"• 입력 경로: {source_dir}")
        print(f"• 출력 경로: {target_dir}")
        print(f"• 메인 폰트: Montserrat (100% 단일 서체 강제)")
        print("=======================================================\n")

        valid_extensions = (".png", ".jpg", ".jpeg", ".jfif", ".gif", ".webp")
        raw_files = [f for f in get_recursive_files(source_dir) if f.lower().endswith(valid_extensions)]
        targets = sorted(raw_files, key=natural_sort_key)

        if not targets:
            print(f"[WARNING] '{source_dir}' 폴더에 처리할 이미지가 없습니다.")
            return

        all_translations = []

        for filename in targets:
            if "_수정번역" in filename or filename.endswith(".txt") or filename.endswith(".md"):
                continue

            in_path = os.path.join(source_dir, filename)
            out_name = f"{os.path.splitext(filename)[0]}_수정번역.png"
            out_path = os.path.join(target_dir, out_name)

            if os.path.exists(out_path):
                print(f"[SKIP] 이미 완료된 파일입니다: {filename}")
                continue

            print(f"\n[RENDER] 처리 시작: {filename}")

            try:
                original_image = Image.open(in_path)
                original_image.load()
                orig_w, orig_h = original_image.size
            except Exception as e:
                print(f"  -> [ERROR] 이미지 로드 실패 ({filename}): {e}")
                continue

            # -------------------------------------------------------------
            # PASS 1: 언어 자동 감지 및 초월번역 / 교정 매핑 생성 (Pro 모델)
            # -------------------------------------------------------------
            print("  -> [PASS 1] 텍스트 OCR 및 언어 감지, 영문 초월번역 생성 중...")
            mapping_data_str = ""
            try:
                response_p1 = await client.aio.models.generate_content(
                    model=MODEL_PRO,
                    contents=[original_image, PASS1_PROMPT],
                    config=types.GenerateContentConfig(
                        system_instruction=GLOBAL_COMPLIANCE_SYSTEM_INSTRUCTION,
                        response_mime_type="application/json",
                        temperature=0.6,
                        top_p=0.9,
                        max_output_tokens=8192
                    ),
                )
                mapping_data_str = response_p1.text.strip()
                parsed_json = json.loads(mapping_data_str)
                mode = parsed_json.get("detected_mode", "UNKNOWN")
                map_count = len(parsed_json.get("translation_map", []))
                print(f"  -> [PASS 1 SUCCESS] 감지 모드: {mode} (매핑 항목: {map_count}개)")

                if "translation_map" in parsed_json:
                    en_reps = load_en_compliance_lexicon()
                    for item in parsed_json["translation_map"]:
                        item["source_file"] = filename
                        item["mode"] = mode
                        cor_en = item.get("corrected_en", "")
                        orig_cor = cor_en
                        for pat, rep in en_reps.items():
                            cor_en = re.sub(pat, rep, cor_en, flags=re.IGNORECASE)
                        if cor_en != orig_cor:
                            item["corrected_en"] = cor_en
                            print(f"     ⚡ [EN 법규 자동 보정] `{orig_cor[:35]}` ➔ `{cor_en[:35]}`")
                    all_translations.extend(parsed_json["translation_map"])
            except Exception as e:
                print(f"  -> [PASS 1 ERROR] 매핑 생성 실패: {e}")
                continue

            # -------------------------------------------------------------
            # PASS 2: 영문 이미지 인페인팅 렌더링 (Flash-Image 모델)
            # -------------------------------------------------------------
            print("  -> [PASS 2] Montserrat 영문 타이포그래피 인페인팅 렌더링 중...")
            max_retries = 3
            final_prompt = PASS2_PROMPT_TEMPLATE.replace("{json_data}", mapping_data_str)

            for attempt in range(max_retries):
                try:
                    response_p2 = await client.aio.models.generate_content(
                        model=MODEL_FLASH_IMAGE,
                        contents=[final_prompt, original_image],
                        config=types.GenerateContentConfig(
                            response_modalities=["IMAGE"],
                            temperature=0.6,
                            top_p=0.9
                        )
                    )

                    img_bytes = extract_image_bytes(response_p2)
                    if img_bytes:
                        img = Image.open(io.BytesIO(img_bytes))
                        # 원본 해상도 1:1 강제 일치 (Aspect Ratio Lock)
                        img = img.resize((orig_w, orig_h), Image.Resampling.LANCZOS)
                        img.save(out_path, format="PNG")
                        print(f"  -> [PASS 2 SUCCESS] {out_name} 저장 완료 ({orig_w}x{orig_h}px)!")
                        break
                    else:
                        print("  -> [RETRY] Pass 2 이미지 응답 없음, 10초 대기 후 재시도...")
                        time.sleep(10)
                except Exception as e:
                    print(f"  -> [PASS 2 ERROR] 렌더링 에러 (시도 {attempt+1}/{max_retries}): {e}")
                    if "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                        wait_sec = 25 * (attempt + 1)
                        print(f"  -> [QUOTA WAIT] 429 쿼터 대기 ({wait_sec}초)...")
                        time.sleep(wait_sec)
                    else:
                        time.sleep(10)

            time.sleep(12)

        # -----------------------------------------------------------------
        # 최종 번역/교정 대조 리포트 발행
        # -----------------------------------------------------------------
        if all_translations:
            print("\n[REPORT] 영문 번역/교정 대조 리포트 생성 중...")
            report_path = os.path.join(target_dir, "EN_Translation_Polish_Report.txt")
            with open(report_path, "w", encoding="utf-8") as f:
                f.write("==============================================================\n")
                f.write("🇺🇸 EN V1 엔진: 국문 원본 vs 최종 영문(Montserrat) 매핑 리포트\n")
                f.write("==============================================================\n\n")
                for t in all_translations:
                    orig = t.get("original_text", "").replace("\n", " ")
                    corr = t.get("corrected_en", "").replace("\n", " ")
                    src = t.get("source_file", "")
                    mode = t.get("mode", "")
                    f.write("--------------------------------------------------------------\n")
                    f.write(f"[파일명]: {src} | [모드]: {mode}\n")
                    f.write(f"[원본 텍스트]: {orig}\n")
                    f.write(f"[교정 영문]: {corr}\n")
                f.write("--------------------------------------------------------------\n")
            print(f"  -> [REPORT SUCCESS] 리포트 저장 완료: {report_path}")

        print("\n[FINISH] EN_Text-In_Image_Translation_Engine_V1 영문 이미지 처리 파이프라인 종료!")


def _generate_docx_file(title: str, text_content: str, out_docx_path: str, target_lang: str = "EN"):
    """MS Word 서식(.docx)으로 4-Core 상세페이지 완성 원고를 렌더링합니다."""
    try:
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        doc = docx.Document()
        for s in doc.sections:
            s.top_margin = Inches(0.8)
            s.bottom_margin = Inches(0.8)
            s.left_margin = Inches(0.8)
            s.right_margin = Inches(0.8)

        # Title Header
        p_title = doc.add_paragraph()
        r_t = p_title.add_run(f"🛒 {title} - E-Commerce PDP Master Copy")
        r_t.font.name = "맑은 고딕"
        r_t.font.size = Pt(16)
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(16, 44, 87)

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(14)
        r_sub = p_sub.add_run("이 문서는 각국 전자상거래 플랫폼(스마트스토어, 쿠팡, 아마존, 큐텐 등) 상세페이지 에디터에 그대로 복사하여 등록할 수 있는 100% 고객 노출용 원고입니다.")
        r_sub.font.name = "맑은 고딕"
        r_sub.font.size = Pt(9.5)
        r_sub.font.color.rgb = RGBColor(100, 116, 139)

        # Parse sections
        lines = [l.strip() for l in text_content.splitlines() if l.strip()]
        cur_sec = 0
        sec_buffers = {1: [], 2: [], 3: [], 4: []}

        for l in lines:
            if l.startswith("1."):
                cur_sec = 1
                sec_buffers[1].append(l)
            elif l.startswith("2."):
                cur_sec = 2
                sec_buffers[2].append(l)
            elif l.startswith("3."):
                cur_sec = 3
                sec_buffers[3].append(l)
            elif l.startswith("4."):
                cur_sec = 4
                sec_buffers[4].append(l)
            else:
                if cur_sec in sec_buffers:
                    sec_buffers[cur_sec].append(l)

        # Section 1: Title
        h1 = doc.add_heading(level=1)
        r_h1 = h1.add_run("1. 공식 상품명 (Official Product Title)")
        r_h1.font.name = "맑은 고딕"
        r_h1.font.size = Pt(12.5)
        r_h1.font.bold = True
        r_h1.font.color.rgb = RGBColor(37, 99, 235)

        t_val = " ".join([l for l in sec_buffers[1] if not l.startswith("1.")]).strip()
        if not t_val and sec_buffers[1]:
            t_val = sec_buffers[1][-1].replace("1.", "").strip()
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_after = Pt(10)
        r_tv = p_t.add_run(t_val)
        r_tv.font.name = "맑은 고딕"
        r_tv.font.size = Pt(10.5)
        r_tv.font.bold = True

        # Section 2: Summary
        h2 = doc.add_heading(level=1)
        r_h2 = h2.add_run("2. 핵심 가치 및 제품 안내 (Core Value & Summary)")
        r_h2.font.name = "맑은 고딕"
        r_h2.font.size = Pt(12.5)
        r_h2.font.bold = True
        r_h2.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[2]:
            if l.startswith("2."):
                continue
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(l.lstrip("- •*"))
            r.font.name = "맑은 고딕"
            r.font.size = Pt(9.5)

        # Section 3: Comparison Table
        h3 = doc.add_heading(level=1)
        r_h3 = h3.add_run("3. 제품 상세 스펙 비교 (Product Specifications & Comparison Table)")
        r_h3.font.name = "맑은 고딕"
        r_h3.font.size = Pt(12.5)
        r_h3.font.bold = True
        r_h3.font.color.rgb = RGBColor(37, 99, 235)

        raw_s3_text = "\n".join(sec_buffers[3])
        if "<table>" in raw_s3_text:
            rows_data = []
            tr_matches = re.findall(r'<tr>(.*?)</tr>', raw_s3_text, flags=re.DOTALL)
            for tr in tr_matches:
                cols = re.findall(r'<t[hd]>(.*?)</t[hd]>', tr, flags=re.DOTALL)
                if cols:
                    rows_data.append([re.sub(r'<[^>]+>', '', c).strip() for c in cols])
            
            if rows_data:
                col_cnt = max(len(r) for r in rows_data)
                tbl = doc.add_table(rows=len(rows_data), cols=col_cnt)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r_idx, r_cols in enumerate(rows_data):
                    for c_idx, c_val in enumerate(r_cols):
                        cell = tbl.rows[r_idx].cells[c_idx]
                        p = cell.paragraphs[0]
                        r = p.add_run(c_val)
                        r.font.name = "맑은 고딕"
                        r.font.size = Pt(9)
                        if r_idx == 0:
                            r.font.bold = True
                            tcPr = cell._element.get_or_add_tcPr()
                            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
                            tcPr.append(shd)
        else:
            for l in sec_buffers[3]:
                if not l.startswith("3."):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    r = p.add_run(l)
                    r.font.name = "맑은 고딕"
                    r.font.size = Pt(9.5)

        # Section 4: FAQ
        h4 = doc.add_heading(level=1)
        r_h4 = h4.add_run("4. 자주 묻는 질문 (Frequently Asked Questions)")
        r_h4.font.name = "맑은 고딕"
        r_h4.font.size = Pt(12.5)
        r_h4.font.bold = True
        r_h4.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[4]:
            if l.startswith("4."):
                continue
            if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = RGBColor(30, 58, 138)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(51, 65, 85)

        doc.save(out_docx_path)
    except Exception as de:
        print(f"  ⚠️ [DOCX 생성 폴백]: {de}", flush=True)


def _generate_web_copier_html_file(title: str, text_content: str, out_html_path: str, target_lang: str = "EN"):
    """지마켓/스마트스토어/쿠팡/해외몰 등 전 이커머스 플랫폼에 맞춘 4-Core '에디터에 붙여넣을 텍스트 복사' 및 'HTML로 텍스트 복사' 뷰어를 생성합니다."""
    
    section_titles = {
        "EN": {
            "sec2": "2. Core Value & Active Ingredient Summary",
            "sec3": "3. Product Specifications & Comparison Table",
            "sec4": "4. Product Usage Guide & Frequently Asked Questions (FAQ)"
        },
        "JP": {
            "sec2": "2. コアバリュー＆成分サイエンス要約",
            "sec3": "3. 製品仕様・他社比較テーブル",
            "sec4": "4. 使用ガイド＆よくある質問 (FAQ)"
        },
        "CN": {
            "sec2": "2. 核心价值与成分科技摘要",
            "sec3": "3. 产品规格与竞品对比表",
            "sec4": "4. 商品使用指南与常见问题解答 (FAQ)"
        },
        "TW": {
            "sec2": "2. 核心價值與成分科技摘要",
            "sec3": "3. 產品規格與競品對比表",
            "sec4": "4. 使用指南與常見問題解答 (FAQ)"
        },
        "KR": {
            "sec2": "2. 제품 핵심 안내 및 성분 요약",
            "sec3": "3. 제품 상세 비교 스펙 테이블",
            "sec4": "4. 자주 묻는 질문 (FAQ)"
        }
    }
    
    detected_lang = target_lang.upper() if target_lang else "EN"
    titles_map = section_titles.get(detected_lang, section_titles["EN"])
    sec2_heading = titles_map["sec2"]
    sec3_heading = titles_map["sec3"]
    sec4_heading = titles_map.get("sec4", "4. FAQ")

    lines = [l.strip() for l in text_content.splitlines() if l.strip()]
    
    s1_lines = []
    s2_lines = []
    s3_lines = []
    s4_lines = []
    cur_sec = 0
    for l in lines:
        if l.startswith("1."):
            cur_sec = 1
            s1_lines.append(l)
        elif l.startswith("2."):
            cur_sec = 2
            s2_lines.append(l)
        elif l.startswith("3."):
            cur_sec = 3
            s3_lines.append(l)
        elif l.startswith("4."):
            cur_sec = 4
            s4_lines.append(l)
        else:
            if cur_sec == 1:
                s1_lines.append(l)
            elif cur_sec == 2:
                s2_lines.append(l)
            elif cur_sec == 3:
                s3_lines.append(l)
            elif cur_sec == 4:
                s4_lines.append(l)

    # 1. Section 1 (Title)
    s1_clean = " ".join([l for l in s1_lines if not l.startswith("1.")]).strip()
    if not s1_clean and s1_lines:
        s1_clean = s1_lines[-1].replace("1.", "").strip()

    # 2. Section 2 (Summary)
    s2_clean_items = [l for l in s2_lines if not l.startswith("2.")]
    s2_text_formatted = f"{sec2_heading}\n\n" + "\n\n".join(s2_clean_items)
    
    s2_html_items = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s2_html_items.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>🔬 {sec2_heading}</h3>")
    for l in s2_clean_items:
        if ":" in l or "：" in l:
            k, v = re.split(r'[:：]', l, 1)
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'><strong style='color:#1e3a8a;'>{k.strip()}:</strong> {v.strip()}</p>")
        else:
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'>{l.strip()}</p>")
    s2_html_items.append("</div>")
    s2_html = "\n".join(s2_html_items)

    # 3. Section 3 (Table)
    s3_clean_items = [l for l in s3_lines if not l.startswith("3.")]
    s3_raw_block = "\n".join(s3_clean_items)
    s3_text_formatted = f"{sec3_heading}\n\n" + s3_raw_block
    
    if "<table>" in s3_raw_block:
        styled_table = s3_raw_block.replace("<table>", "<table style='width:100%; border-collapse:collapse; margin:10px 0; font-size:13.5px;'>")
        styled_table = styled_table.replace("<th>", "<th style='background:#f1f5f9; padding:10px 12px; border:1px solid #cbd5e1; color:#1e3a8a; text-align:left; font-weight:bold;'>")
        styled_table = styled_table.replace("<td>", "<td style='padding:9px 12px; border:1px solid #cbd5e1; color:#334155;'>")
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n{styled_table}\n</div>"
    else:
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n<pre style='white-space:pre-wrap; font-family:inherit; font-size:13.5px;'>{s3_raw_block}</pre>\n</div>"

    # 4. Section 4 (FAQ)
    s4_clean_items = [l for l in s4_lines if not l.startswith("4.")]
    s4_text_blocks = []
    s4_html_blocks = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s4_html_blocks.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 12px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>💬 {sec4_heading}</h3>")
    
    cur_q = ""
    cur_answers = []
    for l in s4_clean_items:
        if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
            if cur_q:
                q_text = cur_q
                a_text = "\n".join(cur_answers)
                s4_text_blocks.append(f"{q_text}\n{a_text}")
                s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
                for ans in cur_answers:
                    s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
            cur_q = l.strip()
            cur_answers = []
        else:
            sentences = [s.strip() for s in re.split(r'(?<=[。！？\.\?!])\s*', l) if s.strip()]
            for s in sentences:
                cur_answers.append(s)
                
    if cur_q:
        q_text = cur_q
        a_text = "\n".join(cur_answers)
        s4_text_blocks.append(f"{q_text}\n{a_text}")
        s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
        for ans in cur_answers:
            s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
    s4_html_blocks.append("</div>")

    s4_text_formatted = f"{sec4_heading}\n\n" + "\n\n".join(s4_text_blocks)
    s4_html = "\n".join(s4_html_blocks)

    full_html_code = f"""<!-- 다국어 E-Commerce 상세페이지 4-Core 마이크로-써머리 & 비교표 & FAQ -->
<div style="font-family:'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; max-width:860px; margin:0 auto; padding:10px 0; color:#1e293b;">
{s2_html}
{s3_html}
{s4_html}
</div>"""

    html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; background: #f1f5f9; color: #0f172a; padding: 25px; margin: 0; line-height: 1.6; }}
  .container {{ max-width: 960px; margin: 0 auto; background: #ffffff; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.06); padding: 35px; border: 1px solid #cbd5e1; }}
  h1 {{ font-size: 22px; color: #0f172a; border-bottom: 2px solid #2563eb; padding-bottom: 12px; margin-top: 0; display: flex; align-items: center; gap: 8px; }}
  .guide-box {{ background: #eff6ff; border: 1px solid #bfdbfe; border-left: 5px solid #2563eb; padding: 18px 20px; border-radius: 8px; margin-bottom: 25px; font-size: 14px; color: #1e40af; line-height: 1.8; }}
  .guide-box strong {{ color: #1e3a8a; }}
  .card {{ background: #f8fafc; border: 1px solid #cbd5e1; border-radius: 10px; padding: 22px; margin-bottom: 25px; }}
  .card-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 14px; border-bottom: 1px solid #e2e8f0; padding-bottom: 10px; }}
  .card-title {{ font-size: 16px; font-weight: bold; color: #1e3a8a; }}
  .btn-group {{ display: flex; gap: 8px; }}
  .copy-btn {{ background: #2563eb; color: #ffffff; border: none; padding: 9px 16px; border-radius: 6px; cursor: pointer; font-size: 13px; font-weight: 700; transition: all 0.2s; box-shadow: 0 2px 4px rgba(37,99,235,0.2); display: inline-flex; align-items: center; gap: 6px; }}
  .copy-btn:hover {{ background: #1d4ed8; transform: translateY(-1px); }}
  .copy-btn.html-mode-btn {{ background: #059669; box-shadow: 0 2px 4px rgba(5,150,105,0.25); }}
  .copy-btn.html-mode-btn:hover {{ background: #047857; }}
  .full-btn {{ background: #7c3aed; padding: 12px 24px; font-size: 15px; width: 100%; justify-content: center; margin-bottom: 20px; box-shadow: 0 3px 6px rgba(124,58,237,0.25); }}
  .full-btn:hover {{ background: #6d28d9; }}
  .text-area {{ width: 100%; box-sizing: border-box; background: #ffffff; border: 1px solid #cbd5e1; border-radius: 6px; padding: 14px 16px; font-size: 13.5px; line-height: 1.75; color: #1e293b; font-family: 'Malgun Gothic', 'Segoe UI', monospace; resize: vertical; outline: none; }}
  .text-area:focus {{ border-color: #2563eb; box-shadow: 0 0 0 3px rgba(37,99,235,0.15); }}
  .toast {{ position: fixed; bottom: 30px; right: 30px; background: #0f172a; color: #ffffff; padding: 14px 28px; border-radius: 8px; font-size: 14px; font-weight: 600; display: none; z-index: 1000; box-shadow: 0 6px 20px rgba(0,0,0,0.25); }}
</style>
</head>
<body>

<div class="container">
  <h1>🌐 {title}</h1>
  
  <div class="guide-box">
    📢 <strong>쇼핑몰 등록 방식별 2대 원클릭 복사 기능 안내:</strong><br>
    • <strong>1. [📋 에디터에 붙여넣을 텍스트 복사] (파란색 버튼)</strong>: 지마켓/스마트스토어/쿠팡/해외몰 <strong>'에디터 작성'</strong> 화면에 붙여넣을 때 사용합니다.<br>
    • <strong>2. [🌐 HTML로 텍스트 복사] (초록색 버튼)</strong>: <strong>'HTML 작성'</strong> 탭이나 HTML 직접 입력 모드에 붙여넣을 때 사용합니다.
  </div>

  <button class="copy-btn full-btn html-mode-btn" onclick="copyFromTextarea('full-html-ta', '🎉 전체 HTML 소스코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🚀 [HTML로 전체 일괄 복사] 4-Core 요약 + 비교표 + FAQ 전체 소스코드 복사</button>
  <textarea id="full-html-ta" style="display:none;">{full_html_code}</textarea>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📌 1. 공식 상품명 (Title)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec1-ta', '✅ 상품명이 복사되었습니다!')">📋 상품명 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec1-ta" rows="2" readonly>{s1_clean}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">🔬 2. 핵심 가치 및 5줄 마이크로 요약</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec2-ta', '✅ 5줄 요약 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec2-html-ta', '🌐 5줄 요약 HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec2-ta" rows="8" readonly>{s2_text_formatted}</textarea>
    <textarea id="sec2-html-ta" style="display:none;">{s2_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📊 3. 제품 상세 스펙 비교표 (HTML Table)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec3-ta', '✅ 비교표 텍스트가 복사되었습니다!')">📋 비교표 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec3-html-ta', '🌐 비교표 HTML 코드가 복사되었습니다!')">🌐 HTML로 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec3-ta" rows="10" readonly>{s3_text_formatted}</textarea>
    <textarea id="sec3-html-ta" style="display:none;">{s3_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">💬 4. 5대 핵심 FAQ & 상세 가이드</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec4-ta', '✅ FAQ 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec4-html-ta', '🌐 FAQ HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec4-ta" rows="18" readonly>{s4_text_formatted}</textarea>
    <textarea id="sec4-html-ta" style="display:none;">{s4_html}</textarea>
  </div>

</div>

<div class="toast" id="toast">✅ 클립보드에 복사되었습니다! 쇼핑몰 에디터에 붙여넣기(Ctrl+V) 하세요.</div>

<script>
function showToast(msg) {{
  const t = document.getElementById('toast');
  t.innerText = msg;
  t.style.display = 'block';
  setTimeout(() => {{ t.style.display = 'none'; }}, 3000);
}}

function copyFromTextarea(id, customMsg) {{
  const ta = document.getElementById(id);
  const text = ta.value;
  
  const temp = document.createElement('textarea');
  temp.value = text;
  temp.style.position = 'fixed';
  temp.style.left = '-9999px';
  document.body.appendChild(temp);
  temp.select();
  temp.setSelectionRange(0, 99999);
  
  try {{
    document.execCommand('copy');
    showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
  }} catch (err) {{
    if (navigator.clipboard) {{
      navigator.clipboard.writeText(text).then(() => {{
        showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
      }});
    }}
  }} finally {{
    document.body.removeChild(temp);
  }}
}}
</script>

</body>
</html>
"""
    with open(out_html_path, "w", encoding="utf-8") as f:
        f.write(html_content)


async def generate_seo_geo_aeo_txt(client: genai.Client, current_source_dir: str, target_dir: str, target_lang: str, product_name: str):
    """4-Core 마이크로-써머리 SEO/GEO/AEO (TXT, HTML 뷰어, DOCX, MD) 4종 파일을 자동 생성합니다. (URL/고시표 듀얼 인제스천 및 4개국 법무 렉시콘 100% 결합)"""
    print(f"\n🌐 [SEO/GEO/AEO 4-Core] 정밀 팩트 인제스천 및 4종 포맷(DOCX/HTML/TXT/MD) 생성 중 ({target_lang})...", flush=True)

    # 1. 듀얼 인제스천: url.txt 실시간 웹 스크래핑
    url_fact_context = ""
    url_file_candidates = [
        os.path.join(current_source_dir, "url.txt"),
        os.path.join(current_source_dir, "product_url.txt"),
        os.path.join(current_source_dir, "URL.txt")
    ]
    for u_path in url_file_candidates:
        if os.path.exists(u_path):
            try:
                with open(u_path, "r", encoding="utf-8") as uf:
                    raw_url = uf.read().strip()
                if raw_url.startswith("http"):
                    print(f"  🔗 [URL INGESTION 감지] {raw_url}", flush=True)
                    import urllib.request
                    req = urllib.request.Request(raw_url, headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"})
                    with urllib.request.urlopen(req, timeout=5) as resp:
                        html_bytes = resp.read()
                        raw_html = html_bytes.decode("utf-8", errors="ignore")
                        clean_text = re.sub(r'<script.*?</script>', '', raw_html, flags=re.DOTALL)
                        clean_text = re.sub(r'<style.*?</style>', '', clean_text, flags=re.DOTALL)
                        clean_text = re.sub(r'<[^>]+>', ' ', clean_text)
                        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                        url_fact_context = f"\n[Live Web Product Page Facts (Ground Truth)]\nURL: {raw_url}\nScraped Text Extract: {clean_text[:1200]}\n"
                        print(f"  ✅ [URL 팩트 스크래핑 완료] {len(clean_text)} 자 추출", flush=True)
                        break
            except Exception as ue:
                print(f"  ⚠️ [URL 스크래핑 실패 / 폴백 진행]: {ue}", flush=True)

    # 2. 이미지 팩트 앵커링 & 초월번역 컨텍스트 로드
    guide_path = os.path.join(current_source_dir, "transcreation_guide.json")
    qa_path = os.path.join(target_dir, "Transcreation_QA_Report.json")
    guide_context = ""
    if os.path.exists(guide_path):
        try:
            with open(guide_path, "r", encoding="utf-8") as f:
                gdata = json.load(f)
                guide_context = json.dumps(gdata.get("transcreation_comparisons", [])[:6], ensure_ascii=False)
        except Exception:
            pass
    elif os.path.exists(qa_path):
        try:
            with open(qa_path, "r", encoding="utf-8") as f:
                qdata = json.load(f)
                guide_context = json.dumps(qdata.get("transcreation_comparisons", [])[:6], ensure_ascii=False)
        except Exception:
            pass

    # 3. 4개국 법무 렉시콘 로드 (COMPLIANCE-FIRST)
    lexicon_rules_text = ""
    lexicon_map = {
        "EN": "en_fda_mocra_lexicon.json",
        "JP": "jp_pmda_pharm_lexicon.json",
        "CN": "cn_nmpa_adlaw_lexicon.json",
        "TW": "tw_tfda_lexicon.json"
    }
    lex_file = lexicon_map.get(target_lang)
    if lex_file:
        lex_path = os.path.join(PROJECT_ROOT, "00_공통자료", "compliance_lexicons", lex_file)
        if os.path.exists(lex_path):
            try:
                with open(lex_path, "r", encoding="utf-8") as lf:
                    lex_data = json.load(lf)
                    lexicon_rules_text = f"\n[MANDATORY COMPLIANCE LEXICON ({lex_data.get('jurisdiction', '')})]\n"
                    cats = lex_data.get("categories", {})
                    for c_name, c_val in cats.items():
                        banned = c_val.get("banned_terms", [])
                        for b in banned[:5]:
                            lexicon_rules_text += f"- Banned: '{b.get('banned')}' -> Must use: '{b.get('preferred')}' ({b.get('reason')})\n"
            except Exception:
                pass

    lang_names = {
        "EN": "English for Amazon / Sephora US",
        "JP": "Japanese for Qoo10 Japan / Cosme",
        "CN": "Simplified Chinese for Tmall / Xiaohongshu",
        "TW": "Traditional Chinese for Shopee Taiwan / Momo",
        "KR": "Korean for Naver Smartstore / Coupang"
    }
    target_lang_desc = lang_names.get(target_lang, "English")

    prompt = f"""[SYSTEM PROMPT] Global E-Commerce SEO/GEO/AEO 4-Core Master Copy Generator
Product Name: {product_name}
Target Market & Language: {target_lang_desc}

{url_fact_context}
[Verified Transcreation Context & Ingredients Data]:
{guide_context}

{lexicon_rules_text}

[CRITICAL INSTRUCTION - ZERO META COMMENTARY]
Never output developer metadata, markdown headers '##', explanation notes, character counters, or words like 'GEO', 'AEO', 'RAG'.
Output purely customer-facing content structured in 4 distinct sections.

Strict 4-Core Structure:
1. Official Product Title
(Under 100 characters. Noun Phrase: Brand 'Logicall Skin' + Product Title + Key Efficacy + Volume)

2. Core Value & Active Ingredient Summary
(5 concise bullet points containing quantitative metrics e.g. ppm, %, non-irritation score 0.00):
- Brand: Logicall Skin
- Core Actives & Concentration: (e.g. Multi-Vitamin 10% / 100,000ppm, Aquatide 3%)
- Key Benefits: (Efficacy claims strictly compliant with target country cosmetics law)
- Texture & Absorption: (Hydra-watery, non-greasy, fast-absorbing)
- Skin Compatibility: (Dermatologist tested, 0.00 irritation index)

3. Product Specifications & Comparison Table
(Output a clean HTML <table> comparing Logicall Skin vs Generic Market Standard):
<table>
  <tr><th>Dimensions</th><th>Logicall Skin</th><th>Standard Market Benchmark</th></tr>
  <tr><td>Active Concentration</td><td>High-Potency Multi-Vitamin 100,000ppm (10%)</td><td>Diluted extract 1,000~5,000ppm</td></tr>
  <tr><td>Patented Science</td><td>Aquatide 5000 (3%)</td><td>Generic purified water base</td></tr>
  <tr><td>Formula Stability</td><td>High-stability oxidation-free formula</td><td>Prone to discoloration / oxidation</td></tr>
  <tr><td>Irritation Index</td><td>0.00 Low-Irritation Certified</td><td>May cause stinging or redness</td></tr>
</table>

4. Product Usage Guide & Frequently Asked Questions (FAQ)
(5 high-conversion B2C customer FAQs):
Q1: When should I apply this serum?
A: ...
Q2: Is it suitable for sensitive skin?
A: ...
Q3: How does the active formula benefit the skin?
A: ...
Q4: What is the main efficacy of the Multi-Vitamin complex?
A: ...
Q5: Can I layer this with other skincare products?
A: ...

Generate the complete 4-Core content in {target_lang_desc} now.
"""
    try:
        resp = await client.aio.models.generate_content(
            model=MODEL_PRO,
            contents=[prompt],
            config=types.GenerateContentConfig(
                temperature=0.6,
                top_p=0.9,
                max_output_tokens=8192
            )
        )
        content_text = resp.text.strip()
    except Exception as e:
        print(f"  ⚠️ [WARN] SEO 텍스트 생성 중 오류 발생 -> 기본 템플릿 대체: {e}")
        content_text = f"""1. Logicall Skin {product_name} Multi Vitamin Daily Care Serum 50ml

2. Core Value & Active Ingredient Summary
- Brand: Logicall Skin
- Core Actives & Concentration: High-Potency Multi-Vitamin 10% (100,000ppm) & Aquatide 5000 (3%)
- Key Benefits: Visibly brightens, refines texture, and reinforces natural moisture barrier
- Texture & Absorption: Refreshingly lightweight hydra-watery formula with instant absorption
- Skin Compatibility: Dermatologist-tested, 0.00 skin irritation index suitable for daily use

3. Product Specifications & Comparison Table
<table>
  <tr><th>Specification</th><th>Logicall Skin Multi-Vitamin Serum</th><th>Standard Vitamin Serum</th></tr>
  <tr><td>Active Concentration</td><td>Multi-Vitamin Complex 100,000ppm (10.0%)</td><td>1,000 ~ 5,000ppm</td></tr>
  <tr><td>Patented Technology</td><td>Aquatide 5000 30,000ppm (3.0%)</td><td>Purified water base</td></tr>
  <tr><td>Formula Stability</td><td>High stability against air and light oxidation</td><td>Vulnerable to discoloration</td></tr>
  <tr><td>Skin Irritation Score</td><td>0.00 (Certified Low-Irritation)</td><td>May cause stinging sensation</td></tr>
</table>

4. Product Usage Guide & Frequently Asked Questions (FAQ)
Q1: When should I apply this serum?
A: Apply 3-4 drops evenly morning and evening after cleansing and toner.
Q2: Is it suitable for sensitive skin?
A: Yes, it is dermatologist-tested with a 0.00 skin irritation index.
Q3: How does Aquatide benefit the skin?
A: It reinforces the natural moisture barrier and revitalizes skin appearance.
Q4: What is the main efficacy of the Multi-Vitamin complex?
A: It provides deep hydration for a resilient-looking complexion and combats the signs of premature aging.
Q5: Can I layer this with other skincare products?
A: Yes, its fast-absorbing texture layers smoothly under creams and sunscreens.
"""

    # 4. 결정론적 법무 후처리 게이트 통과 (apply_deterministic_qa_overrides)
    if target_lang == "EN":
        for b_pat, p_val in [
            (r"\bnutrients for cellular vitality\b", "hydration for a resilient-looking complexion"),
            (r"\bcellular vitality\b", "resilient-looking complexion"),
            (r"\breinforces cellular resilience\b", "reinforces the skin's natural moisture barrier"),
            (r"\bcellular resilience\b", "skin's natural moisture barrier"),
            (r"\bcellular metabolism\b", "natural skin vitality"),
            (r"\bcombats premature aging\b", "combats the signs of premature aging"),
            (r"\bComplex skin issues\b", "Multiple skin concerns"),
            (r"\bTroubled skin\b", "Blemish-prone skin"),
            (r"\bcellular autophagy\b", "targeted skin nourishment")
        ]:
            content_text = re.sub(b_pat, p_val, content_text, flags=re.IGNORECASE)

    # 5. 4종 멀티 포맷 일괄 익스포트 (TXT, HTML, DOCX, MD)
    txt_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO.txt"
    html_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO_VIEWER.html"
    docx_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO.docx"
    md_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO.md"

    txt_path = os.path.join(target_dir, txt_filename)
    html_path = os.path.join(target_dir, html_filename)
    docx_path = os.path.join(target_dir, docx_filename)
    md_path = os.path.join(target_dir, md_filename)

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(content_text)
    print(f"  📄 [TXT 저장 완료]: {txt_path}")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(f"# {product_name} ({target_lang}) SEO/GEO/AEO PDP 원고\n\n" + content_text)
    print(f"  📑 [MD 저장 완료]: {md_path}")

    _generate_web_copier_html_file(f"{product_name} ({target_lang})", content_text, html_path, target_lang=target_lang)
    print(f"  🌐 [HTML 뷰어 저장 완료]: {html_path}")

    _generate_docx_file(f"{product_name} ({target_lang})", content_text, docx_path, target_lang=target_lang)
    print(f"  📄 [DOCX 서식 문서 저장 완료]: {docx_path}")


def _generate_docx_file(title: str, text_content: str, out_docx_path: str, target_lang: str = "EN"):
    """MS Word 서식(.docx)으로 4-Core 상세페이지 완성 원고를 렌더링합니다."""
    try:
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        doc = docx.Document()
        for s in doc.sections:
            s.top_margin = Inches(0.8)
            s.bottom_margin = Inches(0.8)
            s.left_margin = Inches(0.8)
            s.right_margin = Inches(0.8)

        # Title Header
        p_title = doc.add_paragraph()
        r_t = p_title.add_run(f"🛒 {title} - E-Commerce PDP Master Copy")
        r_t.font.name = "맑은 고딕"
        r_t.font.size = Pt(16)
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(16, 44, 87)

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(14)
        r_sub = p_sub.add_run("이 문서는 각국 전자상거래 플랫폼(스마트스토어, 쿠팡, 아마존, 큐텐 등) 상세페이지 에디터에 그대로 복사하여 등록할 수 있는 100% 고객 노출용 원고입니다.")
        r_sub.font.name = "맑은 고딕"
        r_sub.font.size = Pt(9.5)
        r_sub.font.color.rgb = RGBColor(100, 116, 139)

        # Parse sections
        lines = [l.strip() for l in text_content.splitlines() if l.strip()]
        cur_sec = 0
        sec_buffers = {1: [], 2: [], 3: [], 4: []}

        for l in lines:
            if l.startswith("1."):
                cur_sec = 1
                sec_buffers[1].append(l)
            elif l.startswith("2."):
                cur_sec = 2
                sec_buffers[2].append(l)
            elif l.startswith("3."):
                cur_sec = 3
                sec_buffers[3].append(l)
            elif l.startswith("4."):
                cur_sec = 4
                sec_buffers[4].append(l)
            else:
                if cur_sec in sec_buffers:
                    sec_buffers[cur_sec].append(l)

        # Section 1: Title
        h1 = doc.add_heading(level=1)
        r_h1 = h1.add_run("1. 공식 상품명 (Official Product Title)")
        r_h1.font.name = "맑은 고딕"
        r_h1.font.size = Pt(12.5)
        r_h1.font.bold = True
        r_h1.font.color.rgb = RGBColor(37, 99, 235)

        t_val = " ".join([l for l in sec_buffers[1] if not l.startswith("1.")]).strip()
        if not t_val and sec_buffers[1]:
            t_val = sec_buffers[1][-1].replace("1.", "").strip()
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_after = Pt(10)
        r_tv = p_t.add_run(t_val)
        r_tv.font.name = "맑은 고딕"
        r_tv.font.size = Pt(10.5)
        r_tv.font.bold = True

        # Section 2: Summary
        h2 = doc.add_heading(level=1)
        r_h2 = h2.add_run("2. 핵심 가치 및 제품 안내 (Core Value & Summary)")
        r_h2.font.name = "맑은 고딕"
        r_h2.font.size = Pt(12.5)
        r_h2.font.bold = True
        r_h2.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[2]:
            if l.startswith("2."):
                continue
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(l.lstrip("- •*"))
            r.font.name = "맑은 고딕"
            r.font.size = Pt(9.5)

        # Section 3: Comparison Table
        h3 = doc.add_heading(level=1)
        r_h3 = h3.add_run("3. 제품 상세 스펙 비교 (Product Specifications & Comparison Table)")
        r_h3.font.name = "맑은 고딕"
        r_h3.font.size = Pt(12.5)
        r_h3.font.bold = True
        r_h3.font.color.rgb = RGBColor(37, 99, 235)

        raw_s3_text = "\n".join(sec_buffers[3])
        if "<table>" in raw_s3_text:
            rows_data = []
            tr_matches = re.findall(r'<tr>(.*?)</tr>', raw_s3_text, flags=re.DOTALL)
            for tr in tr_matches:
                cols = re.findall(r'<t[hd]>(.*?)</t[hd]>', tr, flags=re.DOTALL)
                if cols:
                    rows_data.append([re.sub(r'<[^>]+>', '', c).strip() for c in cols])
            
            if rows_data:
                col_cnt = max(len(r) for r in rows_data)
                tbl = doc.add_table(rows=len(rows_data), cols=col_cnt)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r_idx, r_cols in enumerate(rows_data):
                    for c_idx, c_val in enumerate(r_cols):
                        cell = tbl.rows[r_idx].cells[c_idx]
                        p = cell.paragraphs[0]
                        r = p.add_run(c_val)
                        r.font.name = "맑은 고딕"
                        r.font.size = Pt(9)
                        if r_idx == 0:
                            r.font.bold = True
                            tcPr = cell._element.get_or_add_tcPr()
                            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
                            tcPr.append(shd)
        else:
            for l in sec_buffers[3]:
                if not l.startswith("3."):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    r = p.add_run(l)
                    r.font.name = "맑은 고딕"
                    r.font.size = Pt(9.5)

        # Section 4: FAQ
        h4 = doc.add_heading(level=1)
        r_h4 = h4.add_run("4. 자주 묻는 질문 (Frequently Asked Questions)")
        r_h4.font.name = "맑은 고딕"
        r_h4.font.size = Pt(12.5)
        r_h4.font.bold = True
        r_h4.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[4]:
            if l.startswith("4."):
                continue
            if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = RGBColor(30, 58, 138)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(51, 65, 85)

        doc.save(out_docx_path)
    except Exception as de:
        print(f"  ⚠️ [DOCX 생성 폴백]: {de}", flush=True)


def _generate_web_copier_html_file(title: str, text_content: str, out_html_path: str, target_lang: str = "EN"):
    """지마켓/스마트스토어/쿠팡/해외몰 등 전 이커머스 플랫폼에 맞춘 4-Core '에디터에 붙여넣을 텍스트 복사' 및 'HTML로 텍스트 복사' 뷰어를 생성합니다."""
    
    section_titles = {
        "EN": {
            "sec2": "2. Core Value & Active Ingredient Summary",
            "sec3": "3. Product Specifications & Comparison Table",
            "sec4": "4. Product Usage Guide & Frequently Asked Questions (FAQ)"
        },
        "JP": {
            "sec2": "2. コアバリュー＆成分サイエンス要約",
            "sec3": "3. 製品仕様・他社比較テーブル",
            "sec4": "4. 使用ガイド＆よくある質問 (FAQ)"
        },
        "CN": {
            "sec2": "2. 核心价值与成分科技摘要",
            "sec3": "3. 产品规格与竞品对比表",
            "sec4": "4. 商品使用指南与常见问题解答 (FAQ)"
        },
        "TW": {
            "sec2": "2. 核心價值與成分科技摘要",
            "sec3": "3. 產品規格與競品對比表",
            "sec4": "4. 使用指南與常見問題解答 (FAQ)"
        },
        "KR": {
            "sec2": "2. 제품 핵심 안내 및 성분 요약",
            "sec3": "3. 제품 상세 비교 스펙 테이블",
            "sec4": "4. 자주 묻는 질문 (FAQ)"
        }
    }
    
    detected_lang = target_lang.upper() if target_lang else "EN"
    titles_map = section_titles.get(detected_lang, section_titles["EN"])
    sec2_heading = titles_map["sec2"]
    sec3_heading = titles_map["sec3"]
    sec4_heading = titles_map.get("sec4", "4. FAQ")

    lines = [l.strip() for l in text_content.splitlines() if l.strip()]
    
    s1_lines = []
    s2_lines = []
    s3_lines = []
    s4_lines = []
    cur_sec = 0
    for l in lines:
        if l.startswith("1."):
            cur_sec = 1
            s1_lines.append(l)
        elif l.startswith("2."):
            cur_sec = 2
            s2_lines.append(l)
        elif l.startswith("3."):
            cur_sec = 3
            s3_lines.append(l)
        elif l.startswith("4."):
            cur_sec = 4
            s4_lines.append(l)
        else:
            if cur_sec == 1:
                s1_lines.append(l)
            elif cur_sec == 2:
                s2_lines.append(l)
            elif cur_sec == 3:
                s3_lines.append(l)
            elif cur_sec == 4:
                s4_lines.append(l)

    # 1. Section 1 (Title)
    s1_clean = " ".join([l for l in s1_lines if not l.startswith("1.")]).strip()
    if not s1_clean and s1_lines:
        s1_clean = s1_lines[-1].replace("1.", "").strip()

    # 2. Section 2 (Summary)
    s2_clean_items = [l for l in s2_lines if not l.startswith("2.")]
    s2_text_formatted = f"{sec2_heading}\n\n" + "\n\n".join(s2_clean_items)
    
    s2_html_items = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s2_html_items.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>🔬 {sec2_heading}</h3>")
    for l in s2_clean_items:
        if ":" in l or "：" in l:
            k, v = re.split(r'[:：]', l, 1)
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'><strong style='color:#1e3a8a;'>{k.strip()}:</strong> {v.strip()}</p>")
        else:
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'>{l.strip()}</p>")
    s2_html_items.append("</div>")
    s2_html = "\n".join(s2_html_items)

    # 3. Section 3 (Table)
    s3_clean_items = [l for l in s3_lines if not l.startswith("3.")]
    s3_raw_block = "\n".join(s3_clean_items)
    s3_text_formatted = f"{sec3_heading}\n\n" + s3_raw_block
    
    if "<table>" in s3_raw_block:
        styled_table = s3_raw_block.replace("<table>", "<table style='width:100%; border-collapse:collapse; margin:10px 0; font-size:13.5px;'>")
        styled_table = styled_table.replace("<th>", "<th style='background:#f1f5f9; padding:10px 12px; border:1px solid #cbd5e1; color:#1e3a8a; text-align:left; font-weight:bold;'>")
        styled_table = styled_table.replace("<td>", "<td style='padding:9px 12px; border:1px solid #cbd5e1; color:#334155;'>")
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n{styled_table}\n</div>"
    else:
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n<pre style='white-space:pre-wrap; font-family:inherit; font-size:13.5px;'>{s3_raw_block}</pre>\n</div>"

    # 4. Section 4 (FAQ)
    s4_clean_items = [l for l in s4_lines if not l.startswith("4.")]
    s4_text_blocks = []
    s4_html_blocks = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s4_html_blocks.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 12px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>💬 {sec4_heading}</h3>")
    
    cur_q = ""
    cur_answers = []
    for l in s4_clean_items:
        if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
            if cur_q:
                q_text = cur_q
                a_text = "\n".join(cur_answers)
                s4_text_blocks.append(f"{q_text}\n{a_text}")
                s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
                for ans in cur_answers:
                    s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
            cur_q = l.strip()
            cur_answers = []
        else:
            sentences = [s.strip() for s in re.split(r'(?<=[。！？\.\?!])\s*', l) if s.strip()]
            for s in sentences:
                cur_answers.append(s)
                
    if cur_q:
        q_text = cur_q
        a_text = "\n".join(cur_answers)
        s4_text_blocks.append(f"{q_text}\n{a_text}")
        s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
        for ans in cur_answers:
            s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
    s4_html_blocks.append("</div>")

    s4_text_formatted = f"{sec4_heading}\n\n" + "\n\n".join(s4_text_blocks)
    s4_html = "\n".join(s4_html_blocks)

    full_html_code = f"""<!-- 다국어 E-Commerce 상세페이지 4-Core 마이크로-써머리 & 비교표 & FAQ -->
<div style="font-family:'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; max-width:860px; margin:0 auto; padding:10px 0; color:#1e293b;">
{s2_html}
{s3_html}
{s4_html}
</div>"""

    html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; background: #f1f5f9; color: #0f172a; padding: 25px; margin: 0; line-height: 1.6; }}
  .container {{ max-width: 960px; margin: 0 auto; background: #ffffff; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.06); padding: 35px; border: 1px solid #cbd5e1; }}
  h1 {{ font-size: 22px; color: #0f172a; border-bottom: 2px solid #2563eb; padding-bottom: 12px; margin-top: 0; display: flex; align-items: center; gap: 8px; }}
  .guide-box {{ background: #eff6ff; border: 1px solid #bfdbfe; border-left: 5px solid #2563eb; padding: 18px 20px; border-radius: 8px; margin-bottom: 25px; font-size: 14px; color: #1e40af; line-height: 1.8; }}
  .guide-box strong {{ color: #1e3a8a; }}
  .card {{ background: #f8fafc; border: 1px solid #cbd5e1; border-radius: 10px; padding: 22px; margin-bottom: 25px; }}
  .card-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 14px; border-bottom: 1px solid #e2e8f0; padding-bottom: 10px; }}
  .card-title {{ font-size: 16px; font-weight: bold; color: #1e3a8a; }}
  .btn-group {{ display: flex; gap: 8px; }}
  .copy-btn {{ background: #2563eb; color: #ffffff; border: none; padding: 9px 16px; border-radius: 6px; cursor: pointer; font-size: 13px; font-weight: 700; transition: all 0.2s; box-shadow: 0 2px 4px rgba(37,99,235,0.2); display: inline-flex; align-items: center; gap: 6px; }}
  .copy-btn:hover {{ background: #1d4ed8; transform: translateY(-1px); }}
  .copy-btn.html-mode-btn {{ background: #059669; box-shadow: 0 2px 4px rgba(5,150,105,0.25); }}
  .copy-btn.html-mode-btn:hover {{ background: #047857; }}
  .full-btn {{ background: #7c3aed; padding: 12px 24px; font-size: 15px; width: 100%; justify-content: center; margin-bottom: 20px; box-shadow: 0 3px 6px rgba(124,58,237,0.25); }}
  .full-btn:hover {{ background: #6d28d9; }}
  .text-area {{ width: 100%; box-sizing: border-box; background: #ffffff; border: 1px solid #cbd5e1; border-radius: 6px; padding: 14px 16px; font-size: 13.5px; line-height: 1.75; color: #1e293b; font-family: 'Malgun Gothic', 'Segoe UI', monospace; resize: vertical; outline: none; }}
  .text-area:focus {{ border-color: #2563eb; box-shadow: 0 0 0 3px rgba(37,99,235,0.15); }}
  .toast {{ position: fixed; bottom: 30px; right: 30px; background: #0f172a; color: #ffffff; padding: 14px 28px; border-radius: 8px; font-size: 14px; font-weight: 600; display: none; z-index: 1000; box-shadow: 0 6px 20px rgba(0,0,0,0.25); }}
</style>
</head>
<body>

<div class="container">
  <h1>🌐 {title}</h1>
  
  <div class="guide-box">
    📢 <strong>쇼핑몰 등록 방식별 2대 원클릭 복사 기능 안내:</strong><br>
    • <strong>1. [📋 에디터에 붙여넣을 텍스트 복사] (파란색 버튼)</strong>: 지마켓/스마트스토어/쿠팡/해외몰 <strong>'에디터 작성'</strong> 화면에 붙여넣을 때 사용합니다.<br>
    • <strong>2. [🌐 HTML로 텍스트 복사] (초록색 버튼)</strong>: <strong>'HTML 작성'</strong> 탭이나 HTML 직접 입력 모드에 붙여넣을 때 사용합니다.
  </div>

  <button class="copy-btn full-btn html-mode-btn" onclick="copyFromTextarea('full-html-ta', '🎉 전체 HTML 소스코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🚀 [HTML로 전체 일괄 복사] 4-Core 요약 + 비교표 + FAQ 전체 소스코드 복사</button>
  <textarea id="full-html-ta" style="display:none;">{full_html_code}</textarea>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📌 1. 공식 상품명 (Title)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec1-ta', '✅ 상품명이 복사되었습니다!')">📋 상품명 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec1-ta" rows="2" readonly>{s1_clean}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">🔬 2. 핵심 가치 및 5줄 마이크로 요약</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec2-ta', '✅ 5줄 요약 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec2-html-ta', '🌐 5줄 요약 HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec2-ta" rows="8" readonly>{s2_text_formatted}</textarea>
    <textarea id="sec2-html-ta" style="display:none;">{s2_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📊 3. 제품 상세 스펙 비교표 (HTML Table)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec3-ta', '✅ 비교표 텍스트가 복사되었습니다!')">📋 비교표 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec3-html-ta', '🌐 비교표 HTML 코드가 복사되었습니다!')">🌐 HTML로 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec3-ta" rows="10" readonly>{s3_text_formatted}</textarea>
    <textarea id="sec3-html-ta" style="display:none;">{s3_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">💬 4. 5대 핵심 FAQ & 상세 가이드</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec4-ta', '✅ FAQ 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec4-html-ta', '🌐 FAQ HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec4-ta" rows="18" readonly>{s4_text_formatted}</textarea>
    <textarea id="sec4-html-ta" style="display:none;">{s4_html}</textarea>
  </div>

</div>

<div class="toast" id="toast">✅ 클립보드에 복사되었습니다! 쇼핑몰 에디터에 붙여넣기(Ctrl+V) 하세요.</div>

<script>
function showToast(msg) {{
  const t = document.getElementById('toast');
  t.innerText = msg;
  t.style.display = 'block';
  setTimeout(() => {{ t.style.display = 'none'; }}, 3000);
}}

function copyFromTextarea(id, customMsg) {{
  const ta = document.getElementById(id);
  const text = ta.value;
  
  const temp = document.createElement('textarea');
  temp.value = text;
  temp.style.position = 'fixed';
  temp.style.left = '-9999px';
  document.body.appendChild(temp);
  temp.select();
  temp.setSelectionRange(0, 99999);
  
  try {{
    document.execCommand('copy');
    showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
  }} catch (err) {{
    if (navigator.clipboard) {{
      navigator.clipboard.writeText(text).then(() => {{
        showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
      }});
    }}
  }} finally {{
    document.body.removeChild(temp);
  }}
}}
</script>

</body>
</html>
"""
    with open(out_html_path, "w", encoding="utf-8") as f:
        f.write(html_content)


if __name__ == '__main__':
    asyncio.run(main_async())



def _generate_docx_file(title: str, text_content: str, out_docx_path: str, target_lang: str = "EN"):
    """MS Word 서식(.docx)으로 4-Core 상세페이지 완성 원고를 렌더링합니다."""
    try:
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        doc = docx.Document()
        for s in doc.sections:
            s.top_margin = Inches(0.8)
            s.bottom_margin = Inches(0.8)
            s.left_margin = Inches(0.8)
            s.right_margin = Inches(0.8)

        # Title Header
        p_title = doc.add_paragraph()
        r_t = p_title.add_run(f"🛒 {title} - E-Commerce PDP Master Copy")
        r_t.font.name = "맑은 고딕"
        r_t.font.size = Pt(16)
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(16, 44, 87)

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(14)
        r_sub = p_sub.add_run("이 문서는 각국 전자상거래 플랫폼(스마트스토어, 쿠팡, 아마존, 큐텐 등) 상세페이지 에디터에 그대로 복사하여 등록할 수 있는 100% 고객 노출용 원고입니다.")
        r_sub.font.name = "맑은 고딕"
        r_sub.font.size = Pt(9.5)
        r_sub.font.color.rgb = RGBColor(100, 116, 139)

        # Parse sections
        lines = [l.strip() for l in text_content.splitlines() if l.strip()]
        cur_sec = 0
        sec_buffers = {1: [], 2: [], 3: [], 4: []}

        for l in lines:
            if l.startswith("1."):
                cur_sec = 1
                sec_buffers[1].append(l)
            elif l.startswith("2."):
                cur_sec = 2
                sec_buffers[2].append(l)
            elif l.startswith("3."):
                cur_sec = 3
                sec_buffers[3].append(l)
            elif l.startswith("4."):
                cur_sec = 4
                sec_buffers[4].append(l)
            else:
                if cur_sec in sec_buffers:
                    sec_buffers[cur_sec].append(l)

        # Section 1: Title
        h1 = doc.add_heading(level=1)
        r_h1 = h1.add_run("1. 공식 상품명 (Official Product Title)")
        r_h1.font.name = "맑은 고딕"
        r_h1.font.size = Pt(12.5)
        r_h1.font.bold = True
        r_h1.font.color.rgb = RGBColor(37, 99, 235)

        t_val = " ".join([l for l in sec_buffers[1] if not l.startswith("1.")]).strip()
        if not t_val and sec_buffers[1]:
            t_val = sec_buffers[1][-1].replace("1.", "").strip()
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_after = Pt(10)
        r_tv = p_t.add_run(t_val)
        r_tv.font.name = "맑은 고딕"
        r_tv.font.size = Pt(10.5)
        r_tv.font.bold = True

        # Section 2: Summary
        h2 = doc.add_heading(level=1)
        r_h2 = h2.add_run("2. 핵심 가치 및 제품 안내 (Core Value & Summary)")
        r_h2.font.name = "맑은 고딕"
        r_h2.font.size = Pt(12.5)
        r_h2.font.bold = True
        r_h2.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[2]:
            if l.startswith("2."):
                continue
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(l.lstrip("- •*"))
            r.font.name = "맑은 고딕"
            r.font.size = Pt(9.5)

        # Section 3: Comparison Table
        h3 = doc.add_heading(level=1)
        r_h3 = h3.add_run("3. 제품 상세 스펙 비교 (Product Specifications & Comparison Table)")
        r_h3.font.name = "맑은 고딕"
        r_h3.font.size = Pt(12.5)
        r_h3.font.bold = True
        r_h3.font.color.rgb = RGBColor(37, 99, 235)

        raw_s3_text = "\n".join(sec_buffers[3])
        if "<table>" in raw_s3_text:
            rows_data = []
            tr_matches = re.findall(r'<tr>(.*?)</tr>', raw_s3_text, flags=re.DOTALL)
            for tr in tr_matches:
                cols = re.findall(r'<t[hd]>(.*?)</t[hd]>', tr, flags=re.DOTALL)
                if cols:
                    rows_data.append([re.sub(r'<[^>]+>', '', c).strip() for c in cols])
            
            if rows_data:
                col_cnt = max(len(r) for r in rows_data)
                tbl = doc.add_table(rows=len(rows_data), cols=col_cnt)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r_idx, r_cols in enumerate(rows_data):
                    for c_idx, c_val in enumerate(r_cols):
                        cell = tbl.rows[r_idx].cells[c_idx]
                        p = cell.paragraphs[0]
                        r = p.add_run(c_val)
                        r.font.name = "맑은 고딕"
                        r.font.size = Pt(9)
                        if r_idx == 0:
                            r.font.bold = True
                            tcPr = cell._element.get_or_add_tcPr()
                            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
                            tcPr.append(shd)
        else:
            for l in sec_buffers[3]:
                if not l.startswith("3."):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    r = p.add_run(l)
                    r.font.name = "맑은 고딕"
                    r.font.size = Pt(9.5)

        # Section 4: FAQ
        h4 = doc.add_heading(level=1)
        r_h4 = h4.add_run("4. 자주 묻는 질문 (Frequently Asked Questions)")
        r_h4.font.name = "맑은 고딕"
        r_h4.font.size = Pt(12.5)
        r_h4.font.bold = True
        r_h4.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[4]:
            if l.startswith("4."):
                continue
            if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = RGBColor(30, 58, 138)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(51, 65, 85)

        doc.save(out_docx_path)
    except Exception as de:
        print(f"  ⚠️ [DOCX 생성 폴백]: {de}", flush=True)


def _generate_web_copier_html_file(title: str, text_content: str, out_html_path: str, target_lang: str = "EN"):
    """지마켓/스마트스토어/쿠팡/해외몰 등 전 이커머스 플랫폼에 맞춘 4-Core '에디터에 붙여넣을 텍스트 복사' 및 'HTML로 텍스트 복사' 뷰어를 생성합니다."""
    
    section_titles = {
        "EN": {
            "sec2": "2. Core Value & Active Ingredient Summary",
            "sec3": "3. Product Specifications & Comparison Table",
            "sec4": "4. Product Usage Guide & Frequently Asked Questions (FAQ)"
        },
        "JP": {
            "sec2": "2. コアバリュー＆成分サイエンス要約",
            "sec3": "3. 製品仕様・他社比較テーブル",
            "sec4": "4. 使用ガイド＆よくある質問 (FAQ)"
        },
        "CN": {
            "sec2": "2. 核心价值与成分科技摘要",
            "sec3": "3. 产品规格与竞品对比表",
            "sec4": "4. 商品使用指南与常见问题解答 (FAQ)"
        },
        "TW": {
            "sec2": "2. 核心價值與成分科技摘要",
            "sec3": "3. 產品規格與競品對比表",
            "sec4": "4. 使用指南與常見問題解答 (FAQ)"
        },
        "KR": {
            "sec2": "2. 제품 핵심 안내 및 성분 요약",
            "sec3": "3. 제품 상세 비교 스펙 테이블",
            "sec4": "4. 자주 묻는 질문 (FAQ)"
        }
    }
    
    detected_lang = target_lang.upper() if target_lang else "EN"
    titles_map = section_titles.get(detected_lang, section_titles["EN"])
    sec2_heading = titles_map["sec2"]
    sec3_heading = titles_map["sec3"]
    sec4_heading = titles_map.get("sec4", "4. FAQ")

    lines = [l.strip() for l in text_content.splitlines() if l.strip()]
    
    s1_lines = []
    s2_lines = []
    s3_lines = []
    s4_lines = []
    cur_sec = 0
    for l in lines:
        if l.startswith("1."):
            cur_sec = 1
            s1_lines.append(l)
        elif l.startswith("2."):
            cur_sec = 2
            s2_lines.append(l)
        elif l.startswith("3."):
            cur_sec = 3
            s3_lines.append(l)
        elif l.startswith("4."):
            cur_sec = 4
            s4_lines.append(l)
        else:
            if cur_sec == 1:
                s1_lines.append(l)
            elif cur_sec == 2:
                s2_lines.append(l)
            elif cur_sec == 3:
                s3_lines.append(l)
            elif cur_sec == 4:
                s4_lines.append(l)

    # 1. Section 1 (Title)
    s1_clean = " ".join([l for l in s1_lines if not l.startswith("1.")]).strip()
    if not s1_clean and s1_lines:
        s1_clean = s1_lines[-1].replace("1.", "").strip()

    # 2. Section 2 (Summary)
    s2_clean_items = [l for l in s2_lines if not l.startswith("2.")]
    s2_text_formatted = f"{sec2_heading}\n\n" + "\n\n".join(s2_clean_items)
    
    s2_html_items = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s2_html_items.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>🔬 {sec2_heading}</h3>")
    for l in s2_clean_items:
        if ":" in l or "：" in l:
            k, v = re.split(r'[:：]', l, 1)
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'><strong style='color:#1e3a8a;'>{k.strip()}:</strong> {v.strip()}</p>")
        else:
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'>{l.strip()}</p>")
    s2_html_items.append("</div>")
    s2_html = "\n".join(s2_html_items)

    # 3. Section 3 (Table)
    s3_clean_items = [l for l in s3_lines if not l.startswith("3.")]
    s3_raw_block = "\n".join(s3_clean_items)
    s3_text_formatted = f"{sec3_heading}\n\n" + s3_raw_block
    
    if "<table>" in s3_raw_block:
        styled_table = s3_raw_block.replace("<table>", "<table style='width:100%; border-collapse:collapse; margin:10px 0; font-size:13.5px;'>")
        styled_table = styled_table.replace("<th>", "<th style='background:#f1f5f9; padding:10px 12px; border:1px solid #cbd5e1; color:#1e3a8a; text-align:left; font-weight:bold;'>")
        styled_table = styled_table.replace("<td>", "<td style='padding:9px 12px; border:1px solid #cbd5e1; color:#334155;'>")
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n{styled_table}\n</div>"
    else:
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n<pre style='white-space:pre-wrap; font-family:inherit; font-size:13.5px;'>{s3_raw_block}</pre>\n</div>"

    # 4. Section 4 (FAQ)
    s4_clean_items = [l for l in s4_lines if not l.startswith("4.")]
    s4_text_blocks = []
    s4_html_blocks = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s4_html_blocks.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 12px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>💬 {sec4_heading}</h3>")
    
    cur_q = ""
    cur_answers = []
    for l in s4_clean_items:
        if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
            if cur_q:
                q_text = cur_q
                a_text = "\n".join(cur_answers)
                s4_text_blocks.append(f"{q_text}\n{a_text}")
                s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
                for ans in cur_answers:
                    s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
            cur_q = l.strip()
            cur_answers = []
        else:
            sentences = [s.strip() for s in re.split(r'(?<=[。！？\.\?!])\s*', l) if s.strip()]
            for s in sentences:
                cur_answers.append(s)
                
    if cur_q:
        q_text = cur_q
        a_text = "\n".join(cur_answers)
        s4_text_blocks.append(f"{q_text}\n{a_text}")
        s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
        for ans in cur_answers:
            s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
    s4_html_blocks.append("</div>")

    s4_text_formatted = f"{sec4_heading}\n\n" + "\n\n".join(s4_text_blocks)
    s4_html = "\n".join(s4_html_blocks)

    full_html_code = f"""<!-- 다국어 E-Commerce 상세페이지 4-Core 마이크로-써머리 & 비교표 & FAQ -->
<div style="font-family:'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; max-width:860px; margin:0 auto; padding:10px 0; color:#1e293b;">
{s2_html}
{s3_html}
{s4_html}
</div>"""

    html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; background: #f1f5f9; color: #0f172a; padding: 25px; margin: 0; line-height: 1.6; }}
  .container {{ max-width: 960px; margin: 0 auto; background: #ffffff; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.06); padding: 35px; border: 1px solid #cbd5e1; }}
  h1 {{ font-size: 22px; color: #0f172a; border-bottom: 2px solid #2563eb; padding-bottom: 12px; margin-top: 0; display: flex; align-items: center; gap: 8px; }}
  .guide-box {{ background: #eff6ff; border: 1px solid #bfdbfe; border-left: 5px solid #2563eb; padding: 18px 20px; border-radius: 8px; margin-bottom: 25px; font-size: 14px; color: #1e40af; line-height: 1.8; }}
  .guide-box strong {{ color: #1e3a8a; }}
  .card {{ background: #f8fafc; border: 1px solid #cbd5e1; border-radius: 10px; padding: 22px; margin-bottom: 25px; }}
  .card-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 14px; border-bottom: 1px solid #e2e8f0; padding-bottom: 10px; }}
  .card-title {{ font-size: 16px; font-weight: bold; color: #1e3a8a; }}
  .btn-group {{ display: flex; gap: 8px; }}
  .copy-btn {{ background: #2563eb; color: #ffffff; border: none; padding: 9px 16px; border-radius: 6px; cursor: pointer; font-size: 13px; font-weight: 700; transition: all 0.2s; box-shadow: 0 2px 4px rgba(37,99,235,0.2); display: inline-flex; align-items: center; gap: 6px; }}
  .copy-btn:hover {{ background: #1d4ed8; transform: translateY(-1px); }}
  .copy-btn.html-mode-btn {{ background: #059669; box-shadow: 0 2px 4px rgba(5,150,105,0.25); }}
  .copy-btn.html-mode-btn:hover {{ background: #047857; }}
  .full-btn {{ background: #7c3aed; padding: 12px 24px; font-size: 15px; width: 100%; justify-content: center; margin-bottom: 20px; box-shadow: 0 3px 6px rgba(124,58,237,0.25); }}
  .full-btn:hover {{ background: #6d28d9; }}
  .text-area {{ width: 100%; box-sizing: border-box; background: #ffffff; border: 1px solid #cbd5e1; border-radius: 6px; padding: 14px 16px; font-size: 13.5px; line-height: 1.75; color: #1e293b; font-family: 'Malgun Gothic', 'Segoe UI', monospace; resize: vertical; outline: none; }}
  .text-area:focus {{ border-color: #2563eb; box-shadow: 0 0 0 3px rgba(37,99,235,0.15); }}
  .toast {{ position: fixed; bottom: 30px; right: 30px; background: #0f172a; color: #ffffff; padding: 14px 28px; border-radius: 8px; font-size: 14px; font-weight: 600; display: none; z-index: 1000; box-shadow: 0 6px 20px rgba(0,0,0,0.25); }}
</style>
</head>
<body>

<div class="container">
  <h1>🌐 {title}</h1>
  
  <div class="guide-box">
    📢 <strong>쇼핑몰 등록 방식별 2대 원클릭 복사 기능 안내:</strong><br>
    • <strong>1. [📋 에디터에 붙여넣을 텍스트 복사] (파란색 버튼)</strong>: 지마켓/스마트스토어/쿠팡/해외몰 <strong>'에디터 작성'</strong> 화면에 붙여넣을 때 사용합니다.<br>
    • <strong>2. [🌐 HTML로 텍스트 복사] (초록색 버튼)</strong>: <strong>'HTML 작성'</strong> 탭이나 HTML 직접 입력 모드에 붙여넣을 때 사용합니다.
  </div>

  <button class="copy-btn full-btn html-mode-btn" onclick="copyFromTextarea('full-html-ta', '🎉 전체 HTML 소스코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🚀 [HTML로 전체 일괄 복사] 4-Core 요약 + 비교표 + FAQ 전체 소스코드 복사</button>
  <textarea id="full-html-ta" style="display:none;">{full_html_code}</textarea>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📌 1. 공식 상품명 (Title)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec1-ta', '✅ 상품명이 복사되었습니다!')">📋 상품명 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec1-ta" rows="2" readonly>{s1_clean}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">🔬 2. 핵심 가치 및 5줄 마이크로 요약</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec2-ta', '✅ 5줄 요약 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec2-html-ta', '🌐 5줄 요약 HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec2-ta" rows="8" readonly>{s2_text_formatted}</textarea>
    <textarea id="sec2-html-ta" style="display:none;">{s2_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📊 3. 제품 상세 스펙 비교표 (HTML Table)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec3-ta', '✅ 비교표 텍스트가 복사되었습니다!')">📋 비교표 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec3-html-ta', '🌐 비교표 HTML 코드가 복사되었습니다!')">🌐 HTML로 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec3-ta" rows="10" readonly>{s3_text_formatted}</textarea>
    <textarea id="sec3-html-ta" style="display:none;">{s3_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">💬 4. 5대 핵심 FAQ & 상세 가이드</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec4-ta', '✅ FAQ 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec4-html-ta', '🌐 FAQ HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec4-ta" rows="18" readonly>{s4_text_formatted}</textarea>
    <textarea id="sec4-html-ta" style="display:none;">{s4_html}</textarea>
  </div>

</div>

<div class="toast" id="toast">✅ 클립보드에 복사되었습니다! 쇼핑몰 에디터에 붙여넣기(Ctrl+V) 하세요.</div>

<script>
function showToast(msg) {{
  const t = document.getElementById('toast');
  t.innerText = msg;
  t.style.display = 'block';
  setTimeout(() => {{ t.style.display = 'none'; }}, 3000);
}}

function copyFromTextarea(id, customMsg) {{
  const ta = document.getElementById(id);
  const text = ta.value;
  
  const temp = document.createElement('textarea');
  temp.value = text;
  temp.style.position = 'fixed';
  temp.style.left = '-9999px';
  document.body.appendChild(temp);
  temp.select();
  temp.setSelectionRange(0, 99999);
  
  try {{
    document.execCommand('copy');
    showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
  }} catch (err) {{
    if (navigator.clipboard) {{
      navigator.clipboard.writeText(text).then(() => {{
        showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
      }});
    }}
  }} finally {{
    document.body.removeChild(temp);
  }}
}}
</script>

</body>
</html>
"""
    with open(out_html_path, "w", encoding="utf-8") as f:
        f.write(html_content)