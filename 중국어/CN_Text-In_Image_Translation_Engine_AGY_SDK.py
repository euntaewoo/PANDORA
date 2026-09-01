import os
def get_recursive_files(base_dir):
    try:
        res = []
        for root, _, files in os.walk(base_dir):
            for f in files:
                res.append(os.path.relpath(os.path.join(root, f), base_dir))
        return res
    except Exception:
        return []

import sys
sys.path.insert(0, r"C:\Users\euntaewoo\Desktop\multilingual_text_in_image_translatio_agy_sdk\multilingual_text_in_image_translatio_agy_sdk_core")
from multilingual_transcreation_qa_evaluator_agy_sdk import evaluate_transcreation, generate_html_report

import os
import io
import sys
import time
import json
import re
import argparse
import asyncio
from google import genai
from google.genai import types
from PIL import Image

sys.stdout.reconfigure(encoding='utf-8')

# ==========================================
# 0. 경로 및 Google Cloud 인증키 설정
# ==========================================
script_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.abspath(os.path.join(script_dir, ".."))
common_dir = os.path.join(project_root, "00_공통자료")
if common_dir not in sys.path:
    sys.path.insert(0, common_dir)

# .env 탐색
env_path = os.path.join(project_root, ".env")
api_key = None
gcp_json_key = None

if os.path.exists(env_path):
    with open(env_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line.startswith("GEMINI_API_KEY="):
                api_key = line.split("=", 1)[1].strip()
            elif line.startswith("GOOGLE_APPLICATION_CREDENTIALS="):
                gcp_json_key = line.split("=", 1)[1].strip().strip('"').strip("'")




GLOBAL_COMPLIANCE_SYSTEM_INSTRUCTION = """[SYSTEM INSTRUCTION: Global Cross-Border E-Commerce Compliance & Prestige Beauty Transcreation Expert (Chinese Mode)]
당신은 중국 NMPA 화장품감독관리조례, 신광고법 및 대만 TFDA 규정을 완벽히 준수하는 15년 차 글로벌 뷰티 법무 감사관이자 샤오홍슈/티몰 럭셔리 수석 카피라이터입니다.

[엄격 실행 대원칙]
1. [신광고법 8대 절대화 금지어 전면 배제]: '最', '第一', '顶级', '极品', '永久', '万能', '100%', '彻底' 등 절대어 사용을 엄격히 금지하고 프리미엄 케어 어휘('优', '前沿', '高端' 등)로 순화하십시오.
2. [의료 및 세포 치료 오인 차단]: '细胞再生', '根除皱纹', '消炎抗敏' 등 의료 클레임을 100% 차단하고 '修护屏障', '淡化细纹', '舒缓修护'로 안전하게 표현하십시오.
3. [고시정보표 법정 조항]: 한국 식약처(MFDS) 심사필, 3대 주의사항, 공정위 분쟁기준, +82 고객상담번호를 표준화하십시오.
"""

async def main_async():
    if not api_key:
        api_key = os.environ.get("GEMINI_API_KEY")
    if not gcp_json_key:
        gcp_json_key = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")

    # 루트 폴더 기준 인증키 fallback 탐색
    key_candidates = [
        os.path.join(project_root, ".env"),
        os.path.join(project_root, "00_공통자료", "APIs_KEY", "인증키_및_계정", "김차장_vertex api_key", "vertex_ai_auth_key.json"),
        os.path.join(project_root, "00_공통자료", "인증키_및_계정", "김차장_vertex api_key", "vertex_ai_auth_key.json"),
    ]

    for kpath in key_candidates:
        if kpath.endswith(".json") and os.path.exists(kpath):
            gcp_json_key = kpath
            break

    if gcp_json_key and os.path.exists(gcp_json_key):
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = gcp_json_key
        print(f"[INFO] Vertex AI 서비스 계정 JSON 키 감지: {gcp_json_key}", flush=True)
        with open(gcp_json_key, 'r', encoding='utf-8') as f:
            key_data = json.load(f)
            project_id = key_data.get('project_id')
        client = genai.Client(vertexai=True, project=project_id, location="global")
        print(f"[INFO] Vertex AI Client 연결 성공 (Project: {project_id}, Location: global)", flush=True)
    elif api_key:
        if api_key.startswith("AQ."):
            print("[INFO] Agent Platform API 키(AQ...) 감지. Vertex AI 모드로 전환합니다.", flush=True)
            client = genai.Client(vertexai=True, api_key=api_key)
        else:
            client = genai.Client(api_key=api_key)
    else:
        print("[ERROR] GEMINI_API_KEY 또는 GOOGLE_APPLICATION_CREDENTIALS가 설정되지 않았습니다.", flush=True)
        sys.exit(1)


    # ==========================================
    # 1. 모델 및 파라미터 설정
    # ==========================================
    MODEL_PRO = "gemini-3.1-pro-preview"       # Pass 1: 텍스트 추출 및 중국 광고법 검열 번역
    MODEL_FLASH_IMAGE = "gemini-3.1-flash-image"  # Pass 2: 시각적 인페인팅 및 알리바바 푸후이체 식자

    # 커맨드라인 인자 파싱
    parser = argparse.ArgumentParser(description="CN Text-In-Image Translation Engine V1")
    parser.add_argument("source_dir", nargs="?", default=os.path.join(project_root, "01_번역대상_원본"), help="원본 이미지 디렉터리")
    parser.add_argument("target_dir", nargs="?", default=None, help="결과물 저장 디렉터리")
    parser.add_argument("--target", choices=["CN", "TW", "HK"], default="CN", help="타겟 권역 (CN: 중국본토 간체, TW: 대만 번체, HK: 홍콩 번체)")

    args = parser.parse_args()

    source_dir = os.path.abspath(args.source_dir)
    target_region = args.target.upper()

    # 디렉터리명이나 파일명에서 자동 권역 힌트 감지
    dir_name_lower = os.path.basename(source_dir).lower()
    if "대만" in dir_name_lower or "tw" in dir_name_lower or "번체" in dir_name_lower:
        target_region = "TW"
    elif "홍콩" in dir_name_lower or "hk" in dir_name_lower:
        target_region = "HK"
    elif "중국" in dir_name_lower or "본토" in dir_name_lower or "간체" in dir_name_lower:
        target_region = "CN"

    if args.target_dir:
        target_dir = os.path.abspath(args.target_dir)
    else:
        folder_name = os.path.basename(source_dir)
        target_dir = os.path.join(project_root, "02_번역결과_최종", f"{folder_name}_CN_{target_region}_Translated")

    os.makedirs(target_dir, exist_ok=True)

    print(f"[START] CN_Text-In_Image_Translation_Engine_V1 (Two-Pass Architecture) 가동...")
    print(f"[CONFIG] 타겟 권역 모드: {target_region} ({'🇨🇳 중국 본토 간체자 (zh-CN)' if target_region == 'CN' else '🇹🇼 대만 번체자 (zh-TW)' if target_region == 'TW' else '🇭🇰 홍콩 번체자 (zh-HK)'})")
    print(f"[CONFIG] 입력 폴더: {source_dir}")
    print(f"[CONFIG] 출력 폴더: {target_dir}")

    # ==========================================
    # 2. 중국 신광고법 및 NMPA 화장품 규정 필터링 정의
    # ==========================================
    # 중국 신광고법 8대 절대화 금지어 및 화장품 위반 어휘 -> 합법적 순화어 매핑
    CN_AD_LAW_FILTERS = {
        r"最": "优",
        r"第一": "前沿",
        r"顶级": "高端",
        r"极品": "优选",
        r"永久": "持久",
        r"彻底": "深层",
        r"万能": "多效",
        r"根除": "改善",
        r"消炎": "舒缓",
        r"镇静": "舒缓修护",
        r"抗衰老": "紧致淡纹",
        r"去皱": "抚平细纹",
        r"再生": "赋活",
        r"无刺激": "温和低敏",
        r"100%安全": "温和配方",
        r"排毒": "净澈肌肤",
        r"美白": "焕亮透光"
    }

    # ==========================================
    # 3. [Pass 1] 권역별 프롬프트 정의
    # ==========================================
    if target_region == "CN":
        lang_name = "Simplified Chinese (简体中文, zh-CN)"
        region_guidelines = """
    [중국 본토 (간체자) 신광고법 및 이커머스 필수 지침]
    1. [간체자 강제] 모든 번역문은 반드시 순수 간체자(Simplified Chinese)로 작성하십시오. 번체자 절대 혼용 금지.
    2. [중국 신(新) 광고법 절대화 표현 전면 금지]
       - '최고/제일/최상급' 표현 절대 불가: 最, 第一, 顶级, 极品, 极致, 独家 등 금지 -> '卓越', '优异', '高端', '精心' 등으로 순화.
       - '영구/완벽/완전' 표현 금지: 永久, 彻底, 根除, 100%, 零刺激 -> '持久', '深层改善', '温和低敏' 등으로 순화.
    3. [NMPA 화장품 효능 표기 가이드]
       - '치료/의약적 효능' 오인 표현 절대 금지: 治疗, 消炎, 镇静, 修复疤痕 -> '舒缓', '修护', '净澈', '维持肌肤稳定' 등으로 대체.
       - '안티에이징/주름제거' -> '紧致淡纹', '改善干纹', '丰盈饱满'.
       - '미백/잡티제거' -> '焕亮肌肤', '改善暗沉', '提亮肤色'.
    4. [이커머스 표준 용어 (타오바오/티몰/샤오홍슈)]
       - 수분/보습: 补水 / 保湿
       - 탄력/리프팅: 紧致 / 提拉 / 弹润
       - 진정/피부장벽: 舒缓 / 屏障修护
    """
    elif target_region == "TW":
        lang_name = "Traditional Chinese for Taiwan (台灣繁體中文, zh-TW)"
        region_guidelines = """
    [대만 번체자 (zh-TW) TFDA 화장품법 및 이커머스 지침]
    1. [대만 표준 번체자] 모든 번역문은 반드시 대만 정체자(Traditional Chinese, zh-TW)로 작성하십시오. 간체자 혼용 절대 금지.
    2. [대만 뷰티 전문 용어 (Shopee TW / momo 최적화)]
       - 토너/스킨: 化妝水
       - 수분잠금/보습: 鎖水 / 保濕 / 補水
       - 에센스/앰플: 精華液 / 安瓶
       - 피부 탄력: 緊緻 / 澎潤 / 彈力
       - 진정/피부결: 舒緩修護 / 調理肌膚
       - 물광/윤광: 水光肌 / 透亮光澤
    3. [TFDA 규정 준수] 의학적 치료, 영구 재생 등 과대광고 표현을 배제하고 부드럽고 신뢰도 높은 어조 적용.
    """
    else: # HK
        lang_name = "Traditional Chinese for Hong Kong (香港繁體中文, zh-HK)"
        region_guidelines = """
    [홍콩 번체자 (zh-HK) 이커머스 지침]
    1. [홍콩 표준 번체자] 모든 번역문은 홍콩 번체자(Traditional Chinese, zh-HK)로 작성하십시오.
    2. [홍콩 뷰티 전문 용어 (HKTVmall / Mannings 최적화)]
       - 토너/스킨: 爽膚水
       - 보습/수분: 補濕 / 保濕
       - 에센스/세럼: 精華素 / 精華
       - 클렌징워터: 卸妝水 / 落妝水
       - 피부결: 提亮 / 緊緻 / 舒緩
    3. 직관적이고 세련된 홍콩 이커머스 마케팅 톤앤매너 적용.
    """

    pass1_prompt = f"""
    [SYSTEM PROMPT] Global Luxury Beauty Transcreation & Compliance Expert (Chinese Engine - {lang_name})

    ## 1. 시스템 역할 및 콘셉트 (Role & Context)
    당신은 에스티로더, 랑콤, 시슬리, SK-II 등 중화권 하이엔드 럭셔리 뷰티 시장을 총괄하는 10년 차 수석 크리에이티브 디렉터이자 엘리트 카피라이터입니다.
    단순 직역을 배제하고, 현지 소비자가 열광하는 프리미엄 뷰티 서사로 '초월번역(Transcreation)'을 수행하세요.

    ## 2. 초월번역 핵심 원칙 (Core Transcreation Principles)
    1. [기계적 직역 및 부사 금지]
       - '确实', '真正', '非常' 등 딱딱한 부사 직역을 전면 금지하고, 럭셔리 뷰티 전문 어휘로 세련되게 재창조하십시오.
    2. [자연스러운 구문 결속 및 활성 성분 연결]
       - "10% LiftDerm" 등 성분 비율이 문맥과 끊기지 않고 제품 효능 및 서사로 매끄럽게 연결되도록 문장 구조를 재조정하십시오.
    3. [4대 기능성 뷰티 전문 어휘 사전]
       - 피부 속/기저층: 肌底 / 肌底深处
       - 토탈 케어/멀티 코렉티브: 多效修护 (번체: 多效修護)
       - 탄력 복원/강화: 赋活肌底弹力 (번체: 賦活肌底彈力)
       - 눈가 잔주름/건조주름: 细纹・干纹 (번체: 細紋・乾紋)
    4. [독자 성분명 영문 보존]
       - 'LiftDerm', 'Lifting Logic for eye' 등 글로벌 독자 성분명은 영문 그대로 유지하되 문맥과 완벽히 융합하십시오.

    ## 3. 타깃 권역별 필수 지침 및 규제 준수
    {region_guidelines}

    [절대 불변 및 규제 검열 원칙 (Regulatory Guardrails)]
    1. [절대적/과대 표현 전면 금지 (Ban on Absolute Claims)]
       - 'World's First', 'No.1', 'Best', 'The Ultimate' 등 검증 불가능한 절대 표현(`全球首創`, `第一`, `最佳`, `終極對策` 등)을 엄격히 금지합니다.
       - 반드시 혁신 기술 및 프리미엄 케어 용어로 의무 순화하십시오 (예: `專為...研發의 創新科技`, `頂級多效`, `精準修護`).
    2. [의료 시술 오인 금지 및 4대 안전 동사 (Compliance-Safe Verbs)]
       - '주름 박멸', '영구 삭제', '보톡스/필러 효과' 등 의료 시술 오인 및 세포 치료/재생 과장 표현 전면 배제.
       - 반드시 화장품 규정 내 안전 동사인 **`撫平` (Smooth), `淡化` (Diminish), `舒緩` (Alleviate), `修護` (Care/Repair)**만을 사용하여 표현하십시오.
    3. [패키지 영문/로고 100% 보존] 제품 본품(용기, 튜브, 단상자 등)에 인쇄된 영문 텍스트(예: LOGICALLY SKIN, 제품 영문명 등)와 브랜드 로고는 절대 번역 매핑에 넣지 마세요. 원본 픽셀을 그대로 유지해야 합니다.
    4. [전수 추출] 이미지 내의 모든 한국어 텍스트는 단 하나도 빠짐없이 100% 추출하여 번역 매핑에 포함시키십시오.

    출력은 반드시 JSON 형식으로 아래 스키마를 엄격히 따르세요:
    {{
      "translation_map": [
        {{
          "kor": "한국어 원문",
          "chn": "광고법 준수 럭셔리 중국어 번역문",
          "violation_reason": "광고법/규제 순화 사유 (수정한 경우 기재, 없으면 빈 문자열)"
        }}
      ],
      "required_footnotes": [
        "필요한 법적 주석 문자열 (없으면 빈 배열)"
      ]
    }}
    """

    # ==========================================
    # 4. [Pass 2] 렌더링 지시 프롬프트 템플릿
    # ==========================================
    if target_region in ["TW", "HK"]:
        pass2_prompt_template = """
    당신은 정밀한 시각적 로컬라이제이션을 수행하는 이미지 인페인팅 AI입니다.
    첨부된 원본 이미지 속의 텍스트 위치, 배경 텍스처, 제품 누끼, 디자인 레이아웃을 1픽셀의 왜곡 없이 그대로 유지하세요.
    아래에 제공된 [번역 매핑 데이터 JSON]을 바탕으로 다음 규칙을 엄격히 적용하여 단일 이미지를 생성하세요.

    [시각적 렌더링 엄격 규칙 - 대만/홍콩 번체자(繁體中文) 전용]
    1. (KOR ERASING) 원본의 한국어 텍스트는 원래 자리에 남겨두지 말고 배경색/텍스처로 완벽하게 덮어써서 100% 지울 것. 병기 절대 금지.
    2. (JSON APPLY) 지워진 그 자리에 오직 [번역 매핑 데이터 JSON]의 'chn' 텍스트만 렌더링할 것. 모델 임의로 글자를 누락하거나 수정하지 말 것.
    3. (FONT STYLE) 폰트는 중화권 최고급 표준 번체 서체인 'Noto Sans TC (스위안헤이티 번체 / 思源黑體 / Source Han Sans TC)' 스타일의 정갈하고 모던한 산세리프로 선명하게 렌더링할 것.
    4. (CHINESE E-COMMERCE LAYOUT RULES - 3대 실전 팁)
       - [본문 크기 줄이기]: 한자는 Em-box를 꽉 채우므로, 한국어 원본 대비 폰트 크기를 약 10~15% 슬림하게 낮추어 여백과 밸런스를 완벽하게 유지할 것.
       - [행간(Line-Height) 15~20% 확장]: 한자가 상하로 빽빽하게 붙어 답답해지지 않도록 한국어 대비 행간을 15~20% 더 넓게 여유 있게 설정할 것.
       - [자간(Letter-Spacing) 여유 확보]: 글자가 뭉쳐 보이지 않도록 자간에 은은한 여백을 주어 고급스러운 프리미엄 뷰티 브랜드 상세페이지 느낌을 극대화할 것.
    5. (🚨 CRITICAL: ABSOLUTE TRADITIONAL GLYPH ENFORCEMENT)
       - 절대 중국 본토 간체자(Simplified Chinese) 획수를 그리지 마십시오! 
       - 확산 모델의 간체자 쏠림(Drift)을 엄격히 차단하고, 반드시 획수가 많은 순수 번체자(繁體字/正體字) 글리프를 정확히 그리십시오:
         * 養 (O) vs 养 (X - 절대금지) -> 保養, 營養, 調理
         * 對 (O) vs 对 (X - 절대금지) -> 對策, 針對
         * 護 (O) vs 护 (X - 절대금지) -> 修護, 護理
         * 創 (O) vs 创 (X - 절대금지) -> 首創, 創造
         * 變 (O) vs 变 (X - 절대금지) -> 變得, 變化
         * 顯 (O) vs 显 (X - 절대금지) -> 明顯
         * 實 (O) vs 实 (X - 절대금지) -> 確實, 實驗
         * 體 (O) vs 体 (X - 절대금지) -> 體驗, 身體
         * 驗 (O) vs 验 (X - 절대금지) -> 體驗, 實驗
         * 緊 (O) vs 紧 (X - 절대금지) -> 緊緻
         * 緻 (O) vs 致 (X - 절대금지) -> 細緻, 緊緻
         * 膚 (O) vs 肤 (X - 절대금지) -> 肌膚, 膚質
         * 雙 (O) vs 双 (X - 절대금지) -> 雙重
         * 氣 (O) vs 气 (X - 절대금지) -> 空氣
         * 隊 (O) vs 队 (X - 절대금지) -> 團隊
         * 劃 (O) vs 划 (X - 절대금지) -> 企劃
         * 歲 (O) vs 岁 (X - 절대금지) -> 30歲
         * 乾 (O) vs 干 (X - 절대금지) -> 乾燥, 乾性
         * 華 (O) vs 华 (X - 절대금지) -> 精華液
         * 濕 (O) vs 湿 (X - 절대금지) -> 保濕
         * 鎖 (O) vs 锁 (X - 절대금지) -> 鎖水
         * 膠 (O) vs 胶 (X - 절대금지) -> 膠原蛋白
         * 纖 (O) vs 纤 (X - 절대금지) -> 纖維
    6. (FULL INPAINTING NO PATCHING) 텍스트 수정 시 오류 부분만 오려내어 덧칠(Patching)하지 말고, 캔버스 전체를 완전히 새롭게 렌더링(Full Inpainting)하여 1픽셀의 이질감도 없는 완벽한 하나의 이미지를 생성할 것.
    7. (PACKAGE PRESERVATION) 제품 본품(용기, 튜브, 박스 등) 표면에 인쇄된 영문 텍스트(예: LOGICALLY SKIN 등) 및 브랜드 로고는 절대 다시 그리거나 훼손하지 말고 100% 완벽하게 보존할 것.
    8. (LAYOUT STRICTNESS) 원본 텍스트의 정렬축(좌/우/중앙), 폰트 두께감, 단락 간격을 정확하게 유지할 것.
    9. (NO EXTRA NOISE) 번역과 무관한 AI 주석이나 영어 설명, 괄호를 이미지에 임의로 추가하지 말 것.

    [번역 매핑 데이터 JSON]
    {json_data}
    """
    else:
        pass2_prompt_template = """
    당신은 정밀한 시각적 로컬라이제이션을 수행하는 이미지 인페인팅 AI입니다.
    첨부된 원본 이미지 속의 텍스트 위치, 배경 텍스처, 제품 누끼, 디자인 레이아웃을 1픽셀의 왜곡 없이 그대로 유지하세요.
    아래에 제공된 [번역 매핑 데이터 JSON]을 바탕으로 다음 규칙을 엄격히 적용하여 단일 이미지를 생성하세요.

    [시각적 렌더링 엄격 규칙 - 중국 본토 간체자(简体中文) 전용]
    1. (KOR ERASING) 원본의 한국어 텍스트는 원래 자리에 남겨두지 말고 배경색/텍스처로 완벽하게 덮어써서 100% 지울 것. 병기 절대 금지.
    2. (JSON APPLY) 지워진 그 자리에 오직 [번역 매핑 데이터 JSON]의 'chn' 텍스트만 렌더링할 것. 모델 임의로 글자를 누락하거나 수정하지 말 것.
    3. (FONT STYLE) 폰트는 중화권 최고급 표준 서체인 'Noto Sans SC (스위안헤이티 / 思源黑体 / Source Han Sans SC)' 스타일의 정갈하고 모던한 산세리프로 선명하게 렌더링할 것.
    4. (CHINESE E-COMMERCE LAYOUT RULES - 3대 실전 팁)
       - [본문 크기 줄이기]: 한자는 Em-box를 꽉 채우므로, 한국어 원본 대비 폰트 크기를 약 10~15% 슬림하게 낮추어 여백과 밸런스를 완벽하게 유지할 것.
       - [행간(Line-Height) 15~20% 확장]: 한자가 상하로 빽빽하게 붙어 답답해지지 않도록 한국어 대비 행간을 15~20% 더 넓게 여유 있게 설정할 것.
       - [자간(Letter-Spacing) 여유 확보]: 글자가 뭉쳐 보이지 않도록 자간에 은은한 여백을 주어 고급스러운 프리미엄 뷰티 브랜드 상세페이지 느낌을 극대화할 것.
    5. (FULL INPAINTING NO PATCHING) 텍스트 수정 시 오류 부분만 오려내어 덧칠(Patching)하지 말고, 캔버스 전체를 완전히 새롭게 렌더링(Full Inpainting)하여 1픽셀의 이질감도 없는 완벽한 하나의 이미지를 생성할 것.
    6. (PACKAGE PRESERVATION) 제품 본품(용기, 튜브, 박스 등) 표면에 인쇄된 영문 텍스트(예: LOGICALLY SKIN 등) 및 브랜드 로고는 절대 다시 그리거나 훼손하지 말고 100% 완벽하게 보존할 것.
    7. (LAYOUT STRICTNESS) 원본 텍스트의 정렬축(좌/우/중앙), 폰트 두께감, 단락 간격을 정확하게 유지할 것.
    8. (NO EXTRA NOISE) 번역과 무관한 AI 주석이나 영어 설명, 괄호를 이미지에 임의로 추가하지 말 것.

    [번역 매핑 데이터 JSON]
    {json_data}
    """

    # ==========================================
    # 5. 이미지 일괄 번역 루프
    # ==========================================
    targets = sorted(
        [f for f in get_recursive_files(source_dir) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.jfif', '.webp'))],
        key=lambda x: [int(c) if c.isdigit() else c.lower() for c in re.split(r'(\d+)', x)]
    )

    if not targets:
        print(f"[WARNING] '{source_dir}' 폴더에 처리할 이미지가 없습니다.")
        sys.exit(0)

    all_translations = []

    for filename in targets:
        # 이미 번역된 파일 또는 고시표 텍스트 파일 등 스킵
        if '_CN_' in filename or '_JP_' in filename or '_EN_' in filename or filename.startswith('01_test_'):
            continue

        # 고시정보 표 이미지 별도 처리 (HTML 렌더러 연동)
        if '고시' in filename or 'KR' in filename or '상세정보' in filename or 'spec' in filename.lower():
            print(f"\n[NOTICE TABLE DETECTED] 고시정보 표 감지: {filename}", flush=True)
            in_path = os.path.join(source_dir, filename)
            out_name = f"{os.path.splitext(filename)[0]}_CN_{target_region}_v1.png"
            out_path = os.path.join(target_dir, out_name)
            if os.path.exists(out_path):
                print(f"  -> [SKIP] 이미 번역 완료된 고시표 파일입니다: {filename}")
                continue
            try:
                import render_notice_table_standard as rnts
                from google.cloud import vision
                v_client = vision.ImageAnnotatorClient()
                with open(in_path, "rb") as f_img:
                    c_data = f_img.read()
                v_img = vision.Image(content=c_data)
                v_res = v_client.document_text_detection(image=v_img)
                ocr_t = v_res.full_text_annotation.text if v_res.full_text_annotation else ""

                t_prompt = f"""
    你是中国国家药监局(NMPA)及新广告法合规专家。
    请将以下韩国化妆品告示表OCR文本翻译并规范化为符合中国法规的简体中文告示表JSON。

    [필수 표준 레이블 매핑 규격]
    - 용량 또는 중량 -> 净含量
    - 제품 주요 사양 -> 适用肤质
    - 사용기한 또는 개봉 후 사용기간 -> 使用期限
    - 사용방법 -> 使用方法
    - 화장품제조업자 및 책임판매업자 -> 化妆品生产企业 / 责任销售商
    - 제조국 -> 原产国
    - 전성분 -> 全成分 (INCI 및 KCID 기반 중국 NMPA 표준 명칭 적용)
    - 기능성 화장품 심사 필 유무 -> 特殊用途化妆品审查状态
    - 사용할 때의 주의사항 -> 使用注意事项
    - 소비자 상담 전화번호 -> 消费者咨询电话 (+82-2-6743-3206)

    [OCR 텍스트 원본]
    {ocr_t}

    반드시 순수 JSON 객체로만 출력하십시오:
    {{
      "title": "商品基本信息",
      "items": [
        {{"label": "净含量", "value": "25ml"}}
      ]
    }}
    """
                t_res = await client.aio.models.generate_content(
                    model=MODEL_PRO,
                    contents=[t_prompt],
                    config=types.GenerateContentConfig(
                        system_instruction=GLOBAL_COMPLIANCE_SYSTEM_INSTRUCTION,
                        response_mime_type="application/json",
                        temperature=0.6,
                        top_p=0.9,
                        max_output_tokens=8192
                    )
                )
                raw_t = t_res.text.strip()
                if raw_t.startswith("```"):
                    raw_t = re.sub(r"^```(?:json)?\s*", "", raw_t)
                    raw_t = re.sub(r"\s*```$", "", raw_t)
                t_data = json.loads(raw_t)
                n_items = t_data.get("items", [])
                n_title = t_data.get("title", "商品基本信息")
                n_items = standardize_notice_table_items(n_items, "CN" if target_region == "Simp" else "TW")
                rnts.render_notice_table_to_png(n_title, n_items, out_path, lang=target_region)
                print(f"  -> [NOTICE TABLE SUCCESS] 고시표 렌더링 완료: {out_name}")
                continue
            except Exception as te:
                print(f"  -> [NOTICE TABLE ERROR] 고시표 렌더링 실패 ({filename}): {te}")
                continue

        in_path = os.path.join(source_dir, filename)
        out_name = f"{os.path.splitext(filename)[0]}_CN_{target_region}_v1.png"
        out_path = os.path.join(target_dir, out_name)


        if os.path.exists(out_path):
            print(f"\n[SKIP] 이미 번역 완료된 파일입니다: {filename}")
            continue

        print(f"\n[RENDER] 변환 시작: {filename}", flush=True)

        try:
            original_image = Image.open(in_path)
            original_image.load()
        except Exception as e:
            print(f"  -> [ERROR] 이미지 로드 실패: {e}", flush=True)
            continue

        # ==========================
        # PASS 1: OCR & 번역 매핑 (Pro 모델)
        # ==========================
        print(f"  -> [PASS 1] 텍스트 매핑 및 {target_region} 권역 규제 검열 중...", flush=True)
        mapping_data_str = None
        for attempt in range(3):
            try:
                response_p1 = await client.aio.models.generate_content(
                    model=MODEL_PRO,
                    contents=[original_image, pass1_prompt],
                    config=types.GenerateContentConfig(
                        system_instruction=GLOBAL_COMPLIANCE_SYSTEM_INSTRUCTION,
                        response_mime_type="application/json",
                        temperature=0.6,
                        top_p=0.9,
                        max_output_tokens=8192
                    )
                )
                mapping_data_str = response_p1.text
                break
            except Exception as e:
                if "429" in str(e) or "ResourceExhausted" in str(e):
                    wait_time = 25 * (attempt + 1)
                    print(f"  -> [RATE LIMIT] 429 감지. {wait_time}초 대기 후 재시도... ({attempt+1}/3)", flush=True)
                    time.sleep(wait_time)
                else:
                    print(f"  -> [PASS 1 ERROR] {e}", flush=True)
                    break

        if not mapping_data_str:
            print("  -> [PASS 1 FAILED] 번역 매핑 데이터 생성 실패.", flush=True)
            continue

        # Python 하드 필터링 (중국 본토 모드일 때 신광고법 추가 검열)
        try:
            parsed_json = json.loads(mapping_data_str)
            if "translation_map" in parsed_json:
                for item in parsed_json["translation_map"]:
                    chn_text = item.get("chn", "")
                    if target_region == "CN":
                        for pattern, safe_word in CN_AD_LAW_FILTERS.items():
                            if re.search(pattern, chn_text):
                                print(f"      [Python Regex Filter] 신광고법 금지어 감지: '{pattern}' -> '{safe_word}' 로 강제 치환", flush=True)
                                chn_text = re.sub(pattern, safe_word, chn_text)
                                item["violation_reason"] = item.get("violation_reason", "") + f" (Python 정규식 치환: {pattern})"
                    item["chn"] = chn_text
                    item["source_file"] = filename
                all_translations.extend(parsed_json["translation_map"])
                mapping_data_str = json.dumps(parsed_json, ensure_ascii=False, indent=2)
            print("  -> [PASS 1 SUCCESS] 매핑 데이터 생성 및 검열 완료.", flush=True)
        except Exception as e:
            print(f"  -> [WARNING] JSON 파싱 경고: {e}", flush=True)

        # ==========================
        # PASS 2: 이미지 인페인팅 렌더링 (Flash-Image 모델)
        # ==========================
        print("  -> [PASS 2] 이미지 인페인팅 및 Noto Sans SC(스위안헤이티) 식자 렌더링 중...", flush=True)
        img_saved = False
        for attempt in range(3):
            try:
                final_prompt = pass2_prompt_template.replace("{json_data}", mapping_data_str)
                response_p2 = await client.aio.models.generate_content(
                    model=MODEL_FLASH_IMAGE,
                    contents=[final_prompt, original_image],
                    config=types.GenerateContentConfig(
                        response_modalities=["IMAGE"],
                        temperature=0.6,
                        top_p=0.9
                    )
                )

                if hasattr(response_p2, 'candidates'):
                    for cand in response_p2.candidates:
                        if hasattr(cand, 'content') and hasattr(cand.content, 'parts'):
                            for part in cand.content.parts:
                                if hasattr(part, 'inline_data') and part.inline_data:
                                    img = Image.open(io.BytesIO(part.inline_data.data))
                                    img = img.resize(original_image.size, Image.Resampling.LANCZOS)
                                    img.save(out_path, format="PNG")
                                    img_saved = True
                                    break
                                elif hasattr(part, 'image') and part.image:
                                    img = Image.open(io.BytesIO(part.image.image_bytes))
                                    img = img.resize(original_image.size, Image.Resampling.LANCZOS)
                                    img.save(out_path, format="PNG")
                                    img_saved = True
                                    break

                if img_saved:
                    print(f"  -> [SUCCESS] {out_name} 최종 저장 완료! (해상도: {original_image.size[0]}x{original_image.size[1]} px)", flush=True)
                    break
                else:
                    print("  -> [FAILED] Pass 2에서 이미지 데이터를 반환받지 못했습니다.", flush=True)
                    break

            except Exception as e:
                if "429" in str(e) or "ResourceExhausted" in str(e):
                    wait_time = 25 * (attempt + 1)
                    print(f"  -> [PASS 2 RATE LIMIT] 429 감지. {wait_time}초 대기 후 자동 재시도... ({attempt+1}/3)", flush=True)
                    time.sleep(wait_time)
                else:
                    print(f"  -> [PASS 2 ERROR] 렌더링 실패: {e}", flush=True)
                    break

        # 429 방지를 위한 8초 안전 대기
        time.sleep(8)



    # ==========================================
    # 6. 중국 광고법 준수 및 번역 비교표 리포트 생성
    # ==========================================
    if all_translations:
        print("\n[REPORT] 중국 광고법 준수 및 번역 비교표 TXT 리포트 생성 중...")
        report_path = os.path.join(target_dir, f"중국어_{target_region}_번역_비교표.txt")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write("==================================================\n")
            f.write(f"🇨🇳 CN_Text-In_Image_Translation_Engine_V1 ({target_region} 모드)\n")
            f.write("중국 신(新) 광고법 준수 및 번역 대조표 리포트\n")
            f.write("==================================================\n")
            f.write(f"타겟 권역: {target_region} ({'중국 본토 간체자 zh-CN' if target_region == 'CN' else '대만 번체자 zh-TW' if target_region == 'TW' else '홍콩 번체자 zh-HK'})\n")
            f.write(f"표준 폰트: Noto Sans SC (스위안헤이티 / 思源黑体 / Source Han Sans SC)\n\n")
            for t in all_translations:
                kor = t.get("kor", "").replace("\n", " ")
                chn = t.get("chn", "").replace("\n", " ")
                reason = t.get("violation_reason", "").strip()
                src = t.get("source_file", "")
                f.write("--------------------------------------------------\n")
                f.write(f"[파일명]: {src}\n")
                f.write(f"[한국어 원문]: {kor}\n")
                if reason:
                    f.write(f"[광고법/규제 수정 사유]: {reason}\n")
                f.write(f"[중국어 번역]: {chn}\n")
            f.write("--------------------------------------------------\n")
        print(f"  -> [SUCCESS] 비교표 저장 완료: {report_path}")

    print(f"\n[FINISH] CN_Text-In_Image_Translation_Engine_V1 모든 작업 완료!")


async def generate_seo_geo_aeo_txt(client: genai.Client, current_source_dir: str, target_dir: str, target_lang: str, product_name: str):
    """4-Core 마이크로-써머리 SEO/GEO/AEO (TXT, HTML 뷰어, DOCX, MD) 4종 파일을 자동 생성합니다. (URL/고시표 듀얼 인제스천 및 4개국 법무 렉시콘 100% 결합)"""
    print(f"\n🌐 [SEO/GEO/AEO 4-Core] 정밀 팩트 인제스천 및 4종 포맷(DOCX/HTML/TXT/MD) 생성 중 ({target_lang})...", flush=True)

    # 1. 듀얼 인제스천: url.txt 실시간 웹 스크래핑
    url_fact_context = ""
    url_file_candidates = [
        os.path.join(current_source_dir, "url.txt"),
        os.path.join(current_source_dir, "product_url.txt"),
        os.path.join(current_source_dir, "URL.txt")
    ]
    for u_path in url_file_candidates:
        if os.path.exists(u_path):
            try:
                with open(u_path, "r", encoding="utf-8") as uf:
                    raw_url = uf.read().strip()
                if raw_url.startswith("http"):
                    print(f"  🔗 [URL INGESTION 감지] {raw_url}", flush=True)
                    import urllib.request
                    req = urllib.request.Request(raw_url, headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"})
                    with urllib.request.urlopen(req, timeout=5) as resp:
                        html_bytes = resp.read()
                        raw_html = html_bytes.decode("utf-8", errors="ignore")
                        clean_text = re.sub(r'<script.*?</script>', '', raw_html, flags=re.DOTALL)
                        clean_text = re.sub(r'<style.*?</style>', '', clean_text, flags=re.DOTALL)
                        clean_text = re.sub(r'<[^>]+>', ' ', clean_text)
                        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                        url_fact_context = f"\n[Live Web Product Page Facts (Ground Truth)]\nURL: {raw_url}\nScraped Text Extract: {clean_text[:1200]}\n"
                        print(f"  ✅ [URL 팩트 스크래핑 완료] {len(clean_text)} 자 추출", flush=True)
                        break
            except Exception as ue:
                print(f"  ⚠️ [URL 스크래핑 실패 / 폴백 진행]: {ue}", flush=True)

    # 2. 이미지 팩트 앵커링 & 초월번역 컨텍스트 로드
    guide_path = os.path.join(current_source_dir, "transcreation_guide.json")
    qa_path = os.path.join(target_dir, "Transcreation_QA_Report.json")
    guide_context = ""
    if os.path.exists(guide_path):
        try:
            with open(guide_path, "r", encoding="utf-8") as f:
                gdata = json.load(f)
                guide_context = json.dumps(gdata.get("transcreation_comparisons", [])[:6], ensure_ascii=False)
        except Exception:
            pass
    elif os.path.exists(qa_path):
        try:
            with open(qa_path, "r", encoding="utf-8") as f:
                qdata = json.load(f)
                guide_context = json.dumps(qdata.get("transcreation_comparisons", [])[:6], ensure_ascii=False)
        except Exception:
            pass

    # 3. 4개국 법무 렉시콘 로드 (COMPLIANCE-FIRST)
    lexicon_rules_text = ""
    lexicon_map = {
        "EN": "en_fda_mocra_lexicon.json",
        "JP": "jp_pmda_pharm_lexicon.json",
        "CN": "cn_nmpa_adlaw_lexicon.json",
        "TW": "tw_tfda_lexicon.json"
    }
    lex_file = lexicon_map.get(target_lang)
    if lex_file:
        lex_path = os.path.join(PROJECT_ROOT, "00_공통자료", "compliance_lexicons", lex_file)
        if os.path.exists(lex_path):
            try:
                with open(lex_path, "r", encoding="utf-8") as lf:
                    lex_data = json.load(lf)
                    lexicon_rules_text = f"\n[MANDATORY COMPLIANCE LEXICON ({lex_data.get('jurisdiction', '')})]\n"
                    cats = lex_data.get("categories", {})
                    for c_name, c_val in cats.items():
                        banned = c_val.get("banned_terms", [])
                        for b in banned[:5]:
                            lexicon_rules_text += f"- Banned: '{b.get('banned')}' -> Must use: '{b.get('preferred')}' ({b.get('reason')})\n"
            except Exception:
                pass

    lang_names = {
        "EN": "English for Amazon / Sephora US",
        "JP": "Japanese for Qoo10 Japan / Cosme",
        "CN": "Simplified Chinese for Tmall / Xiaohongshu",
        "TW": "Traditional Chinese for Shopee Taiwan / Momo",
        "KR": "Korean for Naver Smartstore / Coupang"
    }
    target_lang_desc = lang_names.get(target_lang, "English")

    prompt = f"""[SYSTEM PROMPT] Global E-Commerce SEO/GEO/AEO 4-Core Master Copy Generator
Product Name: {product_name}
Target Market & Language: {target_lang_desc}

{url_fact_context}
[Verified Transcreation Context & Ingredients Data]:
{guide_context}

{lexicon_rules_text}

[CRITICAL INSTRUCTION - ZERO META COMMENTARY]
Never output developer metadata, markdown headers '##', explanation notes, character counters, or words like 'GEO', 'AEO', 'RAG'.
Output purely customer-facing content structured in 4 distinct sections.

Strict 4-Core Structure:
1. Official Product Title
(Under 100 characters. Noun Phrase: Brand 'Logicall Skin' + Product Title + Key Efficacy + Volume)

2. Core Value & Active Ingredient Summary
(5 concise bullet points containing quantitative metrics e.g. ppm, %, non-irritation score 0.00):
- Brand: Logicall Skin
- Core Actives & Concentration: (e.g. Multi-Vitamin 10% / 100,000ppm, Aquatide 3%)
- Key Benefits: (Efficacy claims strictly compliant with target country cosmetics law)
- Texture & Absorption: (Hydra-watery, non-greasy, fast-absorbing)
- Skin Compatibility: (Dermatologist tested, 0.00 irritation index)

3. Product Specifications & Comparison Table
(Output a clean HTML <table> comparing Logicall Skin vs Generic Market Standard):
<table>
  <tr><th>Dimensions</th><th>Logicall Skin</th><th>Standard Market Benchmark</th></tr>
  <tr><td>Active Concentration</td><td>High-Potency Multi-Vitamin 100,000ppm (10%)</td><td>Diluted extract 1,000~5,000ppm</td></tr>
  <tr><td>Patented Science</td><td>Aquatide 5000 (3%)</td><td>Generic purified water base</td></tr>
  <tr><td>Formula Stability</td><td>High-stability oxidation-free formula</td><td>Prone to discoloration / oxidation</td></tr>
  <tr><td>Irritation Index</td><td>0.00 Low-Irritation Certified</td><td>May cause stinging or redness</td></tr>
</table>

4. Product Usage Guide & Frequently Asked Questions (FAQ)
(5 high-conversion B2C customer FAQs):
Q1: When should I apply this serum?
A: ...
Q2: Is it suitable for sensitive skin?
A: ...
Q3: How does the active formula benefit the skin?
A: ...
Q4: What is the main efficacy of the Multi-Vitamin complex?
A: ...
Q5: Can I layer this with other skincare products?
A: ...

Generate the complete 4-Core content in {target_lang_desc} now.
"""
    try:
        resp = await client.aio.models.generate_content(
            model=MODEL_PRO,
            contents=[prompt],
            config=types.GenerateContentConfig(
                temperature=0.6,
                top_p=0.9,
                max_output_tokens=8192
            )
        )
        content_text = resp.text.strip()
    except Exception as e:
        print(f"  ⚠️ [WARN] SEO 텍스트 생성 중 오류 발생 -> 기본 템플릿 대체: {e}")
        content_text = f"""1. Logicall Skin {product_name} Multi Vitamin Daily Care Serum 50ml

2. Core Value & Active Ingredient Summary
- Brand: Logicall Skin
- Core Actives & Concentration: High-Potency Multi-Vitamin 10% (100,000ppm) & Aquatide 5000 (3%)
- Key Benefits: Visibly brightens, refines texture, and reinforces natural moisture barrier
- Texture & Absorption: Refreshingly lightweight hydra-watery formula with instant absorption
- Skin Compatibility: Dermatologist-tested, 0.00 skin irritation index suitable for daily use

3. Product Specifications & Comparison Table
<table>
  <tr><th>Specification</th><th>Logicall Skin Multi-Vitamin Serum</th><th>Standard Vitamin Serum</th></tr>
  <tr><td>Active Concentration</td><td>Multi-Vitamin Complex 100,000ppm (10.0%)</td><td>1,000 ~ 5,000ppm</td></tr>
  <tr><td>Patented Technology</td><td>Aquatide 5000 30,000ppm (3.0%)</td><td>Purified water base</td></tr>
  <tr><td>Formula Stability</td><td>High stability against air and light oxidation</td><td>Vulnerable to discoloration</td></tr>
  <tr><td>Skin Irritation Score</td><td>0.00 (Certified Low-Irritation)</td><td>May cause stinging sensation</td></tr>
</table>

4. Product Usage Guide & Frequently Asked Questions (FAQ)
Q1: When should I apply this serum?
A: Apply 3-4 drops evenly morning and evening after cleansing and toner.
Q2: Is it suitable for sensitive skin?
A: Yes, it is dermatologist-tested with a 0.00 skin irritation index.
Q3: How does Aquatide benefit the skin?
A: It reinforces the natural moisture barrier and revitalizes skin appearance.
Q4: What is the main efficacy of the Multi-Vitamin complex?
A: It provides deep hydration for a resilient-looking complexion and combats the signs of premature aging.
Q5: Can I layer this with other skincare products?
A: Yes, its fast-absorbing texture layers smoothly under creams and sunscreens.
"""

    # 4. 결정론적 법무 후처리 게이트 통과 (apply_deterministic_qa_overrides)
    if target_lang == "EN":
        for b_pat, p_val in [
            (r"\bnutrients for cellular vitality\b", "hydration for a resilient-looking complexion"),
            (r"\bcellular vitality\b", "resilient-looking complexion"),
            (r"\breinforces cellular resilience\b", "reinforces the skin's natural moisture barrier"),
            (r"\bcellular resilience\b", "skin's natural moisture barrier"),
            (r"\bcellular metabolism\b", "natural skin vitality"),
            (r"\bcombats premature aging\b", "combats the signs of premature aging"),
            (r"\bComplex skin issues\b", "Multiple skin concerns"),
            (r"\bTroubled skin\b", "Blemish-prone skin"),
            (r"\bcellular autophagy\b", "targeted skin nourishment")
        ]:
            content_text = re.sub(b_pat, p_val, content_text, flags=re.IGNORECASE)

    # 5. 4종 멀티 포맷 일괄 익스포트 (TXT, HTML, DOCX, MD)
    txt_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO.txt"
    html_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO_VIEWER.html"
    docx_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO.docx"
    md_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO.md"

    txt_path = os.path.join(target_dir, txt_filename)
    html_path = os.path.join(target_dir, html_filename)
    docx_path = os.path.join(target_dir, docx_filename)
    md_path = os.path.join(target_dir, md_filename)

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(content_text)
    print(f"  📄 [TXT 저장 완료]: {txt_path}")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(f"# {product_name} ({target_lang}) SEO/GEO/AEO PDP 원고\n\n" + content_text)
    print(f"  📑 [MD 저장 완료]: {md_path}")

    _generate_web_copier_html_file(f"{product_name} ({target_lang})", content_text, html_path, target_lang=target_lang)
    print(f"  🌐 [HTML 뷰어 저장 완료]: {html_path}")

    _generate_docx_file(f"{product_name} ({target_lang})", content_text, docx_path, target_lang=target_lang)
    print(f"  📄 [DOCX 서식 문서 저장 완료]: {docx_path}")


def _generate_docx_file(title: str, text_content: str, out_docx_path: str, target_lang: str = "EN"):
    """MS Word 서식(.docx)으로 4-Core 상세페이지 완성 원고를 렌더링합니다."""
    try:
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        doc = docx.Document()
        for s in doc.sections:
            s.top_margin = Inches(0.8)
            s.bottom_margin = Inches(0.8)
            s.left_margin = Inches(0.8)
            s.right_margin = Inches(0.8)

        # Title Header
        p_title = doc.add_paragraph()
        r_t = p_title.add_run(f"🛒 {title} - E-Commerce PDP Master Copy")
        r_t.font.name = "맑은 고딕"
        r_t.font.size = Pt(16)
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(16, 44, 87)

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(14)
        r_sub = p_sub.add_run("이 문서는 각국 전자상거래 플랫폼(스마트스토어, 쿠팡, 아마존, 큐텐 등) 상세페이지 에디터에 그대로 복사하여 등록할 수 있는 100% 고객 노출용 원고입니다.")
        r_sub.font.name = "맑은 고딕"
        r_sub.font.size = Pt(9.5)
        r_sub.font.color.rgb = RGBColor(100, 116, 139)

        # Parse sections
        lines = [l.strip() for l in text_content.splitlines() if l.strip()]
        cur_sec = 0
        sec_buffers = {1: [], 2: [], 3: [], 4: []}

        for l in lines:
            if l.startswith("1."):
                cur_sec = 1
                sec_buffers[1].append(l)
            elif l.startswith("2."):
                cur_sec = 2
                sec_buffers[2].append(l)
            elif l.startswith("3."):
                cur_sec = 3
                sec_buffers[3].append(l)
            elif l.startswith("4."):
                cur_sec = 4
                sec_buffers[4].append(l)
            else:
                if cur_sec in sec_buffers:
                    sec_buffers[cur_sec].append(l)

        # Section 1: Title
        h1 = doc.add_heading(level=1)
        r_h1 = h1.add_run("1. 공식 상품명 (Official Product Title)")
        r_h1.font.name = "맑은 고딕"
        r_h1.font.size = Pt(12.5)
        r_h1.font.bold = True
        r_h1.font.color.rgb = RGBColor(37, 99, 235)

        t_val = " ".join([l for l in sec_buffers[1] if not l.startswith("1.")]).strip()
        if not t_val and sec_buffers[1]:
            t_val = sec_buffers[1][-1].replace("1.", "").strip()
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_after = Pt(10)
        r_tv = p_t.add_run(t_val)
        r_tv.font.name = "맑은 고딕"
        r_tv.font.size = Pt(10.5)
        r_tv.font.bold = True

        # Section 2: Summary
        h2 = doc.add_heading(level=1)
        r_h2 = h2.add_run("2. 핵심 가치 및 제품 안내 (Core Value & Summary)")
        r_h2.font.name = "맑은 고딕"
        r_h2.font.size = Pt(12.5)
        r_h2.font.bold = True
        r_h2.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[2]:
            if l.startswith("2."):
                continue
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(l.lstrip("- •*"))
            r.font.name = "맑은 고딕"
            r.font.size = Pt(9.5)

        # Section 3: Comparison Table
        h3 = doc.add_heading(level=1)
        r_h3 = h3.add_run("3. 제품 상세 스펙 비교 (Product Specifications & Comparison Table)")
        r_h3.font.name = "맑은 고딕"
        r_h3.font.size = Pt(12.5)
        r_h3.font.bold = True
        r_h3.font.color.rgb = RGBColor(37, 99, 235)

        raw_s3_text = "\n".join(sec_buffers[3])
        if "<table>" in raw_s3_text:
            rows_data = []
            tr_matches = re.findall(r'<tr>(.*?)</tr>', raw_s3_text, flags=re.DOTALL)
            for tr in tr_matches:
                cols = re.findall(r'<t[hd]>(.*?)</t[hd]>', tr, flags=re.DOTALL)
                if cols:
                    rows_data.append([re.sub(r'<[^>]+>', '', c).strip() for c in cols])
            
            if rows_data:
                col_cnt = max(len(r) for r in rows_data)
                tbl = doc.add_table(rows=len(rows_data), cols=col_cnt)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r_idx, r_cols in enumerate(rows_data):
                    for c_idx, c_val in enumerate(r_cols):
                        cell = tbl.rows[r_idx].cells[c_idx]
                        p = cell.paragraphs[0]
                        r = p.add_run(c_val)
                        r.font.name = "맑은 고딕"
                        r.font.size = Pt(9)
                        if r_idx == 0:
                            r.font.bold = True
                            tcPr = cell._element.get_or_add_tcPr()
                            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
                            tcPr.append(shd)
        else:
            for l in sec_buffers[3]:
                if not l.startswith("3."):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    r = p.add_run(l)
                    r.font.name = "맑은 고딕"
                    r.font.size = Pt(9.5)

        # Section 4: FAQ
        h4 = doc.add_heading(level=1)
        r_h4 = h4.add_run("4. 자주 묻는 질문 (Frequently Asked Questions)")
        r_h4.font.name = "맑은 고딕"
        r_h4.font.size = Pt(12.5)
        r_h4.font.bold = True
        r_h4.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[4]:
            if l.startswith("4."):
                continue
            if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = RGBColor(30, 58, 138)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(51, 65, 85)

        doc.save(out_docx_path)
    except Exception as de:
        print(f"  ⚠️ [DOCX 생성 폴백]: {de}", flush=True)


def _generate_web_copier_html_file(title: str, text_content: str, out_html_path: str, target_lang: str = "EN"):
    """지마켓/스마트스토어/쿠팡/해외몰 등 전 이커머스 플랫폼에 맞춘 4-Core '에디터에 붙여넣을 텍스트 복사' 및 'HTML로 텍스트 복사' 뷰어를 생성합니다."""
    
    section_titles = {
        "EN": {
            "sec2": "2. Core Value & Active Ingredient Summary",
            "sec3": "3. Product Specifications & Comparison Table",
            "sec4": "4. Product Usage Guide & Frequently Asked Questions (FAQ)"
        },
        "JP": {
            "sec2": "2. コアバリュー＆成分サイエンス要約",
            "sec3": "3. 製品仕様・他社比較テーブル",
            "sec4": "4. 使用ガイド＆よくある質問 (FAQ)"
        },
        "CN": {
            "sec2": "2. 核心价值与成分科技摘要",
            "sec3": "3. 产品规格与竞品对比表",
            "sec4": "4. 商品使用指南与常见问题解答 (FAQ)"
        },
        "TW": {
            "sec2": "2. 核心價值與成分科技摘要",
            "sec3": "3. 產品規格與競品對比表",
            "sec4": "4. 使用指南與常見問題解答 (FAQ)"
        },
        "KR": {
            "sec2": "2. 제품 핵심 안내 및 성분 요약",
            "sec3": "3. 제품 상세 비교 스펙 테이블",
            "sec4": "4. 자주 묻는 질문 (FAQ)"
        }
    }
    
    detected_lang = target_lang.upper() if target_lang else "EN"
    titles_map = section_titles.get(detected_lang, section_titles["EN"])
    sec2_heading = titles_map["sec2"]
    sec3_heading = titles_map["sec3"]
    sec4_heading = titles_map.get("sec4", "4. FAQ")

    lines = [l.strip() for l in text_content.splitlines() if l.strip()]
    
    s1_lines = []
    s2_lines = []
    s3_lines = []
    s4_lines = []
    cur_sec = 0
    for l in lines:
        if l.startswith("1."):
            cur_sec = 1
            s1_lines.append(l)
        elif l.startswith("2."):
            cur_sec = 2
            s2_lines.append(l)
        elif l.startswith("3."):
            cur_sec = 3
            s3_lines.append(l)
        elif l.startswith("4."):
            cur_sec = 4
            s4_lines.append(l)
        else:
            if cur_sec == 1:
                s1_lines.append(l)
            elif cur_sec == 2:
                s2_lines.append(l)
            elif cur_sec == 3:
                s3_lines.append(l)
            elif cur_sec == 4:
                s4_lines.append(l)

    # 1. Section 1 (Title)
    s1_clean = " ".join([l for l in s1_lines if not l.startswith("1.")]).strip()
    if not s1_clean and s1_lines:
        s1_clean = s1_lines[-1].replace("1.", "").strip()

    # 2. Section 2 (Summary)
    s2_clean_items = [l for l in s2_lines if not l.startswith("2.")]
    s2_text_formatted = f"{sec2_heading}\n\n" + "\n\n".join(s2_clean_items)
    
    s2_html_items = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s2_html_items.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>🔬 {sec2_heading}</h3>")
    for l in s2_clean_items:
        if ":" in l or "：" in l:
            k, v = re.split(r'[:：]', l, 1)
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'><strong style='color:#1e3a8a;'>{k.strip()}:</strong> {v.strip()}</p>")
        else:
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'>{l.strip()}</p>")
    s2_html_items.append("</div>")
    s2_html = "\n".join(s2_html_items)

    # 3. Section 3 (Table)
    s3_clean_items = [l for l in s3_lines if not l.startswith("3.")]
    s3_raw_block = "\n".join(s3_clean_items)
    s3_text_formatted = f"{sec3_heading}\n\n" + s3_raw_block
    
    if "<table>" in s3_raw_block:
        styled_table = s3_raw_block.replace("<table>", "<table style='width:100%; border-collapse:collapse; margin:10px 0; font-size:13.5px;'>")
        styled_table = styled_table.replace("<th>", "<th style='background:#f1f5f9; padding:10px 12px; border:1px solid #cbd5e1; color:#1e3a8a; text-align:left; font-weight:bold;'>")
        styled_table = styled_table.replace("<td>", "<td style='padding:9px 12px; border:1px solid #cbd5e1; color:#334155;'>")
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n{styled_table}\n</div>"
    else:
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n<pre style='white-space:pre-wrap; font-family:inherit; font-size:13.5px;'>{s3_raw_block}</pre>\n</div>"

    # 4. Section 4 (FAQ)
    s4_clean_items = [l for l in s4_lines if not l.startswith("4.")]
    s4_text_blocks = []
    s4_html_blocks = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s4_html_blocks.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 12px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>💬 {sec4_heading}</h3>")
    
    cur_q = ""
    cur_answers = []
    for l in s4_clean_items:
        if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
            if cur_q:
                q_text = cur_q
                a_text = "\n".join(cur_answers)
                s4_text_blocks.append(f"{q_text}\n{a_text}")
                s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
                for ans in cur_answers:
                    s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
            cur_q = l.strip()
            cur_answers = []
        else:
            sentences = [s.strip() for s in re.split(r'(?<=[。！？\.\?!])\s*', l) if s.strip()]
            for s in sentences:
                cur_answers.append(s)
                
    if cur_q:
        q_text = cur_q
        a_text = "\n".join(cur_answers)
        s4_text_blocks.append(f"{q_text}\n{a_text}")
        s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
        for ans in cur_answers:
            s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
    s4_html_blocks.append("</div>")

    s4_text_formatted = f"{sec4_heading}\n\n" + "\n\n".join(s4_text_blocks)
    s4_html = "\n".join(s4_html_blocks)

    full_html_code = f"""<!-- 다국어 E-Commerce 상세페이지 4-Core 마이크로-써머리 & 비교표 & FAQ -->
<div style="font-family:'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; max-width:860px; margin:0 auto; padding:10px 0; color:#1e293b;">
{s2_html}
{s3_html}
{s4_html}
</div>"""

    html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; background: #f1f5f9; color: #0f172a; padding: 25px; margin: 0; line-height: 1.6; }}
  .container {{ max-width: 960px; margin: 0 auto; background: #ffffff; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.06); padding: 35px; border: 1px solid #cbd5e1; }}
  h1 {{ font-size: 22px; color: #0f172a; border-bottom: 2px solid #2563eb; padding-bottom: 12px; margin-top: 0; display: flex; align-items: center; gap: 8px; }}
  .guide-box {{ background: #eff6ff; border: 1px solid #bfdbfe; border-left: 5px solid #2563eb; padding: 18px 20px; border-radius: 8px; margin-bottom: 25px; font-size: 14px; color: #1e40af; line-height: 1.8; }}
  .guide-box strong {{ color: #1e3a8a; }}
  .card {{ background: #f8fafc; border: 1px solid #cbd5e1; border-radius: 10px; padding: 22px; margin-bottom: 25px; }}
  .card-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 14px; border-bottom: 1px solid #e2e8f0; padding-bottom: 10px; }}
  .card-title {{ font-size: 16px; font-weight: bold; color: #1e3a8a; }}
  .btn-group {{ display: flex; gap: 8px; }}
  .copy-btn {{ background: #2563eb; color: #ffffff; border: none; padding: 9px 16px; border-radius: 6px; cursor: pointer; font-size: 13px; font-weight: 700; transition: all 0.2s; box-shadow: 0 2px 4px rgba(37,99,235,0.2); display: inline-flex; align-items: center; gap: 6px; }}
  .copy-btn:hover {{ background: #1d4ed8; transform: translateY(-1px); }}
  .copy-btn.html-mode-btn {{ background: #059669; box-shadow: 0 2px 4px rgba(5,150,105,0.25); }}
  .copy-btn.html-mode-btn:hover {{ background: #047857; }}
  .full-btn {{ background: #7c3aed; padding: 12px 24px; font-size: 15px; width: 100%; justify-content: center; margin-bottom: 20px; box-shadow: 0 3px 6px rgba(124,58,237,0.25); }}
  .full-btn:hover {{ background: #6d28d9; }}
  .text-area {{ width: 100%; box-sizing: border-box; background: #ffffff; border: 1px solid #cbd5e1; border-radius: 6px; padding: 14px 16px; font-size: 13.5px; line-height: 1.75; color: #1e293b; font-family: 'Malgun Gothic', 'Segoe UI', monospace; resize: vertical; outline: none; }}
  .text-area:focus {{ border-color: #2563eb; box-shadow: 0 0 0 3px rgba(37,99,235,0.15); }}
  .toast {{ position: fixed; bottom: 30px; right: 30px; background: #0f172a; color: #ffffff; padding: 14px 28px; border-radius: 8px; font-size: 14px; font-weight: 600; display: none; z-index: 1000; box-shadow: 0 6px 20px rgba(0,0,0,0.25); }}
</style>
</head>
<body>

<div class="container">
  <h1>🌐 {title}</h1>
  
  <div class="guide-box">
    📢 <strong>쇼핑몰 등록 방식별 2대 원클릭 복사 기능 안내:</strong><br>
    • <strong>1. [📋 에디터에 붙여넣을 텍스트 복사] (파란색 버튼)</strong>: 지마켓/스마트스토어/쿠팡/해외몰 <strong>'에디터 작성'</strong> 화면에 붙여넣을 때 사용합니다.<br>
    • <strong>2. [🌐 HTML로 텍스트 복사] (초록색 버튼)</strong>: <strong>'HTML 작성'</strong> 탭이나 HTML 직접 입력 모드에 붙여넣을 때 사용합니다.
  </div>

  <button class="copy-btn full-btn html-mode-btn" onclick="copyFromTextarea('full-html-ta', '🎉 전체 HTML 소스코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🚀 [HTML로 전체 일괄 복사] 4-Core 요약 + 비교표 + FAQ 전체 소스코드 복사</button>
  <textarea id="full-html-ta" style="display:none;">{full_html_code}</textarea>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📌 1. 공식 상품명 (Title)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec1-ta', '✅ 상품명이 복사되었습니다!')">📋 상품명 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec1-ta" rows="2" readonly>{s1_clean}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">🔬 2. 핵심 가치 및 5줄 마이크로 요약</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec2-ta', '✅ 5줄 요약 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec2-html-ta', '🌐 5줄 요약 HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec2-ta" rows="8" readonly>{s2_text_formatted}</textarea>
    <textarea id="sec2-html-ta" style="display:none;">{s2_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📊 3. 제품 상세 스펙 비교표 (HTML Table)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec3-ta', '✅ 비교표 텍스트가 복사되었습니다!')">📋 비교표 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec3-html-ta', '🌐 비교표 HTML 코드가 복사되었습니다!')">🌐 HTML로 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec3-ta" rows="10" readonly>{s3_text_formatted}</textarea>
    <textarea id="sec3-html-ta" style="display:none;">{s3_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">💬 4. 5대 핵심 FAQ & 상세 가이드</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec4-ta', '✅ FAQ 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec4-html-ta', '🌐 FAQ HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec4-ta" rows="18" readonly>{s4_text_formatted}</textarea>
    <textarea id="sec4-html-ta" style="display:none;">{s4_html}</textarea>
  </div>

</div>

<div class="toast" id="toast">✅ 클립보드에 복사되었습니다! 쇼핑몰 에디터에 붙여넣기(Ctrl+V) 하세요.</div>

<script>
function showToast(msg) {{
  const t = document.getElementById('toast');
  t.innerText = msg;
  t.style.display = 'block';
  setTimeout(() => {{ t.style.display = 'none'; }}, 3000);
}}

function copyFromTextarea(id, customMsg) {{
  const ta = document.getElementById(id);
  const text = ta.value;
  
  const temp = document.createElement('textarea');
  temp.value = text;
  temp.style.position = 'fixed';
  temp.style.left = '-9999px';
  document.body.appendChild(temp);
  temp.select();
  temp.setSelectionRange(0, 99999);
  
  try {{
    document.execCommand('copy');
    showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
  }} catch (err) {{
    if (navigator.clipboard) {{
      navigator.clipboard.writeText(text).then(() => {{
        showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
      }});
    }}
  }} finally {{
    document.body.removeChild(temp);
  }}
}}
</script>

</body>
</html>
"""
    with open(out_html_path, "w", encoding="utf-8") as f:
        f.write(html_content)


if __name__ == '__main__':
    asyncio.run(main_async())
