import os
import sys
import json
import re
import time
import io
import argparse
import asyncio
from typing import Any, Dict, List, Optional, Tuple
from google import genai
from google.genai import types
from PIL import Image

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, ".."))
script_dir = SCRIPT_DIR
project_root = PROJECT_ROOT
CORE_DIR = os.path.join(PROJECT_ROOT, "multilingual_text_in_image_translatio_agy_sdk_core")
if os.path.exists(CORE_DIR) and CORE_DIR not in sys.path:
    sys.path.insert(0, CORE_DIR)

def get_recursive_files(base_dir):
    try:
        res = []
        for root, _, files in os.walk(base_dir):
            for f in files:
                res.append(os.path.relpath(os.path.join(root, f), base_dir))
        return res
    except Exception:
        return []


def load_credentials() -> genai.Client:
    env_paths = [
        os.path.join(PROJECT_ROOT, ".env"),
        os.path.join(PROJECT_ROOT, "영어", ".env"),
        os.path.join(PROJECT_ROOT, "일본어", ".env"),
        os.path.join(PROJECT_ROOT, "중국어", ".env"),
        os.path.join(SCRIPT_DIR, ".env"),
    ]
    api_key = os.environ.get("GEMINI_API_KEY")
    gcp_json_key = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")

    for p in env_paths:
        if os.path.exists(p):
            with open(p, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line.startswith("GEMINI_API_KEY="):
                        api_key = line.split("=", 1)[1].strip().strip('"').strip("'")
                    elif line.startswith("GOOGLE_APPLICATION_CREDENTIALS="):
                        gcp_json_key = line.split("=", 1)[1].strip().strip('"').strip("'")

    key_candidates = [
        gcp_json_key,
        os.path.join(PROJECT_ROOT, "00_공통자료", "APIs_KEY", "인증키_및_계정", "김차장_vertex api_key", "vertex_ai_auth_key.json"),
        os.path.join(PROJECT_ROOT, "00_공통자료", "인증키_및_계정", "김차장_vertex api_key", "vertex_ai_auth_key.json"),
    ]
    for kpath in key_candidates:
        if kpath and os.path.exists(kpath) and kpath.endswith(".json"):
            gcp_json_key = kpath
            break

    if gcp_json_key and os.path.exists(gcp_json_key):
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = gcp_json_key
        try:
            with open(gcp_json_key, "r", encoding="utf-8") as f:
                project_id = json.load(f).get("project_id")
            client = genai.Client(vertexai=True, project=project_id, location="global")
            print(f"[AUTH SUCCESS] Vertex AI Client 연결 완료 (Project: {project_id}, Location: global)", flush=True)
            return client
        except Exception as e:
            print(f"[AUTH WARN] Vertex AI 키 로드 실패, API 키로 대체 시도: {e}", flush=True)

    if api_key:
        if api_key.startswith("AQ."):
            client = genai.Client(vertexai=True, api_key=api_key)
        else:
            client = genai.Client(api_key=api_key)
        print("[AUTH SUCCESS] Gemini API Key Client 연결 완료", flush=True)
        return client

    raise RuntimeError("GEMINI_API_KEY 또는 GOOGLE_APPLICATION_CREDENTIALS가 설정되지 않았습니다.")

from multilingual_transcreation_qa_evaluator_agy_sdk import evaluate_transcreation, generate_html_report

from PIL import Image

sys.stdout.reconfigure(encoding='utf-8')

# 클라우드 인증은 load_credentials() 함수에서 동적으로 처리됩니다.

# 3.1 최신 모델 설정
MODEL_PRO = "gemini-3.1-pro-preview" # 텍스트 및 로직용 (최신 Pro 모델)
MODEL_FLASH_IMAGE = "gemini-3.1-flash-image"

GLOBAL_COMPLIANCE_SYSTEM_INSTRUCTION = """[SYSTEM INSTRUCTION: Global Cross-Border E-Commerce Compliance & Prestige Beauty Transcreation Expert (Japanese Mode)]
당신은 일본 후생노동성(MHLW) 약기법 및 @cosme 럭셔리 뷰티 가이드라인을 완벽히 준수하는 15년 차 글로벌 뷰티 법무 감사관이자 시슬리/SK-II급 수석 카피라이터입니다.

[엄격 실행 대원칙]
1. [약기법 56종 포지티브 리스트 엄격 준수]: 치료/재생/세포활성화/소염 등 의약품 오인 클레임을 100% 차단하고 '肌を整える', 'うるおいを与える', '肌荒れを防ぐ' 등 공인된 56종 허용 효능으로 순화하십시오.
2. [절대 표현 전면 금지]: '世界初', 'No.1', '最高', '究極' 등 검증 불가능한 절대 표현을 배제하고 프리미엄 케어 표현으로 격상하십시오.
3. [고시정보표 법정 조항]: 한국 식약처(MFDS) 심사필, 3대 주의사항, 공정위 분쟁기준, +82 고객상담번호를 표준화하십시오.
"""

def load_jp_compliance_lexicon() -> Dict[str, str]:
    fpath = os.path.join(PROJECT_ROOT, "00_공통자료", "compliance_lexicons", "jp_pmda_pharm_lexicon.json")
    replacements = {
        r"治療": "肌を整えるケア",
        r"再生": "すこやかに保つ",
        r"消炎": "肌荒れを防ぐ",
        r"無刺激": "低刺激処方",
        r"細胞活性化": "肌にハリとうるおいを与える",
        r"美白": "うるおいによる透明感",
        r"世界初": "先進テクノロジー",
        r"No\.1": "こだわり抜いた",
        r"最高": "優れた",
        r"究極": "高機能"
    }
    if os.path.exists(fpath):
        try:
            with open(fpath, "r", encoding="utf-8") as f:
                data = json.load(f)
                for cat in data.get("categories", {}).values():
                    for it in cat.get("banned_terms", []):
                        b, p = it.get("banned", ""), it.get("preferred", "")
                        if b and p:
                            replacements[rf"{re.escape(b)}"] = p
        except Exception:
            pass
    return replacements
 # 이미지 인페인팅용 (특화)

# 커맨드라인 파라미터 파싱
if len(sys.argv) > 1:
    source_dir = sys.argv[1]
    if len(sys.argv) > 2:
        base_target_dir = sys.argv[2]
    else:
        base_target_dir = os.path.join(os.path.dirname(source_dir), os.path.normpath(source_dir).split(os.sep)[-1] + "_JP_Translated")
else:
    # 기본 경로 세팅 (개발/테스트용)
    source_dir = os.path.join(PROJECT_ROOT, "01_번역대상_원본")
    base_target_dir = os.path.join(PROJECT_ROOT, "02_번역결과_최종")

if len(sys.argv) > 2:
    target_dir = base_target_dir
else:
    if len(sys.argv) == 1:
        target_dir = base_target_dir
    else:
        folder_name = os.path.normpath(source_dir).split(os.sep)[-1]
        target_dir = os.path.join(base_target_dir, folder_name)
os.makedirs(target_dir, exist_ok=True)

# [Pass 1] 약기법 강제 번역 매핑 프롬프트
# [일본 후생노동성 공인 56종 허용 표현 (Positive List) 동적 로드]
efficacy_json_path = os.path.join(script_dir, "cosmetics_efficacy_56.json")
efficacy_list_str = ""
if os.path.exists(efficacy_json_path):
    try:
        with open(efficacy_json_path, "r", encoding="utf-8") as ef:
            efficacy_data = json.load(ef)
            efficacy_list_str = "\n".join([f"{item['id']}. {item['claim_jp']} ({item['claim_ko']})" for item in efficacy_data])
            print(f"[INFO] 일본 후생노동성 공인 56종 허용 효능 규격 로드 완료 ({len(efficacy_data)}종)")
    except Exception as e:
        print(f"[WARNING] 56종 규격 JSON 로드 중 에러 발생: {e}")

if not efficacy_list_str:
    efficacy_list_str = "1. 肌を整える\n2. 肌荒れを防ぐ\n3. 皮膚にうるおいを与える 등 56종"

# [Pass 1] 럭셔리 뷰티 초월번역 및 약기법 강제 매핑 프롬프트
pass1_prompt = f"""
[SYSTEM PROMPT] Global Luxury Beauty Transcreation & Compliance Expert (PROTO Japanese Engine)

## 1. 시스템 역할 및 콘셉트 (Role & Context)
당신은 시슬리, SK-II, 데코르테 등 일본 하이엔드 프레스티지 뷰티 시장을 총괄하는 10년 차 수석 크리에이티브 디렉터이자 @cosme 전문 엘리트 카피라이터입니다.
단순 직역을 배제하고, 일본 소비자의 감성을 깊게 자극하는 정중하고 품격 있는 뷰티 카피(美肌, ハリ, 潤い)로 초월번역(Transcreation)하세요.

## 2. 초월번역 핵심 원칙 (Core Transcreation Principles)
1. [기계적 직역 및 부사 금지]
   - '確実に', '本当に', '絶対に' 등 딱딱한 부사 직역을 전면 금지하고, 피부 감촉과 효능을 섬세하게 묘사하는 프리미엄 어휘로 재창조하십시오.
2. [자연스러운 구문 결속 및 제형 감성 묘사]
   - "10% LiftDerm" 등 성분 비율이 문맥과 끊기지 않고 매끄러운 뷰티 서사로 이어지도록 구조를 재조정하십시오.
3. [4대 기능성 뷰티 전문 어휘 사전 채택]
   - 피부 속/기저층: 肌の奥・角質層のすみずみまで
   - 토탈 케어/멀티 코렉티브: 高機能トータルリペア / 多機能エイジングケア
   - 탄력 복원/강화: ハリ・弾力を呼び覚ます / 弾むようなハリ感
   - 눈가 잔주름/건조주름: 目元の小ジワ・乾燥ジワ
4. [독자 성분명 영문 보존]
   - 'LiftDerm', 'Lifting Logic for eye' 등은 억지로 가타카나로 뭉개지 않고 영문 고유 표기를 유지하여 임상적 신뢰도를 극대화하십시오.
5. [절대적/과대 표현 전면 금지 (Ban on Absolute Claims)]
   - '世界初', 'No.1', '最高', '究極' 등 검증 불가능한 절대 표현 사용을 금지하고, '目元のために開発された先進テクノロジー', '高機能トータルケア' 등 프리미엄 케어 표현으로 순화하십시오.

## 3. 후생노동성 약기법(약사법) 규제 준수 가이드
[일본 후생노동성 공인 56종 허용 효능 목록 (Positive List)]
{efficacy_list_str}

[약기법 필수 준수 지침]
1. [기본 원칙] 일본 화장품법은 포지티브 리스트(56가지 허용 표현) 방식입니다. '치료/효능'이 아닌 '세정/관리/느낌' 위주로 순화해야 합니다. 위 [일본 후생노동성 공인 56종 허용 효능 목록]에 등재된 일본어 표현만을 정확히 사용하여 번역하십시오.
2. '자극 없이', '무자극' -> '피부에 순하게(肌にやさしく)' 또는 '저자극 처방(低刺激処方)'.
3. '진정(鎮静)' -> 샴푸/두피 제품의 경우 '지하다 케어(地肌ケア, 두피 관리)' 또는 '청결하게 유지(清潔に保つ)'. 피부의 경우 '피부를 정돈하다(肌を整える)'.
4. '탈모 방지' -> '두피 환경을 정돈(頭皮環境を整える)'.
5. '모공 축소', '피지 조절' -> '세정으로 모공 노폐물 제거(洗浄により毛穴の汚れを落とす)'.
6. '흡수' 또는 '침투' -> 범위를 명시하여 '각질층까지 침退(角質層まで浸透)'로 번역하거나 주석 `*浸透は角質層まで` 추가.
7. '치료' -> 'ケア', '개선' -> '整える', '재생' -> 'いきいき'. '적당량' -> '適量'.
8. '미백', '화이트닝', '잡티 제거' -> '(메이크업 효과로) 화사하게 연출' 또는 '수분을 주어 투명감 있는 피부로 케어(うるおいを与え透明感のある肌へ)'.
9. '리프팅', '안티에이징', '주름 개선/방지', '처짐 방지', '젊어짐' 등은 절대 금지. 허용 대체 문구: '肌にはりを与える(탄력을 주다)', '肌にツヤを与える(윤기를 주다)', '肌を引き締め、ハリのある印象へ', 'ハリを感じる肌へ導きます', 'スッキリとした印象の肌へ', '引き締まった印象で、若々しい肌へ', '肌を守りながらふっくらハリ肌へ'.
10. '잔주름' 표현은 일본 임상 시험 데이터가 없으면 불법이므로 '乾燥による小ジワを目立たなくする(건조로 인한 잔주름을 눈에 띄지 않게 함)'도 원본에 명시된 임상 근거가 없다면 자의적으로 쓰지 말고, 9번의 '탄력' 관련 문구로 일괄 우회할 것.
11. '에이징 케어(エイジングケア)'라는 단어를 쓸 경우, 주석 `*エイジングケアとは、年齢に応じたお手入れのこと` 를 반드시 표기할 것.
12. '노벨상' -> 'オートファジー技術'. '강력한', '제일' -> '高い保湿感', '優れた保湿力'.
13. [상품 정보 고시(Notice Table) 번역 표준 규격 (약기법 및 MFDS 식약처 심사필 강제 매핑)]
    - 기능성화장품 심사 유무: 韓国化粧品法に基づき韓国食品医薬品安全処(MFDS)の機能性化粧品審査(または報告)済
    - 사용상 주의사항: 1) お肌に異常が生じていないかよく注意して使用... 2) 傷やはれもの... 3) 保管上の注意...
    - 품질보증기준: 本商品に異常がある場合、公正取引委員会告示（消費者紛争解決基準）に基づき補償いたします。
    - 고객상담실: +82-2-6743-3206
14. [중요] 제품 패키지/용기에 적힌 영문 텍스트(예: LOGICALLY SKIN, MULTI CORRECTIVE EYE CREAM 등)는 절대 번역하거나 매핑 딕셔너리에 포함시키지 마세요. 이는 이미지 렌더링 시 AI가 해당 영문을 훼손(오타 등)하는 것을 방지하기 위함입니다. 오직 '한국어'만 번역 대상으로 삼으세요.
15. '디톡스', '해독', '배출(排出)', '피로(疲労)', '지침(疲れ)' -> 의약품/건강기능식품 용어이므로 화장품 사용 절대 불가. '肌を整える', '肌荒れを防ぐ', '不要なものをすっきりと', 'すこやかに保つ' 등으로 대체할 것.
16. [엄격 주의] 이미지 내의 모든 한국어 텍스트는 단 하나도 빠짐없이 100% 추출하여 번역 매핑에 포함시켜야 합니다.
17. [안전망 규칙 - TABLE HTML TO PNG] 이미지 내에 정보 고시 표나 복잡한 표(테이블) 레이아웃이 포함된 경우, 본 V6 인페인팅 방식으로 렌더링하지 마십시오.

[Few-Shot 매핑 사례 (반드시 아래의 번역 톤을 따를 것)]
사례 1: "피부 진정 효과가 뛰어난 티트리 추출물" -> "肌を整えるティーツリーエキス" (진정 -> 정돈하다)
사례 2: "눈가 주름 개선 및 안티에이징 기능" -> "目元にハリを与え、若々しい印象へ" (주름 개선/안티에이징 -> 탄력을 주다, 젊어보이는 인상)
사례 3: "피부 속까지 깊숙이 흡수되어" -> "角質層まで浸透し" (속까지 흡수 -> 각질층까지 침투)
사례 4: "부작용 없이 안전한 무자극 화장품" -> "肌にやさしい低刺激処方の化粧品" (무자극/부작용 없이 -> 저자극/피부에 순한)

출력은 반드시 JSON 형식으로 아래 스키마를 엄격히 따르세요:
{{
  "translation_map": [
    {{
      "kor": "한국어 원문", 
      "rule_check_reasoning": "이 문구에 안티에이징, 진정, 재생 등의 불법 표현이 포함되어 있는가? (검열 사고 과정 작성)",
      "jpn": "약기법 준수 일본어 번역문",
      "violation_reason": "수정 사유 (56종 위반 내용 등, 수정한 경우에만 기재, 수정 안했으면 빈 문자열)"
    }}
  ],
  "required_footnotes": [
    "필요한 법적 주석 문자열 (없으면 빈 배열)"
  ]
}}
"""

# [Pass 2] 렌더링 지시 프롬프트 템플릿
pass2_prompt_template = """
당신은 정밀한 시각적 로컬라이제이션을 수행하는 이미지 인페인팅 AI입니다.
첨부된 원본 이미지 속의 텍스트 위치, 배경 텍스처, 디자인 레이아웃을 1픽셀의 왜곡 없이 그대로 유지하세요.
아래에 제공된 [번역 매핑 데이터 JSON]을 바탕으로 다음 규칙을 엄격히 적용하여 단일 이미지를 생성하세요.

[시각적 렌더링 엄격 규칙]
1. (KOR ERASING) 원본의 한국어 텍스트는 원래 자리에 남겨두지 말고 배경색으로 덮어써서 100% 지울 것. 병기(한글+일본어) 절대 금지.
2. (JSON APPLY) 지워진 그 자리에 오직 [번역 매핑 데이터 JSON]의 'jpn' 텍스트만 렌더링할 것. 모델 임의로 번역을 수정하지 말 것.
3. (FULL INPAINTING NO PATCHING) [중요] 텍스트 수정 지시가 있더라도 오류 부분만 오려내어 덧칠(Patching)하지 마세요. 반드시 첨부된 원본 이미지를 기반으로 캔버스 전체를 완전히 새롭게 렌더링(Full Inpainting)하여 1픽셀의 이질감도 없는 완벽한 하나의 이미지를 생성하십시오.
4. (VISUAL BALANCE) [중요] 전체를 새로 렌더링할 때, 원본의 상단 텍스트와 하단 텍스트 간의 '폰트 사이즈', '폰트 두께(Weight)', '자간'을 임의로 다르게 그리지 마십시오. 시각적으로 완벽하게 동일한 폰트 규격과 두께감으로 통일성 있게 식자해야 합니다.
5. (LAYOUT STRICTNESS) [중요] 빈 공간이 넓다고 해서 텍스트를 거대하게 채우지 마십시오. 주변 이미지(로고, 선 등)의 구도와 여백을 철저히 계산하여 원래 텍스트가 있던 단락의 정렬축(좌/우/중앙)을 1픽셀의 오차 없이 그대로 유지하십시오.
6. (PACKAGE PRESERVATION) 제품 본품(용기, 튜브 등) 표면에 인쇄된 영문 텍스트(예: 브랜드명 'LOGICALLY SKIN', 'MULTI CORRECTIVE EYE CREAM')는 절대 다시 그리거나(redraw) 인페인팅 하지 마세요. 원본 픽셀 구성을 단 1픽셀도 건드리지 말고 100% 완벽하게 보존해야 합니다. 오타(예: Logtcally 등)가 발생하면 즉시 실패로 간주됩니다.
7. (NO EXTRA NOISE) 번역과 무관한 AI 주석이나 영어 설명, 괄호를 이미지에 추가하지 말 것.
8. (FOOTNOTES) 'required_footnotes' 배열에 내용이 있다면, 이미지 최하단 또는 적절한 여백에 해당 주석 텍스트를 아주 작은 글씨로 삽입할 것.
9. (HTML TO PNG RULE) [중요] 테이블 표 속에 긴 문장과 텍스트가 빽빽하게 나열된 경우, AI 이미지 렌더링(인페인팅) 방식으로는 글씨가 뭉개지거나 격자가 파괴되는 치명적 한계가 발생하므로, 이를 코딩(HTML/CSS) 기반 렌더링으로 우회하여 완벽한 고화질 PNG를 얻어냅니다. 따라서 표 레이아웃이 감지되면 어떠한 텍스트 덮어쓰기 작업도 강행하지 마세요.

[번역 매핑 데이터 JSON]
{json_data}
"""



async def main_async():
    print("[START] PROTO_Text-In_Image_Translation_Engine_V7 (Two-Pass Architecture) 코어 가동...")
    print(f"[INFO] 타겟 스캔 폴더: {source_dir}")
    print(f"[INFO] 결과 저장 폴더: {target_dir}")

    targets = [f for f in get_recursive_files(source_dir) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.jfif', '.gif'))]

    if not targets:
        print(f"[WARNING] '{source_dir}' 폴더에 처리할 이미지가 없습니다.")
        sys.exit(0)

    all_translations = []

    for filename in targets:
        if 'JP' in filename:
            continue
        
        if '상세정보안내' in filename or '상품정보제공고시' in filename or '상품상세정보' in filename or filename.endswith('.txt'):
            continue
        
        in_path = os.path.join(source_dir, filename)
        out_name = f"{os.path.splitext(filename)[0]}_JP_Surgical_v7.png"
        out_path = os.path.join(target_dir, out_name)
    
        if os.path.exists(out_path):
            print(f"\n[SKIP] 이미 번역 완료된 파일입니다: {filename}")
            continue

        print(f"\n[RENDER] 변환 시작: {filename}")
    
        try:
            original_image = Image.open(in_path)
            original_image.load()
        except Exception as e:
            print(f"  -> [ERROR] 이미지 로드 실패: {e}")
            continue

        # ==========================
        # PASS 1: 텍스트 추출 및 번역 (pro 모델)
        # ==========================
        print("  -> [PASS 1] 텍스트 매핑 및 약기법 검열 중...")
        try:
            response_p1 = await client.aio.models.generate_content(
                model=MODEL_PRO,
                contents=[original_image, pass1_prompt],
                config=types.GenerateContentConfig(
                    system_instruction=GLOBAL_COMPLIANCE_SYSTEM_INSTRUCTION,
                    response_mime_type="application/json",
                    temperature=0.6,
                    top_p=0.9,
                    max_output_tokens=8192
                )
            )
            mapping_data_str = response_p1.text
            forbidden_patterns = {
                r"鎮静": "肌を整える",
                r"アンチエイジング": "年齢に応じたケア",
                r"副作用なし": "低刺激処方",
                r"無刺激": "低刺激処方",
                r"再生": "いきいき",
                r"シワ改善": "ハリを与える",
                r"ニキビ跡": "肌荒れを防ぐ",
                r"排出": "不要なものをすっきりと",
                r"デトックス": "肌を整える",
                r"解毒": "肌を整える",
                r"100%安全": "肌にやさしい",
                r"疲労": "すこやかに保つ",
                r"疲れ": "すこやかに保つ",
                r"最高": "優れた",
                r"一番": "優れた",
                r"強力な保湿": "優れた保湿力",
                r"強力な": "優れた",
                r"肌荒れ予防": "肌荒れを防ぐ",
                r"予防": "防ぐ",
                r"抗酸化作用": "肌を整える",
                r"抗酸化": "肌を整える"
            }
            # JSON 유효성 테스트
            parsed_json = json.loads(mapping_data_str)
            if "translation_map" in parsed_json:
                for item in parsed_json["translation_map"]:
                    jpn_text = item.get("jpn", "")
                    # Python 하드 필터링 (최후의 보루)
                    for pattern, safe_word in forbidden_patterns.items():
                        if re.search(pattern, jpn_text):
                            print(f"      [Python Regex Filter] 금지어 감지: '{pattern}' -> '{safe_word}' 로 강제 치환")
                            jpn_text = re.sub(pattern, safe_word, jpn_text)
                            item["violation_reason"] = item.get("violation_reason", "") + f" (Python 강제 필터링: {pattern} 적발)"
                    item["jpn"] = jpn_text
                    # 파일명도 출처로 함께 저장
                    item["source_file"] = filename
                all_translations.extend(parsed_json["translation_map"])
            print("  -> [PASS 1 SUCCESS] 매핑 데이터 생성 완료.")
        except Exception as e:
            print(f"  -> [PASS 1 ERROR] 약기법 매핑 실패: {e}")
            continue

        # ==========================
        # PASS 2: 이미지 렌더링 (flash-image 모델)
        # ==========================
        print("  -> [PASS 2] 이미지 인페인팅 렌더링 중...")
        try:
            final_prompt = pass2_prompt_template.replace("{json_data}", mapping_data_str)
            response_p2 = await client.aio.models.generate_content(
                model=MODEL_FLASH_IMAGE,
                contents=[final_prompt, original_image],
                config=types.GenerateContentConfig(
                    response_modalities=["IMAGE"],
                    temperature=0.6,
                    top_p=0.9
                )
            )
        
            img_saved = False
            if hasattr(response_p2, 'candidates'):
                for cand in response_p2.candidates:
                    if hasattr(cand, 'content') and hasattr(cand.content, 'parts'):
                        for part in cand.content.parts:
                            if hasattr(part, 'inline_data') and part.inline_data:
                                img = Image.open(io.BytesIO(part.inline_data.data))
                                img = img.resize(original_image.size, Image.Resampling.LANCZOS)
                                img.save(out_path, format="PNG")
                                img_saved = True
                                break
                            elif hasattr(part, 'image') and part.image:
                                img = Image.open(io.BytesIO(part.image.image_bytes))
                                img = img.resize(original_image.size, Image.Resampling.LANCZOS)
                                img.save(out_path, format="PNG")
                                img_saved = True
                                break
                            
            if img_saved:
                print(f"  -> [SUCCESS] {out_name} 최종 저장 완료!")
            else:
                print("  -> [FAILED] Pass 2에서 이미지 데이터를 반환받지 못했습니다.")
            
        except Exception as e:
            print(f"  -> [PASS 2 ERROR] 렌더링 실패: {e}")
    
        # 두 번 호출하므로 레이트 리밋 관리를 위해 8초 대기
        time.sleep(8)

    if all_translations:
        print("\n[REPORT] 약기법 56종 위반/대체 비교표 TXT 문서 생성 중...")
        report_path = os.path.join(target_dir, "약기법_번역_비교표_56종.txt")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write("==================================================\n")
            f.write("🇯🇵 V7 엔진: 약기법(56종) 위반 문구 vs 대체 허용 문구 비교표\n")
            f.write("==================================================\n")
            f.write("본 비교표는 V7 엔진이 일본 후생노동성 허용 56종(포지티브 리스트)을 강제 적용하여 변환한 데이터입니다.\n\n")
            for t in all_translations:
                reason = t.get("violation_reason", "").strip()
                kor = t.get("kor", "").replace("\n", " ")
                jpn = t.get("jpn", "").replace("\n", " ")
                src = t.get("source_file", "")
                f.write("--------------------------------------------------\n")
                f.write(f"[파일명]: {src}\n")
                f.write(f"[한국어 원본]: {kor}\n")
                f.write(f"[수정 사유]: {reason}\n")
                f.write(f"[대체 일본어]: {jpn}\n")
            f.write("--------------------------------------------------\n")
        print(f"  -> [SUCCESS] 비교표 저장 완료: {report_path}")

    print("\n[FINISH] PROTO_Text-In_Image_Translation_Engine_V7 이미지 번역 완료!")



    async def main_async():
        asyncio.run(main_async())


def _generate_docx_file(title: str, text_content: str, out_docx_path: str, target_lang: str = "EN"):
    """MS Word 서식(.docx)으로 4-Core 상세페이지 완성 원고를 렌더링합니다."""
    try:
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        doc = docx.Document()
        for s in doc.sections:
            s.top_margin = Inches(0.8)
            s.bottom_margin = Inches(0.8)
            s.left_margin = Inches(0.8)
            s.right_margin = Inches(0.8)

        # Title Header
        p_title = doc.add_paragraph()
        r_t = p_title.add_run(f"🛒 {title} - E-Commerce PDP Master Copy")
        r_t.font.name = "맑은 고딕"
        r_t.font.size = Pt(16)
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(16, 44, 87)

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(14)
        r_sub = p_sub.add_run("이 문서는 각국 전자상거래 플랫폼(스마트스토어, 쿠팡, 아마존, 큐텐 등) 상세페이지 에디터에 그대로 복사하여 등록할 수 있는 100% 고객 노출용 원고입니다.")
        r_sub.font.name = "맑은 고딕"
        r_sub.font.size = Pt(9.5)
        r_sub.font.color.rgb = RGBColor(100, 116, 139)

        # Parse sections
        lines = [l.strip() for l in text_content.splitlines() if l.strip()]
        cur_sec = 0
        sec_buffers = {1: [], 2: [], 3: [], 4: []}

        for l in lines:
            if l.startswith("1."):
                cur_sec = 1
                sec_buffers[1].append(l)
            elif l.startswith("2."):
                cur_sec = 2
                sec_buffers[2].append(l)
            elif l.startswith("3."):
                cur_sec = 3
                sec_buffers[3].append(l)
            elif l.startswith("4."):
                cur_sec = 4
                sec_buffers[4].append(l)
            else:
                if cur_sec in sec_buffers:
                    sec_buffers[cur_sec].append(l)

        # Section 1: Title
        h1 = doc.add_heading(level=1)
        r_h1 = h1.add_run("1. 공식 상품명 (Official Product Title)")
        r_h1.font.name = "맑은 고딕"
        r_h1.font.size = Pt(12.5)
        r_h1.font.bold = True
        r_h1.font.color.rgb = RGBColor(37, 99, 235)

        t_val = " ".join([l for l in sec_buffers[1] if not l.startswith("1.")]).strip()
        if not t_val and sec_buffers[1]:
            t_val = sec_buffers[1][-1].replace("1.", "").strip()
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_after = Pt(10)
        r_tv = p_t.add_run(t_val)
        r_tv.font.name = "맑은 고딕"
        r_tv.font.size = Pt(10.5)
        r_tv.font.bold = True

        # Section 2: Summary
        h2 = doc.add_heading(level=1)
        r_h2 = h2.add_run("2. 핵심 가치 및 제품 안내 (Core Value & Summary)")
        r_h2.font.name = "맑은 고딕"
        r_h2.font.size = Pt(12.5)
        r_h2.font.bold = True
        r_h2.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[2]:
            if l.startswith("2."):
                continue
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(l.lstrip("- •*"))
            r.font.name = "맑은 고딕"
            r.font.size = Pt(9.5)

        # Section 3: Comparison Table
        h3 = doc.add_heading(level=1)
        r_h3 = h3.add_run("3. 제품 상세 스펙 비교 (Product Specifications & Comparison Table)")
        r_h3.font.name = "맑은 고딕"
        r_h3.font.size = Pt(12.5)
        r_h3.font.bold = True
        r_h3.font.color.rgb = RGBColor(37, 99, 235)

        raw_s3_text = "\n".join(sec_buffers[3])
        if "<table>" in raw_s3_text:
            rows_data = []
            tr_matches = re.findall(r'<tr>(.*?)</tr>', raw_s3_text, flags=re.DOTALL)
            for tr in tr_matches:
                cols = re.findall(r'<t[hd]>(.*?)</t[hd]>', tr, flags=re.DOTALL)
                if cols:
                    rows_data.append([re.sub(r'<[^>]+>', '', c).strip() for c in cols])
            
            if rows_data:
                col_cnt = max(len(r) for r in rows_data)
                tbl = doc.add_table(rows=len(rows_data), cols=col_cnt)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r_idx, r_cols in enumerate(rows_data):
                    for c_idx, c_val in enumerate(r_cols):
                        cell = tbl.rows[r_idx].cells[c_idx]
                        p = cell.paragraphs[0]
                        r = p.add_run(c_val)
                        r.font.name = "맑은 고딕"
                        r.font.size = Pt(9)
                        if r_idx == 0:
                            r.font.bold = True
                            tcPr = cell._element.get_or_add_tcPr()
                            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
                            tcPr.append(shd)
        else:
            for l in sec_buffers[3]:
                if not l.startswith("3."):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    r = p.add_run(l)
                    r.font.name = "맑은 고딕"
                    r.font.size = Pt(9.5)

        # Section 4: FAQ
        h4 = doc.add_heading(level=1)
        r_h4 = h4.add_run("4. 자주 묻는 질문 (Frequently Asked Questions)")
        r_h4.font.name = "맑은 고딕"
        r_h4.font.size = Pt(12.5)
        r_h4.font.bold = True
        r_h4.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[4]:
            if l.startswith("4."):
                continue
            if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = RGBColor(30, 58, 138)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(51, 65, 85)

        doc.save(out_docx_path)
    except Exception as de:
        print(f"  ⚠️ [DOCX 생성 폴백]: {de}", flush=True)


def _generate_web_copier_html_file(title: str, text_content: str, out_html_path: str, target_lang: str = "EN"):
    """지마켓/스마트스토어/쿠팡/해외몰 등 전 이커머스 플랫폼에 맞춘 4-Core '에디터에 붙여넣을 텍스트 복사' 및 'HTML로 텍스트 복사' 뷰어를 생성합니다."""
    
    section_titles = {
        "EN": {
            "sec2": "2. Core Value & Active Ingredient Summary",
            "sec3": "3. Product Specifications & Comparison Table",
            "sec4": "4. Product Usage Guide & Frequently Asked Questions (FAQ)"
        },
        "JP": {
            "sec2": "2. コアバリュー＆成分サイエンス要約",
            "sec3": "3. 製品仕様・他社比較テーブル",
            "sec4": "4. 使用ガイド＆よくある質問 (FAQ)"
        },
        "CN": {
            "sec2": "2. 核心价值与成分科技摘要",
            "sec3": "3. 产品规格与竞品对比表",
            "sec4": "4. 商品使用指南与常见问题解答 (FAQ)"
        },
        "TW": {
            "sec2": "2. 核心價值與成分科技摘要",
            "sec3": "3. 產品規格與競品對比表",
            "sec4": "4. 使用指南與常見問題解答 (FAQ)"
        },
        "KR": {
            "sec2": "2. 제품 핵심 안내 및 성분 요약",
            "sec3": "3. 제품 상세 비교 스펙 테이블",
            "sec4": "4. 자주 묻는 질문 (FAQ)"
        }
    }
    
    detected_lang = target_lang.upper() if target_lang else "EN"
    titles_map = section_titles.get(detected_lang, section_titles["EN"])
    sec2_heading = titles_map["sec2"]
    sec3_heading = titles_map["sec3"]
    sec4_heading = titles_map.get("sec4", "4. FAQ")

    lines = [l.strip() for l in text_content.splitlines() if l.strip()]
    
    s1_lines = []
    s2_lines = []
    s3_lines = []
    s4_lines = []
    cur_sec = 0
    for l in lines:
        if l.startswith("1."):
            cur_sec = 1
            s1_lines.append(l)
        elif l.startswith("2."):
            cur_sec = 2
            s2_lines.append(l)
        elif l.startswith("3."):
            cur_sec = 3
            s3_lines.append(l)
        elif l.startswith("4."):
            cur_sec = 4
            s4_lines.append(l)
        else:
            if cur_sec == 1:
                s1_lines.append(l)
            elif cur_sec == 2:
                s2_lines.append(l)
            elif cur_sec == 3:
                s3_lines.append(l)
            elif cur_sec == 4:
                s4_lines.append(l)

    # 1. Section 1 (Title)
    s1_clean = " ".join([l for l in s1_lines if not l.startswith("1.")]).strip()
    if not s1_clean and s1_lines:
        s1_clean = s1_lines[-1].replace("1.", "").strip()

    # 2. Section 2 (Summary)
    s2_clean_items = [l for l in s2_lines if not l.startswith("2.")]
    s2_text_formatted = f"{sec2_heading}\n\n" + "\n\n".join(s2_clean_items)
    
    s2_html_items = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s2_html_items.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>🔬 {sec2_heading}</h3>")
    for l in s2_clean_items:
        if ":" in l or "：" in l:
            k, v = re.split(r'[:：]', l, 1)
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'><strong style='color:#1e3a8a;'>{k.strip()}:</strong> {v.strip()}</p>")
        else:
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'>{l.strip()}</p>")
    s2_html_items.append("</div>")
    s2_html = "\n".join(s2_html_items)

    # 3. Section 3 (Table)
    s3_clean_items = [l for l in s3_lines if not l.startswith("3.")]
    s3_raw_block = "\n".join(s3_clean_items)
    s3_text_formatted = f"{sec3_heading}\n\n" + s3_raw_block
    
    if "<table>" in s3_raw_block:
        styled_table = s3_raw_block.replace("<table>", "<table style='width:100%; border-collapse:collapse; margin:10px 0; font-size:13.5px;'>")
        styled_table = styled_table.replace("<th>", "<th style='background:#f1f5f9; padding:10px 12px; border:1px solid #cbd5e1; color:#1e3a8a; text-align:left; font-weight:bold;'>")
        styled_table = styled_table.replace("<td>", "<td style='padding:9px 12px; border:1px solid #cbd5e1; color:#334155;'>")
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n{styled_table}\n</div>"
    else:
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n<pre style='white-space:pre-wrap; font-family:inherit; font-size:13.5px;'>{s3_raw_block}</pre>\n</div>"

    # 4. Section 4 (FAQ)
    s4_clean_items = [l for l in s4_lines if not l.startswith("4.")]
    s4_text_blocks = []
    s4_html_blocks = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s4_html_blocks.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 12px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>💬 {sec4_heading}</h3>")
    
    cur_q = ""
    cur_answers = []
    for l in s4_clean_items:
        if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
            if cur_q:
                q_text = cur_q
                a_text = "\n".join(cur_answers)
                s4_text_blocks.append(f"{q_text}\n{a_text}")
                s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
                for ans in cur_answers:
                    s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
            cur_q = l.strip()
            cur_answers = []
        else:
            sentences = [s.strip() for s in re.split(r'(?<=[。！？\.\?!])\s*', l) if s.strip()]
            for s in sentences:
                cur_answers.append(s)
                
    if cur_q:
        q_text = cur_q
        a_text = "\n".join(cur_answers)
        s4_text_blocks.append(f"{q_text}\n{a_text}")
        s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
        for ans in cur_answers:
            s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
    s4_html_blocks.append("</div>")

    s4_text_formatted = f"{sec4_heading}\n\n" + "\n\n".join(s4_text_blocks)
    s4_html = "\n".join(s4_html_blocks)

    full_html_code = f"""<!-- 다국어 E-Commerce 상세페이지 4-Core 마이크로-써머리 & 비교표 & FAQ -->
<div style="font-family:'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; max-width:860px; margin:0 auto; padding:10px 0; color:#1e293b;">
{s2_html}
{s3_html}
{s4_html}
</div>"""

    html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; background: #f1f5f9; color: #0f172a; padding: 25px; margin: 0; line-height: 1.6; }}
  .container {{ max-width: 960px; margin: 0 auto; background: #ffffff; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.06); padding: 35px; border: 1px solid #cbd5e1; }}
  h1 {{ font-size: 22px; color: #0f172a; border-bottom: 2px solid #2563eb; padding-bottom: 12px; margin-top: 0; display: flex; align-items: center; gap: 8px; }}
  .guide-box {{ background: #eff6ff; border: 1px solid #bfdbfe; border-left: 5px solid #2563eb; padding: 18px 20px; border-radius: 8px; margin-bottom: 25px; font-size: 14px; color: #1e40af; line-height: 1.8; }}
  .guide-box strong {{ color: #1e3a8a; }}
  .card {{ background: #f8fafc; border: 1px solid #cbd5e1; border-radius: 10px; padding: 22px; margin-bottom: 25px; }}
  .card-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 14px; border-bottom: 1px solid #e2e8f0; padding-bottom: 10px; }}
  .card-title {{ font-size: 16px; font-weight: bold; color: #1e3a8a; }}
  .btn-group {{ display: flex; gap: 8px; }}
  .copy-btn {{ background: #2563eb; color: #ffffff; border: none; padding: 9px 16px; border-radius: 6px; cursor: pointer; font-size: 13px; font-weight: 700; transition: all 0.2s; box-shadow: 0 2px 4px rgba(37,99,235,0.2); display: inline-flex; align-items: center; gap: 6px; }}
  .copy-btn:hover {{ background: #1d4ed8; transform: translateY(-1px); }}
  .copy-btn.html-mode-btn {{ background: #059669; box-shadow: 0 2px 4px rgba(5,150,105,0.25); }}
  .copy-btn.html-mode-btn:hover {{ background: #047857; }}
  .full-btn {{ background: #7c3aed; padding: 12px 24px; font-size: 15px; width: 100%; justify-content: center; margin-bottom: 20px; box-shadow: 0 3px 6px rgba(124,58,237,0.25); }}
  .full-btn:hover {{ background: #6d28d9; }}
  .text-area {{ width: 100%; box-sizing: border-box; background: #ffffff; border: 1px solid #cbd5e1; border-radius: 6px; padding: 14px 16px; font-size: 13.5px; line-height: 1.75; color: #1e293b; font-family: 'Malgun Gothic', 'Segoe UI', monospace; resize: vertical; outline: none; }}
  .text-area:focus {{ border-color: #2563eb; box-shadow: 0 0 0 3px rgba(37,99,235,0.15); }}
  .toast {{ position: fixed; bottom: 30px; right: 30px; background: #0f172a; color: #ffffff; padding: 14px 28px; border-radius: 8px; font-size: 14px; font-weight: 600; display: none; z-index: 1000; box-shadow: 0 6px 20px rgba(0,0,0,0.25); }}
</style>
</head>
<body>

<div class="container">
  <h1>🌐 {title}</h1>
  
  <div class="guide-box">
    📢 <strong>쇼핑몰 등록 방식별 2대 원클릭 복사 기능 안내:</strong><br>
    • <strong>1. [📋 에디터에 붙여넣을 텍스트 복사] (파란색 버튼)</strong>: 지마켓/스마트스토어/쿠팡/해외몰 <strong>'에디터 작성'</strong> 화면에 붙여넣을 때 사용합니다.<br>
    • <strong>2. [🌐 HTML로 텍스트 복사] (초록색 버튼)</strong>: <strong>'HTML 작성'</strong> 탭이나 HTML 직접 입력 모드에 붙여넣을 때 사용합니다.
  </div>

  <button class="copy-btn full-btn html-mode-btn" onclick="copyFromTextarea('full-html-ta', '🎉 전체 HTML 소스코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🚀 [HTML로 전체 일괄 복사] 4-Core 요약 + 비교표 + FAQ 전체 소스코드 복사</button>
  <textarea id="full-html-ta" style="display:none;">{full_html_code}</textarea>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📌 1. 공식 상품명 (Title)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec1-ta', '✅ 상품명이 복사되었습니다!')">📋 상품명 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec1-ta" rows="2" readonly>{s1_clean}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">🔬 2. 핵심 가치 및 5줄 마이크로 요약</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec2-ta', '✅ 5줄 요약 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec2-html-ta', '🌐 5줄 요약 HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec2-ta" rows="8" readonly>{s2_text_formatted}</textarea>
    <textarea id="sec2-html-ta" style="display:none;">{s2_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📊 3. 제품 상세 스펙 비교표 (HTML Table)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec3-ta', '✅ 비교표 텍스트가 복사되었습니다!')">📋 비교표 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec3-html-ta', '🌐 비교표 HTML 코드가 복사되었습니다!')">🌐 HTML로 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec3-ta" rows="10" readonly>{s3_text_formatted}</textarea>
    <textarea id="sec3-html-ta" style="display:none;">{s3_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">💬 4. 5대 핵심 FAQ & 상세 가이드</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec4-ta', '✅ FAQ 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec4-html-ta', '🌐 FAQ HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec4-ta" rows="18" readonly>{s4_text_formatted}</textarea>
    <textarea id="sec4-html-ta" style="display:none;">{s4_html}</textarea>
  </div>

</div>

<div class="toast" id="toast">✅ 클립보드에 복사되었습니다! 쇼핑몰 에디터에 붙여넣기(Ctrl+V) 하세요.</div>

<script>
function showToast(msg) {{
  const t = document.getElementById('toast');
  t.innerText = msg;
  t.style.display = 'block';
  setTimeout(() => {{ t.style.display = 'none'; }}, 3000);
}}

function copyFromTextarea(id, customMsg) {{
  const ta = document.getElementById(id);
  const text = ta.value;
  
  const temp = document.createElement('textarea');
  temp.value = text;
  temp.style.position = 'fixed';
  temp.style.left = '-9999px';
  document.body.appendChild(temp);
  temp.select();
  temp.setSelectionRange(0, 99999);
  
  try {{
    document.execCommand('copy');
    showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
  }} catch (err) {{
    if (navigator.clipboard) {{
      navigator.clipboard.writeText(text).then(() => {{
        showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
      }});
    }}
  }} finally {{
    document.body.removeChild(temp);
  }}
}}
</script>

</body>
</html>
"""
    with open(out_html_path, "w", encoding="utf-8") as f:
        f.write(html_content)


def _generate_docx_file(title: str, text_content: str, out_docx_path: str, target_lang: str = "EN"):
    """MS Word 서식(.docx)으로 4-Core 상세페이지 완성 원고를 렌더링합니다."""
    try:
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        doc = docx.Document()
        for s in doc.sections:
            s.top_margin = Inches(0.8)
            s.bottom_margin = Inches(0.8)
            s.left_margin = Inches(0.8)
            s.right_margin = Inches(0.8)

        # Title Header
        p_title = doc.add_paragraph()
        r_t = p_title.add_run(f"🛒 {title} - E-Commerce PDP Master Copy")
        r_t.font.name = "맑은 고딕"
        r_t.font.size = Pt(16)
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(16, 44, 87)

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(14)
        r_sub = p_sub.add_run("이 문서는 각국 전자상거래 플랫폼(스마트스토어, 쿠팡, 아마존, 큐텐 등) 상세페이지 에디터에 그대로 복사하여 등록할 수 있는 100% 고객 노출용 원고입니다.")
        r_sub.font.name = "맑은 고딕"
        r_sub.font.size = Pt(9.5)
        r_sub.font.color.rgb = RGBColor(100, 116, 139)

        # Parse sections
        lines = [l.strip() for l in text_content.splitlines() if l.strip()]
        cur_sec = 0
        sec_buffers = {1: [], 2: [], 3: [], 4: []}

        for l in lines:
            if l.startswith("1."):
                cur_sec = 1
                sec_buffers[1].append(l)
            elif l.startswith("2."):
                cur_sec = 2
                sec_buffers[2].append(l)
            elif l.startswith("3."):
                cur_sec = 3
                sec_buffers[3].append(l)
            elif l.startswith("4."):
                cur_sec = 4
                sec_buffers[4].append(l)
            else:
                if cur_sec in sec_buffers:
                    sec_buffers[cur_sec].append(l)

        # Section 1: Title
        h1 = doc.add_heading(level=1)
        r_h1 = h1.add_run("1. 공식 상품명 (Official Product Title)")
        r_h1.font.name = "맑은 고딕"
        r_h1.font.size = Pt(12.5)
        r_h1.font.bold = True
        r_h1.font.color.rgb = RGBColor(37, 99, 235)

        t_val = " ".join([l for l in sec_buffers[1] if not l.startswith("1.")]).strip()
        if not t_val and sec_buffers[1]:
            t_val = sec_buffers[1][-1].replace("1.", "").strip()
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_after = Pt(10)
        r_tv = p_t.add_run(t_val)
        r_tv.font.name = "맑은 고딕"
        r_tv.font.size = Pt(10.5)
        r_tv.font.bold = True

        # Section 2: Summary
        h2 = doc.add_heading(level=1)
        r_h2 = h2.add_run("2. 핵심 가치 및 제품 안내 (Core Value & Summary)")
        r_h2.font.name = "맑은 고딕"
        r_h2.font.size = Pt(12.5)
        r_h2.font.bold = True
        r_h2.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[2]:
            if l.startswith("2."):
                continue
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(l.lstrip("- •*"))
            r.font.name = "맑은 고딕"
            r.font.size = Pt(9.5)

        # Section 3: Comparison Table
        h3 = doc.add_heading(level=1)
        r_h3 = h3.add_run("3. 제품 상세 스펙 비교 (Product Specifications & Comparison Table)")
        r_h3.font.name = "맑은 고딕"
        r_h3.font.size = Pt(12.5)
        r_h3.font.bold = True
        r_h3.font.color.rgb = RGBColor(37, 99, 235)

        raw_s3_text = "\n".join(sec_buffers[3])
        if "<table>" in raw_s3_text:
            rows_data = []
            tr_matches = re.findall(r'<tr>(.*?)</tr>', raw_s3_text, flags=re.DOTALL)
            for tr in tr_matches:
                cols = re.findall(r'<t[hd]>(.*?)</t[hd]>', tr, flags=re.DOTALL)
                if cols:
                    rows_data.append([re.sub(r'<[^>]+>', '', c).strip() for c in cols])
            
            if rows_data:
                col_cnt = max(len(r) for r in rows_data)
                tbl = doc.add_table(rows=len(rows_data), cols=col_cnt)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r_idx, r_cols in enumerate(rows_data):
                    for c_idx, c_val in enumerate(r_cols):
                        cell = tbl.rows[r_idx].cells[c_idx]
                        p = cell.paragraphs[0]
                        r = p.add_run(c_val)
                        r.font.name = "맑은 고딕"
                        r.font.size = Pt(9)
                        if r_idx == 0:
                            r.font.bold = True
                            tcPr = cell._element.get_or_add_tcPr()
                            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
                            tcPr.append(shd)
        else:
            for l in sec_buffers[3]:
                if not l.startswith("3."):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    r = p.add_run(l)
                    r.font.name = "맑은 고딕"
                    r.font.size = Pt(9.5)

        # Section 4: FAQ
        h4 = doc.add_heading(level=1)
        r_h4 = h4.add_run("4. 자주 묻는 질문 (Frequently Asked Questions)")
        r_h4.font.name = "맑은 고딕"
        r_h4.font.size = Pt(12.5)
        r_h4.font.bold = True
        r_h4.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[4]:
            if l.startswith("4."):
                continue
            if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = RGBColor(30, 58, 138)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(51, 65, 85)

        doc.save(out_docx_path)
    except Exception as de:
        print(f"  ⚠️ [DOCX 생성 폴백]: {de}", flush=True)


def _generate_web_copier_html_file(title: str, text_content: str, out_html_path: str, target_lang: str = "EN"):
    """지마켓/스마트스토어/쿠팡/해외몰 등 전 이커머스 플랫폼에 맞춘 4-Core '에디터에 붙여넣을 텍스트 복사' 및 'HTML로 텍스트 복사' 뷰어를 생성합니다."""
    
    section_titles = {
        "EN": {
            "sec2": "2. Core Value & Active Ingredient Summary",
            "sec3": "3. Product Specifications & Comparison Table",
            "sec4": "4. Product Usage Guide & Frequently Asked Questions (FAQ)"
        },
        "JP": {
            "sec2": "2. コアバリュー＆成分サイエンス要約",
            "sec3": "3. 製品仕様・他社比較テーブル",
            "sec4": "4. 使用ガイド＆よくある質問 (FAQ)"
        },
        "CN": {
            "sec2": "2. 核心价值与成分科技摘要",
            "sec3": "3. 产品规格与竞品对比表",
            "sec4": "4. 商品使用指南与常见问题解答 (FAQ)"
        },
        "TW": {
            "sec2": "2. 核心價值與成分科技摘要",
            "sec3": "3. 產品規格與競品對比表",
            "sec4": "4. 使用指南與常見問題解答 (FAQ)"
        },
        "KR": {
            "sec2": "2. 제품 핵심 안내 및 성분 요약",
            "sec3": "3. 제품 상세 비교 스펙 테이블",
            "sec4": "4. 자주 묻는 질문 (FAQ)"
        }
    }
    
    detected_lang = target_lang.upper() if target_lang else "EN"
    titles_map = section_titles.get(detected_lang, section_titles["EN"])
    sec2_heading = titles_map["sec2"]
    sec3_heading = titles_map["sec3"]
    sec4_heading = titles_map.get("sec4", "4. FAQ")

    lines = [l.strip() for l in text_content.splitlines() if l.strip()]
    
    s1_lines = []
    s2_lines = []
    s3_lines = []
    s4_lines = []
    cur_sec = 0
    for l in lines:
        if l.startswith("1."):
            cur_sec = 1
            s1_lines.append(l)
        elif l.startswith("2."):
            cur_sec = 2
            s2_lines.append(l)
        elif l.startswith("3."):
            cur_sec = 3
            s3_lines.append(l)
        elif l.startswith("4."):
            cur_sec = 4
            s4_lines.append(l)
        else:
            if cur_sec == 1:
                s1_lines.append(l)
            elif cur_sec == 2:
                s2_lines.append(l)
            elif cur_sec == 3:
                s3_lines.append(l)
            elif cur_sec == 4:
                s4_lines.append(l)

    # 1. Section 1 (Title)
    s1_clean = " ".join([l for l in s1_lines if not l.startswith("1.")]).strip()
    if not s1_clean and s1_lines:
        s1_clean = s1_lines[-1].replace("1.", "").strip()

    # 2. Section 2 (Summary)
    s2_clean_items = [l for l in s2_lines if not l.startswith("2.")]
    s2_text_formatted = f"{sec2_heading}\n\n" + "\n\n".join(s2_clean_items)
    
    s2_html_items = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s2_html_items.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>🔬 {sec2_heading}</h3>")
    for l in s2_clean_items:
        if ":" in l or "：" in l:
            k, v = re.split(r'[:：]', l, 1)
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'><strong style='color:#1e3a8a;'>{k.strip()}:</strong> {v.strip()}</p>")
        else:
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'>{l.strip()}</p>")
    s2_html_items.append("</div>")
    s2_html = "\n".join(s2_html_items)

    # 3. Section 3 (Table)
    s3_clean_items = [l for l in s3_lines if not l.startswith("3.")]
    s3_raw_block = "\n".join(s3_clean_items)
    s3_text_formatted = f"{sec3_heading}\n\n" + s3_raw_block
    
    if "<table>" in s3_raw_block:
        styled_table = s3_raw_block.replace("<table>", "<table style='width:100%; border-collapse:collapse; margin:10px 0; font-size:13.5px;'>")
        styled_table = styled_table.replace("<th>", "<th style='background:#f1f5f9; padding:10px 12px; border:1px solid #cbd5e1; color:#1e3a8a; text-align:left; font-weight:bold;'>")
        styled_table = styled_table.replace("<td>", "<td style='padding:9px 12px; border:1px solid #cbd5e1; color:#334155;'>")
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n{styled_table}\n</div>"
    else:
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n<pre style='white-space:pre-wrap; font-family:inherit; font-size:13.5px;'>{s3_raw_block}</pre>\n</div>"

    # 4. Section 4 (FAQ)
    s4_clean_items = [l for l in s4_lines if not l.startswith("4.")]
    s4_text_blocks = []
    s4_html_blocks = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s4_html_blocks.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 12px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>💬 {sec4_heading}</h3>")
    
    cur_q = ""
    cur_answers = []
    for l in s4_clean_items:
        if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
            if cur_q:
                q_text = cur_q
                a_text = "\n".join(cur_answers)
                s4_text_blocks.append(f"{q_text}\n{a_text}")
                s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
                for ans in cur_answers:
                    s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
            cur_q = l.strip()
            cur_answers = []
        else:
            sentences = [s.strip() for s in re.split(r'(?<=[。！？\.\?!])\s*', l) if s.strip()]
            for s in sentences:
                cur_answers.append(s)
                
    if cur_q:
        q_text = cur_q
        a_text = "\n".join(cur_answers)
        s4_text_blocks.append(f"{q_text}\n{a_text}")
        s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
        for ans in cur_answers:
            s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
    s4_html_blocks.append("</div>")

    s4_text_formatted = f"{sec4_heading}\n\n" + "\n\n".join(s4_text_blocks)
    s4_html = "\n".join(s4_html_blocks)

    full_html_code = f"""<!-- 다국어 E-Commerce 상세페이지 4-Core 마이크로-써머리 & 비교표 & FAQ -->
<div style="font-family:'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; max-width:860px; margin:0 auto; padding:10px 0; color:#1e293b;">
{s2_html}
{s3_html}
{s4_html}
</div>"""

    html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; background: #f1f5f9; color: #0f172a; padding: 25px; margin: 0; line-height: 1.6; }}
  .container {{ max-width: 960px; margin: 0 auto; background: #ffffff; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.06); padding: 35px; border: 1px solid #cbd5e1; }}
  h1 {{ font-size: 22px; color: #0f172a; border-bottom: 2px solid #2563eb; padding-bottom: 12px; margin-top: 0; display: flex; align-items: center; gap: 8px; }}
  .guide-box {{ background: #eff6ff; border: 1px solid #bfdbfe; border-left: 5px solid #2563eb; padding: 18px 20px; border-radius: 8px; margin-bottom: 25px; font-size: 14px; color: #1e40af; line-height: 1.8; }}
  .guide-box strong {{ color: #1e3a8a; }}
  .card {{ background: #f8fafc; border: 1px solid #cbd5e1; border-radius: 10px; padding: 22px; margin-bottom: 25px; }}
  .card-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 14px; border-bottom: 1px solid #e2e8f0; padding-bottom: 10px; }}
  .card-title {{ font-size: 16px; font-weight: bold; color: #1e3a8a; }}
  .btn-group {{ display: flex; gap: 8px; }}
  .copy-btn {{ background: #2563eb; color: #ffffff; border: none; padding: 9px 16px; border-radius: 6px; cursor: pointer; font-size: 13px; font-weight: 700; transition: all 0.2s; box-shadow: 0 2px 4px rgba(37,99,235,0.2); display: inline-flex; align-items: center; gap: 6px; }}
  .copy-btn:hover {{ background: #1d4ed8; transform: translateY(-1px); }}
  .copy-btn.html-mode-btn {{ background: #059669; box-shadow: 0 2px 4px rgba(5,150,105,0.25); }}
  .copy-btn.html-mode-btn:hover {{ background: #047857; }}
  .full-btn {{ background: #7c3aed; padding: 12px 24px; font-size: 15px; width: 100%; justify-content: center; margin-bottom: 20px; box-shadow: 0 3px 6px rgba(124,58,237,0.25); }}
  .full-btn:hover {{ background: #6d28d9; }}
  .text-area {{ width: 100%; box-sizing: border-box; background: #ffffff; border: 1px solid #cbd5e1; border-radius: 6px; padding: 14px 16px; font-size: 13.5px; line-height: 1.75; color: #1e293b; font-family: 'Malgun Gothic', 'Segoe UI', monospace; resize: vertical; outline: none; }}
  .text-area:focus {{ border-color: #2563eb; box-shadow: 0 0 0 3px rgba(37,99,235,0.15); }}
  .toast {{ position: fixed; bottom: 30px; right: 30px; background: #0f172a; color: #ffffff; padding: 14px 28px; border-radius: 8px; font-size: 14px; font-weight: 600; display: none; z-index: 1000; box-shadow: 0 6px 20px rgba(0,0,0,0.25); }}
</style>
</head>
<body>

<div class="container">
  <h1>🌐 {title}</h1>
  
  <div class="guide-box">
    📢 <strong>쇼핑몰 등록 방식별 2대 원클릭 복사 기능 안내:</strong><br>
    • <strong>1. [📋 에디터에 붙여넣을 텍스트 복사] (파란색 버튼)</strong>: 지마켓/스마트스토어/쿠팡/해외몰 <strong>'에디터 작성'</strong> 화면에 붙여넣을 때 사용합니다.<br>
    • <strong>2. [🌐 HTML로 텍스트 복사] (초록색 버튼)</strong>: <strong>'HTML 작성'</strong> 탭이나 HTML 직접 입력 모드에 붙여넣을 때 사용합니다.
  </div>

  <button class="copy-btn full-btn html-mode-btn" onclick="copyFromTextarea('full-html-ta', '🎉 전체 HTML 소스코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🚀 [HTML로 전체 일괄 복사] 4-Core 요약 + 비교표 + FAQ 전체 소스코드 복사</button>
  <textarea id="full-html-ta" style="display:none;">{full_html_code}</textarea>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📌 1. 공식 상품명 (Title)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec1-ta', '✅ 상품명이 복사되었습니다!')">📋 상품명 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec1-ta" rows="2" readonly>{s1_clean}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">🔬 2. 핵심 가치 및 5줄 마이크로 요약</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec2-ta', '✅ 5줄 요약 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec2-html-ta', '🌐 5줄 요약 HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec2-ta" rows="8" readonly>{s2_text_formatted}</textarea>
    <textarea id="sec2-html-ta" style="display:none;">{s2_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📊 3. 제품 상세 스펙 비교표 (HTML Table)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec3-ta', '✅ 비교표 텍스트가 복사되었습니다!')">📋 비교표 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec3-html-ta', '🌐 비교표 HTML 코드가 복사되었습니다!')">🌐 HTML로 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec3-ta" rows="10" readonly>{s3_text_formatted}</textarea>
    <textarea id="sec3-html-ta" style="display:none;">{s3_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">💬 4. 5대 핵심 FAQ & 상세 가이드</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec4-ta', '✅ FAQ 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec4-html-ta', '🌐 FAQ HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec4-ta" rows="18" readonly>{s4_text_formatted}</textarea>
    <textarea id="sec4-html-ta" style="display:none;">{s4_html}</textarea>
  </div>

</div>

<div class="toast" id="toast">✅ 클립보드에 복사되었습니다! 쇼핑몰 에디터에 붙여넣기(Ctrl+V) 하세요.</div>

<script>
function showToast(msg) {{
  const t = document.getElementById('toast');
  t.innerText = msg;
  t.style.display = 'block';
  setTimeout(() => {{ t.style.display = 'none'; }}, 3000);
}}

function copyFromTextarea(id, customMsg) {{
  const ta = document.getElementById(id);
  const text = ta.value;
  
  const temp = document.createElement('textarea');
  temp.value = text;
  temp.style.position = 'fixed';
  temp.style.left = '-9999px';
  document.body.appendChild(temp);
  temp.select();
  temp.setSelectionRange(0, 99999);
  
  try {{
    document.execCommand('copy');
    showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
  }} catch (err) {{
    if (navigator.clipboard) {{
      navigator.clipboard.writeText(text).then(() => {{
        showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
      }});
    }}
  }} finally {{
    document.body.removeChild(temp);
  }}
}}
</script>

</body>
</html>
"""
    with open(out_html_path, "w", encoding="utf-8") as f:
        f.write(html_content)


async def generate_seo_geo_aeo_txt(client: genai.Client, current_source_dir: str, target_dir: str, target_lang: str, product_name: str):
    """4-Core 마이크로-써머리 SEO/GEO/AEO (TXT, HTML 뷰어, DOCX, MD) 4종 파일을 자동 생성합니다. (URL/고시표 듀얼 인제스천 및 4개국 법무 렉시콘 100% 결합)"""
    print(f"\n🌐 [SEO/GEO/AEO 4-Core] 정밀 팩트 인제스천 및 4종 포맷(DOCX/HTML/TXT/MD) 생성 중 ({target_lang})...", flush=True)

    # 1. 듀얼 인제스천: url.txt 실시간 웹 스크래핑
    url_fact_context = ""
    url_file_candidates = [
        os.path.join(current_source_dir, "url.txt"),
        os.path.join(current_source_dir, "product_url.txt"),
        os.path.join(current_source_dir, "URL.txt")
    ]
    for u_path in url_file_candidates:
        if os.path.exists(u_path):
            try:
                with open(u_path, "r", encoding="utf-8") as uf:
                    raw_url = uf.read().strip()
                if raw_url.startswith("http"):
                    print(f"  🔗 [URL INGESTION 감지] {raw_url}", flush=True)
                    import urllib.request
                    req = urllib.request.Request(raw_url, headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"})
                    with urllib.request.urlopen(req, timeout=5) as resp:
                        html_bytes = resp.read()
                        raw_html = html_bytes.decode("utf-8", errors="ignore")
                        clean_text = re.sub(r'<script.*?</script>', '', raw_html, flags=re.DOTALL)
                        clean_text = re.sub(r'<style.*?</style>', '', clean_text, flags=re.DOTALL)
                        clean_text = re.sub(r'<[^>]+>', ' ', clean_text)
                        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                        url_fact_context = f"\n[Live Web Product Page Facts (Ground Truth)]\nURL: {raw_url}\nScraped Text Extract: {clean_text[:1200]}\n"
                        print(f"  ✅ [URL 팩트 스크래핑 완료] {len(clean_text)} 자 추출", flush=True)
                        break
            except Exception as ue:
                print(f"  ⚠️ [URL 스크래핑 실패 / 폴백 진행]: {ue}", flush=True)

    # 2. 이미지 팩트 앵커링 & 초월번역 컨텍스트 로드
    guide_path = os.path.join(current_source_dir, "transcreation_guide.json")
    qa_path = os.path.join(target_dir, "Transcreation_QA_Report.json")
    guide_context = ""
    if os.path.exists(guide_path):
        try:
            with open(guide_path, "r", encoding="utf-8") as f:
                gdata = json.load(f)
                guide_context = json.dumps(gdata.get("transcreation_comparisons", [])[:6], ensure_ascii=False)
        except Exception:
            pass
    elif os.path.exists(qa_path):
        try:
            with open(qa_path, "r", encoding="utf-8") as f:
                qdata = json.load(f)
                guide_context = json.dumps(qdata.get("transcreation_comparisons", [])[:6], ensure_ascii=False)
        except Exception:
            pass

    # 3. 4개국 법무 렉시콘 로드 (COMPLIANCE-FIRST)
    lexicon_rules_text = ""
    lexicon_map = {
        "EN": "en_fda_mocra_lexicon.json",
        "JP": "jp_pmda_pharm_lexicon.json",
        "CN": "cn_nmpa_adlaw_lexicon.json",
        "TW": "tw_tfda_lexicon.json"
    }
    lex_file = lexicon_map.get(target_lang)
    if lex_file:
        lex_path = os.path.join(PROJECT_ROOT, "00_공통자료", "compliance_lexicons", lex_file)
        if os.path.exists(lex_path):
            try:
                with open(lex_path, "r", encoding="utf-8") as lf:
                    lex_data = json.load(lf)
                    lexicon_rules_text = f"\n[MANDATORY COMPLIANCE LEXICON ({lex_data.get('jurisdiction', '')})]\n"
                    cats = lex_data.get("categories", {})
                    for c_name, c_val in cats.items():
                        banned = c_val.get("banned_terms", [])
                        for b in banned[:5]:
                            lexicon_rules_text += f"- Banned: '{b.get('banned')}' -> Must use: '{b.get('preferred')}' ({b.get('reason')})\n"
            except Exception:
                pass

    lang_names = {
        "EN": "English for Amazon / Sephora US",
        "JP": "Japanese for Qoo10 Japan / Cosme",
        "CN": "Simplified Chinese for Tmall / Xiaohongshu",
        "TW": "Traditional Chinese for Shopee Taiwan / Momo",
        "KR": "Korean for Naver Smartstore / Coupang"
    }
    target_lang_desc = lang_names.get(target_lang, "English")

    prompt = f"""[SYSTEM PROMPT] Global E-Commerce SEO/GEO/AEO 4-Core Master Copy Generator
Product Name: {product_name}
Target Market & Language: {target_lang_desc}

{url_fact_context}
[Verified Transcreation Context & Ingredients Data]:
{guide_context}

{lexicon_rules_text}

[CRITICAL INSTRUCTION - ZERO META COMMENTARY]
Never output developer metadata, markdown headers '##', explanation notes, character counters, or words like 'GEO', 'AEO', 'RAG'.
Output purely customer-facing content structured in 4 distinct sections.

Strict 4-Core Structure:
1. Official Product Title
(Under 100 characters. Noun Phrase: Brand 'Logicall Skin' + Product Title + Key Efficacy + Volume)

2. Core Value & Active Ingredient Summary
(5 concise bullet points containing quantitative metrics e.g. ppm, %, non-irritation score 0.00):
- Brand: Logicall Skin
- Core Actives & Concentration: (e.g. Multi-Vitamin 10% / 100,000ppm, Aquatide 3%)
- Key Benefits: (Efficacy claims strictly compliant with target country cosmetics law)
- Texture & Absorption: (Hydra-watery, non-greasy, fast-absorbing)
- Skin Compatibility: (Dermatologist tested, 0.00 irritation index)

3. Product Specifications & Comparison Table
(Output a clean HTML <table> comparing Logicall Skin vs Generic Market Standard):
<table>
  <tr><th>Dimensions</th><th>Logicall Skin</th><th>Standard Market Benchmark</th></tr>
  <tr><td>Active Concentration</td><td>High-Potency Multi-Vitamin 100,000ppm (10%)</td><td>Diluted extract 1,000~5,000ppm</td></tr>
  <tr><td>Patented Science</td><td>Aquatide 5000 (3%)</td><td>Generic purified water base</td></tr>
  <tr><td>Formula Stability</td><td>High-stability oxidation-free formula</td><td>Prone to discoloration / oxidation</td></tr>
  <tr><td>Irritation Index</td><td>0.00 Low-Irritation Certified</td><td>May cause stinging or redness</td></tr>
</table>

4. Product Usage Guide & Frequently Asked Questions (FAQ)
(5 high-conversion B2C customer FAQs):
Q1: When should I apply this serum?
A: ...
Q2: Is it suitable for sensitive skin?
A: ...
Q3: How does the active formula benefit the skin?
A: ...
Q4: What is the main efficacy of the Multi-Vitamin complex?
A: ...
Q5: Can I layer this with other skincare products?
A: ...

Generate the complete 4-Core content in {target_lang_desc} now.
"""
    try:
        resp = await client.aio.models.generate_content(
            model=MODEL_PRO,
            contents=[prompt],
            config=types.GenerateContentConfig(
                temperature=0.6,
                top_p=0.9,
                max_output_tokens=8192
            )
        )
        content_text = resp.text.strip()
    except Exception as e:
        print(f"  ⚠️ [WARN] SEO 텍스트 생성 중 오류 발생 -> 기본 템플릿 대체: {e}")
        content_text = f"""1. Logicall Skin {product_name} Multi Vitamin Daily Care Serum 50ml

2. Core Value & Active Ingredient Summary
- Brand: Logicall Skin
- Core Actives & Concentration: High-Potency Multi-Vitamin 10% (100,000ppm) & Aquatide 5000 (3%)
- Key Benefits: Visibly brightens, refines texture, and reinforces natural moisture barrier
- Texture & Absorption: Refreshingly lightweight hydra-watery formula with instant absorption
- Skin Compatibility: Dermatologist-tested, 0.00 skin irritation index suitable for daily use

3. Product Specifications & Comparison Table
<table>
  <tr><th>Specification</th><th>Logicall Skin Multi-Vitamin Serum</th><th>Standard Vitamin Serum</th></tr>
  <tr><td>Active Concentration</td><td>Multi-Vitamin Complex 100,000ppm (10.0%)</td><td>1,000 ~ 5,000ppm</td></tr>
  <tr><td>Patented Technology</td><td>Aquatide 5000 30,000ppm (3.0%)</td><td>Purified water base</td></tr>
  <tr><td>Formula Stability</td><td>High stability against air and light oxidation</td><td>Vulnerable to discoloration</td></tr>
  <tr><td>Skin Irritation Score</td><td>0.00 (Certified Low-Irritation)</td><td>May cause stinging sensation</td></tr>
</table>

4. Product Usage Guide & Frequently Asked Questions (FAQ)
Q1: When should I apply this serum?
A: Apply 3-4 drops evenly morning and evening after cleansing and toner.
Q2: Is it suitable for sensitive skin?
A: Yes, it is dermatologist-tested with a 0.00 skin irritation index.
Q3: How does Aquatide benefit the skin?
A: It reinforces the natural moisture barrier and revitalizes skin appearance.
Q4: What is the main efficacy of the Multi-Vitamin complex?
A: It provides deep hydration for a resilient-looking complexion and combats the signs of premature aging.
Q5: Can I layer this with other skincare products?
A: Yes, its fast-absorbing texture layers smoothly under creams and sunscreens.
"""

    # 4. 결정론적 법무 후처리 게이트 통과 (apply_deterministic_qa_overrides)
    if target_lang == "EN":
        for b_pat, p_val in [
            (r"\bnutrients for cellular vitality\b", "hydration for a resilient-looking complexion"),
            (r"\bcellular vitality\b", "resilient-looking complexion"),
            (r"\breinforces cellular resilience\b", "reinforces the skin's natural moisture barrier"),
            (r"\bcellular resilience\b", "skin's natural moisture barrier"),
            (r"\bcellular metabolism\b", "natural skin vitality"),
            (r"\bcombats premature aging\b", "combats the signs of premature aging"),
            (r"\bComplex skin issues\b", "Multiple skin concerns"),
            (r"\bTroubled skin\b", "Blemish-prone skin"),
            (r"\bcellular autophagy\b", "targeted skin nourishment")
        ]:
            content_text = re.sub(b_pat, p_val, content_text, flags=re.IGNORECASE)

    # 5. 4종 멀티 포맷 일괄 익스포트 (TXT, HTML, DOCX, MD)
    txt_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO.txt"
    html_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO_VIEWER.html"
    docx_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO.docx"
    md_filename = f"{product_name}_{target_lang}_SEO_GEO_AEO.md"

    txt_path = os.path.join(target_dir, txt_filename)
    html_path = os.path.join(target_dir, html_filename)
    docx_path = os.path.join(target_dir, docx_filename)
    md_path = os.path.join(target_dir, md_filename)

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(content_text)
    print(f"  📄 [TXT 저장 완료]: {txt_path}")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(f"# {product_name} ({target_lang}) SEO/GEO/AEO PDP 원고\n\n" + content_text)
    print(f"  📑 [MD 저장 완료]: {md_path}")

    _generate_web_copier_html_file(f"{product_name} ({target_lang})", content_text, html_path, target_lang=target_lang)
    print(f"  🌐 [HTML 뷰어 저장 완료]: {html_path}")

    _generate_docx_file(f"{product_name} ({target_lang})", content_text, docx_path, target_lang=target_lang)
    print(f"  📄 [DOCX 서식 문서 저장 완료]: {docx_path}")


def _generate_docx_file(title: str, text_content: str, out_docx_path: str, target_lang: str = "EN"):
    """MS Word 서식(.docx)으로 4-Core 상세페이지 완성 원고를 렌더링합니다."""
    try:
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        doc = docx.Document()
        for s in doc.sections:
            s.top_margin = Inches(0.8)
            s.bottom_margin = Inches(0.8)
            s.left_margin = Inches(0.8)
            s.right_margin = Inches(0.8)

        # Title Header
        p_title = doc.add_paragraph()
        r_t = p_title.add_run(f"🛒 {title} - E-Commerce PDP Master Copy")
        r_t.font.name = "맑은 고딕"
        r_t.font.size = Pt(16)
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(16, 44, 87)

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(14)
        r_sub = p_sub.add_run("이 문서는 각국 전자상거래 플랫폼(스마트스토어, 쿠팡, 아마존, 큐텐 등) 상세페이지 에디터에 그대로 복사하여 등록할 수 있는 100% 고객 노출용 원고입니다.")
        r_sub.font.name = "맑은 고딕"
        r_sub.font.size = Pt(9.5)
        r_sub.font.color.rgb = RGBColor(100, 116, 139)

        # Parse sections
        lines = [l.strip() for l in text_content.splitlines() if l.strip()]
        cur_sec = 0
        sec_buffers = {1: [], 2: [], 3: [], 4: []}

        for l in lines:
            if l.startswith("1."):
                cur_sec = 1
                sec_buffers[1].append(l)
            elif l.startswith("2."):
                cur_sec = 2
                sec_buffers[2].append(l)
            elif l.startswith("3."):
                cur_sec = 3
                sec_buffers[3].append(l)
            elif l.startswith("4."):
                cur_sec = 4
                sec_buffers[4].append(l)
            else:
                if cur_sec in sec_buffers:
                    sec_buffers[cur_sec].append(l)

        # Section 1: Title
        h1 = doc.add_heading(level=1)
        r_h1 = h1.add_run("1. 공식 상품명 (Official Product Title)")
        r_h1.font.name = "맑은 고딕"
        r_h1.font.size = Pt(12.5)
        r_h1.font.bold = True
        r_h1.font.color.rgb = RGBColor(37, 99, 235)

        t_val = " ".join([l for l in sec_buffers[1] if not l.startswith("1.")]).strip()
        if not t_val and sec_buffers[1]:
            t_val = sec_buffers[1][-1].replace("1.", "").strip()
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_after = Pt(10)
        r_tv = p_t.add_run(t_val)
        r_tv.font.name = "맑은 고딕"
        r_tv.font.size = Pt(10.5)
        r_tv.font.bold = True

        # Section 2: Summary
        h2 = doc.add_heading(level=1)
        r_h2 = h2.add_run("2. 핵심 가치 및 제품 안내 (Core Value & Summary)")
        r_h2.font.name = "맑은 고딕"
        r_h2.font.size = Pt(12.5)
        r_h2.font.bold = True
        r_h2.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[2]:
            if l.startswith("2."):
                continue
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(l.lstrip("- •*"))
            r.font.name = "맑은 고딕"
            r.font.size = Pt(9.5)

        # Section 3: Comparison Table
        h3 = doc.add_heading(level=1)
        r_h3 = h3.add_run("3. 제품 상세 스펙 비교 (Product Specifications & Comparison Table)")
        r_h3.font.name = "맑은 고딕"
        r_h3.font.size = Pt(12.5)
        r_h3.font.bold = True
        r_h3.font.color.rgb = RGBColor(37, 99, 235)

        raw_s3_text = "\n".join(sec_buffers[3])
        if "<table>" in raw_s3_text:
            rows_data = []
            tr_matches = re.findall(r'<tr>(.*?)</tr>', raw_s3_text, flags=re.DOTALL)
            for tr in tr_matches:
                cols = re.findall(r'<t[hd]>(.*?)</t[hd]>', tr, flags=re.DOTALL)
                if cols:
                    rows_data.append([re.sub(r'<[^>]+>', '', c).strip() for c in cols])
            
            if rows_data:
                col_cnt = max(len(r) for r in rows_data)
                tbl = doc.add_table(rows=len(rows_data), cols=col_cnt)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r_idx, r_cols in enumerate(rows_data):
                    for c_idx, c_val in enumerate(r_cols):
                        cell = tbl.rows[r_idx].cells[c_idx]
                        p = cell.paragraphs[0]
                        r = p.add_run(c_val)
                        r.font.name = "맑은 고딕"
                        r.font.size = Pt(9)
                        if r_idx == 0:
                            r.font.bold = True
                            tcPr = cell._element.get_or_add_tcPr()
                            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
                            tcPr.append(shd)
        else:
            for l in sec_buffers[3]:
                if not l.startswith("3."):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    r = p.add_run(l)
                    r.font.name = "맑은 고딕"
                    r.font.size = Pt(9.5)

        # Section 4: FAQ
        h4 = doc.add_heading(level=1)
        r_h4 = h4.add_run("4. 자주 묻는 질문 (Frequently Asked Questions)")
        r_h4.font.name = "맑은 고딕"
        r_h4.font.size = Pt(12.5)
        r_h4.font.bold = True
        r_h4.font.color.rgb = RGBColor(37, 99, 235)

        for l in sec_buffers[4]:
            if l.startswith("4."):
                continue
            if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = RGBColor(30, 58, 138)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(l)
                r.font.name = "맑은 고딕"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(51, 65, 85)

        doc.save(out_docx_path)
    except Exception as de:
        print(f"  ⚠️ [DOCX 생성 폴백]: {de}", flush=True)


def _generate_web_copier_html_file(title: str, text_content: str, out_html_path: str, target_lang: str = "EN"):
    """지마켓/스마트스토어/쿠팡/해외몰 등 전 이커머스 플랫폼에 맞춘 4-Core '에디터에 붙여넣을 텍스트 복사' 및 'HTML로 텍스트 복사' 뷰어를 생성합니다."""
    
    section_titles = {
        "EN": {
            "sec2": "2. Core Value & Active Ingredient Summary",
            "sec3": "3. Product Specifications & Comparison Table",
            "sec4": "4. Product Usage Guide & Frequently Asked Questions (FAQ)"
        },
        "JP": {
            "sec2": "2. コアバリュー＆成分サイエンス要約",
            "sec3": "3. 製品仕様・他社比較テーブル",
            "sec4": "4. 使用ガイド＆よくある質問 (FAQ)"
        },
        "CN": {
            "sec2": "2. 核心价值与成分科技摘要",
            "sec3": "3. 产品规格与竞品对比表",
            "sec4": "4. 商品使用指南与常见问题解答 (FAQ)"
        },
        "TW": {
            "sec2": "2. 核心價值與成分科技摘要",
            "sec3": "3. 產品規格與競品對比表",
            "sec4": "4. 使用指南與常見問題解答 (FAQ)"
        },
        "KR": {
            "sec2": "2. 제품 핵심 안내 및 성분 요약",
            "sec3": "3. 제품 상세 비교 스펙 테이블",
            "sec4": "4. 자주 묻는 질문 (FAQ)"
        }
    }
    
    detected_lang = target_lang.upper() if target_lang else "EN"
    titles_map = section_titles.get(detected_lang, section_titles["EN"])
    sec2_heading = titles_map["sec2"]
    sec3_heading = titles_map["sec3"]
    sec4_heading = titles_map.get("sec4", "4. FAQ")

    lines = [l.strip() for l in text_content.splitlines() if l.strip()]
    
    s1_lines = []
    s2_lines = []
    s3_lines = []
    s4_lines = []
    cur_sec = 0
    for l in lines:
        if l.startswith("1."):
            cur_sec = 1
            s1_lines.append(l)
        elif l.startswith("2."):
            cur_sec = 2
            s2_lines.append(l)
        elif l.startswith("3."):
            cur_sec = 3
            s3_lines.append(l)
        elif l.startswith("4."):
            cur_sec = 4
            s4_lines.append(l)
        else:
            if cur_sec == 1:
                s1_lines.append(l)
            elif cur_sec == 2:
                s2_lines.append(l)
            elif cur_sec == 3:
                s3_lines.append(l)
            elif cur_sec == 4:
                s4_lines.append(l)

    # 1. Section 1 (Title)
    s1_clean = " ".join([l for l in s1_lines if not l.startswith("1.")]).strip()
    if not s1_clean and s1_lines:
        s1_clean = s1_lines[-1].replace("1.", "").strip()

    # 2. Section 2 (Summary)
    s2_clean_items = [l for l in s2_lines if not l.startswith("2.")]
    s2_text_formatted = f"{sec2_heading}\n\n" + "\n\n".join(s2_clean_items)
    
    s2_html_items = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s2_html_items.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>🔬 {sec2_heading}</h3>")
    for l in s2_clean_items:
        if ":" in l or "：" in l:
            k, v = re.split(r'[:：]', l, 1)
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'><strong style='color:#1e3a8a;'>{k.strip()}:</strong> {v.strip()}</p>")
        else:
            s2_html_items.append(f"<p style='margin:6px 0; font-size:14px; color:#334155; line-height:1.6;'>{l.strip()}</p>")
    s2_html_items.append("</div>")
    s2_html = "\n".join(s2_html_items)

    # 3. Section 3 (Table)
    s3_clean_items = [l for l in s3_lines if not l.startswith("3.")]
    s3_raw_block = "\n".join(s3_clean_items)
    s3_text_formatted = f"{sec3_heading}\n\n" + s3_raw_block
    
    if "<table>" in s3_raw_block:
        styled_table = s3_raw_block.replace("<table>", "<table style='width:100%; border-collapse:collapse; margin:10px 0; font-size:13.5px;'>")
        styled_table = styled_table.replace("<th>", "<th style='background:#f1f5f9; padding:10px 12px; border:1px solid #cbd5e1; color:#1e3a8a; text-align:left; font-weight:bold;'>")
        styled_table = styled_table.replace("<td>", "<td style='padding:9px 12px; border:1px solid #cbd5e1; color:#334155;'>")
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n{styled_table}\n</div>"
    else:
        s3_html = f"<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>\n<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 10px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>📊 {sec3_heading}</h3>\n<pre style='white-space:pre-wrap; font-family:inherit; font-size:13.5px;'>{s3_raw_block}</pre>\n</div>"

    # 4. Section 4 (FAQ)
    s4_clean_items = [l for l in s4_lines if not l.startswith("4.")]
    s4_text_blocks = []
    s4_html_blocks = ["<div style='margin-bottom:20px; padding:16px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;'>"]
    s4_html_blocks.append(f"<h3 style='font-size:15px; color:#1e3a8a; margin:0 0 12px 0; border-bottom:1px solid #cbd5e1; padding-bottom:6px;'>💬 {sec4_heading}</h3>")
    
    cur_q = ""
    cur_answers = []
    for l in s4_clean_items:
        if l.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q.")):
            if cur_q:
                q_text = cur_q
                a_text = "\n".join(cur_answers)
                s4_text_blocks.append(f"{q_text}\n{a_text}")
                s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
                for ans in cur_answers:
                    s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
            cur_q = l.strip()
            cur_answers = []
        else:
            sentences = [s.strip() for s in re.split(r'(?<=[。！？\.\?!])\s*', l) if s.strip()]
            for s in sentences:
                cur_answers.append(s)
                
    if cur_q:
        q_text = cur_q
        a_text = "\n".join(cur_answers)
        s4_text_blocks.append(f"{q_text}\n{a_text}")
        s4_html_blocks.append(f"<p style='font-weight:bold; font-size:14.5px; color:#1e3a8a; margin:14px 0 4px 0;'>{cur_q}</p>")
        for ans in cur_answers:
            s4_html_blocks.append(f"<p style='font-size:14px; color:#475569; margin:2px 0 4px 0; line-height:1.65;'>{ans}</p>")
    s4_html_blocks.append("</div>")

    s4_text_formatted = f"{sec4_heading}\n\n" + "\n\n".join(s4_text_blocks)
    s4_html = "\n".join(s4_html_blocks)

    full_html_code = f"""<!-- 다국어 E-Commerce 상세페이지 4-Core 마이크로-써머리 & 비교표 & FAQ -->
<div style="font-family:'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; max-width:860px; margin:0 auto; padding:10px 0; color:#1e293b;">
{s2_html}
{s3_html}
{s4_html}
</div>"""

    html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Malgun Gothic', 'Segoe UI', Roboto, sans-serif; background: #f1f5f9; color: #0f172a; padding: 25px; margin: 0; line-height: 1.6; }}
  .container {{ max-width: 960px; margin: 0 auto; background: #ffffff; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.06); padding: 35px; border: 1px solid #cbd5e1; }}
  h1 {{ font-size: 22px; color: #0f172a; border-bottom: 2px solid #2563eb; padding-bottom: 12px; margin-top: 0; display: flex; align-items: center; gap: 8px; }}
  .guide-box {{ background: #eff6ff; border: 1px solid #bfdbfe; border-left: 5px solid #2563eb; padding: 18px 20px; border-radius: 8px; margin-bottom: 25px; font-size: 14px; color: #1e40af; line-height: 1.8; }}
  .guide-box strong {{ color: #1e3a8a; }}
  .card {{ background: #f8fafc; border: 1px solid #cbd5e1; border-radius: 10px; padding: 22px; margin-bottom: 25px; }}
  .card-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 14px; border-bottom: 1px solid #e2e8f0; padding-bottom: 10px; }}
  .card-title {{ font-size: 16px; font-weight: bold; color: #1e3a8a; }}
  .btn-group {{ display: flex; gap: 8px; }}
  .copy-btn {{ background: #2563eb; color: #ffffff; border: none; padding: 9px 16px; border-radius: 6px; cursor: pointer; font-size: 13px; font-weight: 700; transition: all 0.2s; box-shadow: 0 2px 4px rgba(37,99,235,0.2); display: inline-flex; align-items: center; gap: 6px; }}
  .copy-btn:hover {{ background: #1d4ed8; transform: translateY(-1px); }}
  .copy-btn.html-mode-btn {{ background: #059669; box-shadow: 0 2px 4px rgba(5,150,105,0.25); }}
  .copy-btn.html-mode-btn:hover {{ background: #047857; }}
  .full-btn {{ background: #7c3aed; padding: 12px 24px; font-size: 15px; width: 100%; justify-content: center; margin-bottom: 20px; box-shadow: 0 3px 6px rgba(124,58,237,0.25); }}
  .full-btn:hover {{ background: #6d28d9; }}
  .text-area {{ width: 100%; box-sizing: border-box; background: #ffffff; border: 1px solid #cbd5e1; border-radius: 6px; padding: 14px 16px; font-size: 13.5px; line-height: 1.75; color: #1e293b; font-family: 'Malgun Gothic', 'Segoe UI', monospace; resize: vertical; outline: none; }}
  .text-area:focus {{ border-color: #2563eb; box-shadow: 0 0 0 3px rgba(37,99,235,0.15); }}
  .toast {{ position: fixed; bottom: 30px; right: 30px; background: #0f172a; color: #ffffff; padding: 14px 28px; border-radius: 8px; font-size: 14px; font-weight: 600; display: none; z-index: 1000; box-shadow: 0 6px 20px rgba(0,0,0,0.25); }}
</style>
</head>
<body>

<div class="container">
  <h1>🌐 {title}</h1>
  
  <div class="guide-box">
    📢 <strong>쇼핑몰 등록 방식별 2대 원클릭 복사 기능 안내:</strong><br>
    • <strong>1. [📋 에디터에 붙여넣을 텍스트 복사] (파란색 버튼)</strong>: 지마켓/스마트스토어/쿠팡/해외몰 <strong>'에디터 작성'</strong> 화면에 붙여넣을 때 사용합니다.<br>
    • <strong>2. [🌐 HTML로 텍스트 복사] (초록색 버튼)</strong>: <strong>'HTML 작성'</strong> 탭이나 HTML 직접 입력 모드에 붙여넣을 때 사용합니다.
  </div>

  <button class="copy-btn full-btn html-mode-btn" onclick="copyFromTextarea('full-html-ta', '🎉 전체 HTML 소스코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🚀 [HTML로 전체 일괄 복사] 4-Core 요약 + 비교표 + FAQ 전체 소스코드 복사</button>
  <textarea id="full-html-ta" style="display:none;">{full_html_code}</textarea>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📌 1. 공식 상품명 (Title)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec1-ta', '✅ 상품명이 복사되었습니다!')">📋 상품명 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec1-ta" rows="2" readonly>{s1_clean}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">🔬 2. 핵심 가치 및 5줄 마이크로 요약</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec2-ta', '✅ 5줄 요약 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec2-html-ta', '🌐 5줄 요약 HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec2-ta" rows="8" readonly>{s2_text_formatted}</textarea>
    <textarea id="sec2-html-ta" style="display:none;">{s2_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">📊 3. 제품 상세 스펙 비교표 (HTML Table)</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec3-ta', '✅ 비교표 텍스트가 복사되었습니다!')">📋 비교표 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec3-html-ta', '🌐 비교표 HTML 코드가 복사되었습니다!')">🌐 HTML로 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec3-ta" rows="10" readonly>{s3_text_formatted}</textarea>
    <textarea id="sec3-html-ta" style="display:none;">{s3_html}</textarea>
  </div>

  <div class="card">
    <div class="card-header">
      <div class="card-title">💬 4. 5대 핵심 FAQ & 상세 가이드</div>
      <div class="btn-group">
        <button class="copy-btn" onclick="copyFromTextarea('sec4-ta', '✅ FAQ 텍스트가 복사되었습니다! 에디터에 붙여넣기(Ctrl+V) 하세요.')">📋 에디터에 붙여넣을 텍스트 복사</button>
        <button class="copy-btn html-mode-btn" onclick="copyFromTextarea('sec4-html-ta', '🌐 FAQ HTML 코드가 복사되었습니다! [HTML 작성] 탭에 붙여넣기 하세요.')">🌐 HTML로 텍스트 복사</button>
      </div>
    </div>
    <textarea class="text-area" id="sec4-ta" rows="18" readonly>{s4_text_formatted}</textarea>
    <textarea id="sec4-html-ta" style="display:none;">{s4_html}</textarea>
  </div>

</div>

<div class="toast" id="toast">✅ 클립보드에 복사되었습니다! 쇼핑몰 에디터에 붙여넣기(Ctrl+V) 하세요.</div>

<script>
function showToast(msg) {{
  const t = document.getElementById('toast');
  t.innerText = msg;
  t.style.display = 'block';
  setTimeout(() => {{ t.style.display = 'none'; }}, 3000);
}}

function copyFromTextarea(id, customMsg) {{
  const ta = document.getElementById(id);
  const text = ta.value;
  
  const temp = document.createElement('textarea');
  temp.value = text;
  temp.style.position = 'fixed';
  temp.style.left = '-9999px';
  document.body.appendChild(temp);
  temp.select();
  temp.setSelectionRange(0, 99999);
  
  try {{
    document.execCommand('copy');
    showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
  }} catch (err) {{
    if (navigator.clipboard) {{
      navigator.clipboard.writeText(text).then(() => {{
        showToast(customMsg || '✅ 클립보드에 완벽하게 복사되었습니다! (Ctrl+V)');
      }});
    }}
  }} finally {{
    document.body.removeChild(temp);
  }}
}}
</script>

</body>
</html>
"""
    with open(out_html_path, "w", encoding="utf-8") as f:
        f.write(html_content)


if __name__ == '__main__':
    asyncio.run(main_async())


